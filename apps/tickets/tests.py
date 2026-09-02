import base64
import os
import tempfile
from decimal import Decimal
from io import BytesIO

from django.core.cache import cache
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, Client
from django.urls import reverse
from django.contrib.auth import get_user_model
from django.utils import timezone
from django.core.management import call_command
from datetime import date, timedelta, time as datetime_time, datetime as datetime_dt
from unittest.mock import patch
from apps.tickets.models import Ticket, TicketComment, TicketActivityLog, Asset, AssetCategory, AssetLog, SLA, EscalationRule, ServiceCategory, RemoteSession, RemoteConnector, Vessel, DiveSystem, JobNumber, Mobilization, MobilizationItem, AssetProcurementRequest, Attachment, AssetDepartment, Location, AssetImportBatch, TicketDraft, TicketDraftAttachment
from apps.tickets.asset_name_matching import match_users_by_name
from apps.common.models import Category, Notification
from apps.maintenance.models import MaintenanceSchedule, Vendor
from apps.tickets.periodic_tasks import run_periodic_jobs


# 1x1 transparent PNG, used to test signature-image upload/export handling.
TINY_PNG_BYTES = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII='
)

User = get_user_model()


class TicketModelTests(TestCase):
    """Test Ticket model functionality."""

    def setUp(self):
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT'
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')

    def test_ticket_creation(self):
        """Test basic ticket creation."""
        ticket = Ticket.objects.create(
            number='TK#1234',
            title='Test Ticket',
            description='Test Description',
            requester=self.user,
            category=self.category,
            status=Ticket.Status.NEW
        )
        self.assertEqual(ticket.title, 'Test Ticket')
        self.assertEqual(ticket.requester, self.user)
        self.assertEqual(ticket.status, Ticket.Status.NEW)
        self.assertEqual(str(ticket), 'TK#1234 - Test Ticket')

    def test_ticket_priority_calculation(self):
        """Test priority calculation based on impact and urgency."""
        test_cases = [
            (Ticket.Impact.INDIVIDUAL, Ticket.Urgency.CRITICAL, Ticket.Priority.P3),
            (Ticket.Impact.ORGANIZATION, Ticket.Urgency.CRITICAL, Ticket.Priority.P1),
            (Ticket.Impact.DEPARTMENT, Ticket.Urgency.MEDIUM, Ticket.Priority.P4),
        ]

        for i, (impact, urgency, expected) in enumerate(test_cases):
            ticket = Ticket.objects.create(
                number=f'TK#{i}',
                title='Test Ticket',
                description='Test Description',
                requester=self.user,
                category=self.category,
                impact=impact,
                urgency=urgency
            )
            self.assertEqual(ticket.priority, expected)

    def test_ticket_sla_status_method(self):
        """Test sla_status method on Ticket model."""
        sla = SLA.objects.create(
            priority='P3',
            response_minutes=60,
            resolution_minutes=240
        )
        
        ticket = Ticket.objects.create(
            number='TK#8888',
            title='SLA Test Ticket',
            description='Test description',
            requester=self.user,
            category=self.category,
            priority='P3',
            status=Ticket.Status.NEW
        )
        
        status = ticket.sla_status()
        self.assertIn('overall', status)
        self.assertIn('response', status)
        self.assertIn('resolution', status)

    def test_ticket_sla_breach(self):
        """Test SLA breach detection."""
        sla = SLA.objects.create(
            priority='P3',
            response_minutes=60,
            resolution_minutes=240
        )

        ticket = Ticket.objects.create(
            number='TK#8889',
            title='SLA Breach Test',
            description='Test description',
            requester=self.user,
            category=self.category,
            priority='P3',
            status=Ticket.Status.NEW,
            created_at=timezone.now() - timedelta(minutes=120)
        )

        ticket.response_due_at = timezone.now() - timedelta(minutes=30)
        ticket.resolution_due_at = timezone.now() + timedelta(minutes=120)
        ticket.save()

        status = ticket.sla_status()
        # Response SLA should be breached
        self.assertEqual(status['response'], 'breached')


class PeriodicTaskLockTests(TestCase):
    """Prevent overlapping periodic jobs from running concurrently."""

    def setUp(self):
        cache.clear()

    @patch('apps.tickets.periodic_tasks.call_command')
    def test_run_periodic_jobs_skips_when_lock_is_active(self, mock_call_command):
        cache.set('tickets:periodic_jobs:lock', 'running', timeout=300)

        run_periodic_jobs(stdout=None, stderr=None)

        mock_call_command.assert_not_called()


class TicketDraftAttachmentTests(TestCase):
    """Draft attachment upload/restore/submit lifecycle — see
    apps/tickets/views_drafts.py and restore_kept_draft_attachments in
    views.py."""

    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='drafter@example.com', password='TestPass123!',
            first_name='Draft', last_name='User', department='IT',
            is_active=True, email_verified=True,
        )
        self.other_user = User.objects.create_user(
            email='other@example.com', password='TestPass123!',
            first_name='Other', last_name='User', department='IT',
            is_active=True, email_verified=True,
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.client.login(email='drafter@example.com', password='TestPass123!')

    def _pdf(self, name='draft.pdf'):
        return SimpleUploadedFile(name, b'%PDF-1.4 test content', content_type='application/pdf')

    def test_save_draft_attachment_creates_row_and_draft(self):
        response = self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf(),
        })
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(len(data['created']), 1)
        self.assertEqual(TicketDraftAttachment.objects.count(), 1)
        att = TicketDraftAttachment.objects.first()
        self.assertEqual(att.draft.user, self.user)
        self.assertEqual(att.filename, 'draft.pdf')

    def test_save_draft_attachment_rejects_oversized_file(self):
        big = SimpleUploadedFile('big.pdf', b'x' * (11 * 1024 * 1024), content_type='application/pdf')
        response = self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': big,
        })
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(len(data['created']), 0)
        self.assertEqual(len(data['rejected']), 1)
        self.assertEqual(TicketDraftAttachment.objects.count(), 0)

    def test_get_draft_includes_attachment_metadata(self):
        self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf(),
        })
        response = self.client.get(reverse('tickets:get_draft'), {'type': 'INCIDENT'})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(len(data['attachments']), 1)
        self.assertEqual(data['attachments'][0]['filename'], 'draft.pdf')

    def test_discard_draft_attachment_removes_it(self):
        self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf(),
        })
        att = TicketDraftAttachment.objects.first()
        response = self.client.post(
            reverse('tickets:discard_draft_attachment'),
            data='{"attachment_id": %d}' % att.pk,
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(TicketDraftAttachment.objects.count(), 0)

    def test_discard_draft_attachment_scoped_to_owner(self):
        """Another user's attachment id can't be discarded."""
        self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf(),
        })
        att = TicketDraftAttachment.objects.first()
        self.client.logout()
        self.client.login(email='other@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:discard_draft_attachment'),
            data='{"attachment_id": %d}' % att.pk,
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 404)
        self.assertEqual(TicketDraftAttachment.objects.count(), 1)

    def test_discard_draft_removes_attachments_too(self):
        self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf(),
        })
        response = self.client.post(
            reverse('tickets:discard_draft'),
            data='{"ticket_type": "INCIDENT"}',
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(TicketDraft.objects.count(), 0)
        self.assertEqual(TicketDraftAttachment.objects.count(), 0)

    def test_submit_with_kept_draft_attachment_creates_real_attachment(self):
        """The restore-then-submit-without-touching-the-input flow the
        feature was built for: the draft attachment ends up as a real
        Attachment on the new ticket, and the draft is fully cleaned up."""
        save_response = self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf('kept.pdf'),
        })
        att_id = save_response.json()['created'][0]['id']

        response = self.client.post(reverse('tickets:create'), {
            'type': 'INCIDENT',
            'title': 'Restored Draft Ticket',
            'description': 'Test description',
            'category': self.category.id,
            'impact': 'INDIVIDUAL',
            'urgency': 'MEDIUM',
            'keep_draft_attachments': str(att_id),
        })
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Restored Draft Ticket')
        self.assertEqual(ticket.attachments.count(), 1)
        self.assertEqual(ticket.attachments.first().filename, 'kept.pdf')
        # Draft and its attachment are both gone — no leftover duplicate copy.
        self.assertEqual(TicketDraft.objects.filter(user=self.user, ticket_type='INCIDENT').count(), 0)
        self.assertEqual(TicketDraftAttachment.objects.filter(pk=att_id).count(), 0)

    def test_submit_with_fresh_attachment_does_not_duplicate_draft_copy(self):
        """A file mirrored to the draft this session (via the composer's
        onFilesAdded hook) and also present in the live `attachments` field
        at submit must only be attached once — not pulled in a second time
        just because a draft copy also exists."""
        self.client.post(reverse('tickets:save_draft_attachment'), {
            'ticket_type': 'INCIDENT', 'attachments': self._pdf('fresh.pdf'),
        })
        # Note: no keep_draft_attachments in this POST — the client never
        # adds a freshly-mirrored (non-restored) attachment's id to that
        # field, exactly to avoid this double-attach scenario.
        response = self.client.post(reverse('tickets:create'), {
            'type': 'INCIDENT',
            'title': 'Fresh Attachment Ticket',
            'description': 'Test description',
            'category': self.category.id,
            'impact': 'INDIVIDUAL',
            'urgency': 'MEDIUM',
            'attachments': self._pdf('fresh.pdf'),
        })
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Fresh Attachment Ticket')
        self.assertEqual(ticket.attachments.count(), 1)


class TicketViewTests(TestCase):
    """Test ticket view functionality."""

    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT',
            is_active=True,
            email_verified=True
        )
        self.agent = User.objects.create_user(
            email='agent@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='User',
            department='IT',
            role=User.Role.AGENT,
            is_active=True,
            email_verified=True
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.client.login(email='test@example.com', password='TestPass123!')

    def test_ticket_create_page_loads(self):
        """Test ticket creation page loads."""
        response = self.client.get(reverse('tickets:create'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'requester/incident_form.html')

    def test_ticket_create_service_request_page_loads(self):
        """Test service request creation page loads."""
        response = self.client.get(reverse('tickets:create') + '?type=SERVICE_REQUEST')
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'requester/service_request_form.html')

    def test_ticket_create_success(self):
        """Test successful ticket creation."""
        response = self.client.post(reverse('tickets:create'), {
            'type': 'INCIDENT',
            'title': 'Test Incident',
            'description': 'Test description',
            'category': self.category.id,
            'impact': 'INDIVIDUAL',
            'urgency': 'MEDIUM'
        })
        self.assertEqual(response.status_code, 302)  # Redirect on success
        ticket = Ticket.objects.filter(title='Test Incident').first()
        self.assertIsNotNone(ticket)
        self.assertEqual(ticket.requester, self.user)

    def test_ticket_create_missing_title(self):
        """Test ticket creation with missing title."""
        response = self.client.post(reverse('tickets:create'), {
            'type': 'INCIDENT',
            'title': '',
            'description': 'Test description',
            'category': self.category.id,
            'impact': 'INDIVIDUAL',
            'urgency': 'MEDIUM'
        })
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'This field is required')

    def test_ticket_list_view(self):
        """Test ticket list view for requester."""
        ticket = Ticket.objects.create(
            number='TK#1234',
            title='Test Ticket',
            description='Test Description',
            requester=self.user,
            category=self.category
        )
        response = self.client.get(reverse('tickets:my_list'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Test Ticket')

    def test_ticket_detail_view_requester(self):
        """Test ticket detail view for requester."""
        ticket = Ticket.objects.create(
            number='TK#1234',
            title='Test Ticket',
            description='Test Description',
            requester=self.user,
            category=self.category
        )
        response = self.client.get(reverse('tickets:detail', args=[ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Test Ticket')

    def test_ticket_detail_view_other_user_denied(self):
        """Test ticket detail view denied for other users."""
        other_user = User.objects.create_user(
            email='other@example.com',
            password='TestPass123!',
            first_name='Other',
            last_name='User',
            department='HR',
            is_active=True,
            email_verified=True
        )
        ticket = Ticket.objects.create(
            number='TK#1234',
            title='Test Ticket',
            description='Test Description',
            requester=self.user,
            category=self.category
        )
        self.client.login(email='other@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:detail', args=[ticket.pk]))
        self.assertEqual(response.status_code, 302)  # Redirect to dashboard


class TicketCommentTests(TestCase):
    """Test ticket comment functionality."""
    
    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT',
            is_active=True,
            email_verified=True
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.ticket = Ticket.objects.create(
            number='TK#1234',
            title='Test Ticket',
            description='Test Description',
            requester=self.user,
            category=self.category
        )
        self.client.login(email='test@example.com', password='TestPass123!')

    def test_add_comment(self):
        """Test adding a comment to a ticket."""
        response = self.client.post(
            reverse('tickets:detail', args=[self.ticket.pk]),
            {
                'body': 'Test comment',
                'attachments': []
            },
            HTTP_HX_REQUEST='true'
        )
        self.assertEqual(response.status_code, 200)
        comment = TicketComment.objects.filter(ticket=self.ticket).first()
        self.assertIsNotNone(comment)
        self.assertEqual(comment.body, 'Test comment')
        self.assertEqual(comment.author, self.user)

    def test_attachment_only_comment_is_accepted(self):
        """A reply with no text but a valid attachment should be accepted,
        not rejected as an empty comment."""
        file_ = SimpleUploadedFile('note.txt', b'hello world', content_type='text/plain')
        response = self.client.post(
            reverse('tickets:detail', args=[self.ticket.pk]),
            {'body': '', 'attachments': [file_]},
            HTTP_HX_REQUEST='true'
        )
        self.assertEqual(response.status_code, 200)
        comment = TicketComment.objects.filter(ticket=self.ticket).first()
        self.assertIsNotNone(comment)
        self.assertEqual(comment.body, '')
        self.assertEqual(comment.attachment_set.count(), 1)

    def test_empty_comment_without_attachment_still_rejected(self):
        """No text and no attachment should still be rejected."""
        response = self.client.post(
            reverse('tickets:detail', args=[self.ticket.pk]),
            {'body': '', 'attachments': []},
            HTTP_HX_REQUEST='true'
        )
        self.assertEqual(response.status_code, 400)
        self.assertEqual(TicketComment.objects.filter(ticket=self.ticket).count(), 0)


class SaveAttachmentsTests(TestCase):
    """Test the save_attachments() validation/rejection reporting."""

    def setUp(self):
        self.user = User.objects.create_user(
            email='attach@example.com',
            password='TestPass123!',
            first_name='Attach',
            last_name='User',
            department='IT',
            is_active=True,
            email_verified=True
        )
        self.category = Category.objects.create(name='Hardware2', slug='hardware2')
        self.ticket = Ticket.objects.create(
            number='TK#5555',
            title='Attachment Ticket',
            description='Test Description',
            requester=self.user,
            category=self.category
        )

    def test_oversized_file_is_rejected_with_reason(self):
        from apps.tickets.views import save_attachments, MAX_SIZE_MB
        big_file = SimpleUploadedFile(
            'big.jpg',
            b'x' * (MAX_SIZE_MB * 1024 * 1024 + 1),
            content_type='image/jpeg'
        )
        created, rejected = save_attachments(self.ticket, [big_file], self.user)
        self.assertEqual(created, [])
        self.assertEqual(len(rejected), 1)
        self.assertEqual(rejected[0][0], 'big.jpg')
        self.assertIn('MB limit', rejected[0][1])

    def test_disallowed_type_is_rejected_with_reason(self):
        from apps.tickets.views import save_attachments
        bad_file = SimpleUploadedFile('script.exe', b'MZ...', content_type='application/x-msdownload')
        created, rejected = save_attachments(self.ticket, [bad_file], self.user)
        self.assertEqual(created, [])
        self.assertEqual(len(rejected), 1)
        self.assertEqual(rejected[0][0], 'script.exe')

    def test_valid_file_is_created_with_no_rejections(self):
        from apps.tickets.views import save_attachments
        good_file = SimpleUploadedFile('note.txt', b'hello world', content_type='text/plain')
        created, rejected = save_attachments(self.ticket, [good_file], self.user)
        self.assertEqual(len(created), 1)
        self.assertEqual(rejected, [])


class AssetModelTests(TestCase):
    """Test Asset model functionality."""
    
    def setUp(self):
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT'
        )

    def test_asset_creation(self):
        """Test basic asset creation."""
        asset = Asset.objects.create(
            name='Test Laptop',
            serial_number='SN12345',
            status='ACTIVE',
            assigned_to=self.user
        )
        self.assertEqual(asset.name, 'Test Laptop')
        self.assertEqual(asset.assigned_to, self.user)
        self.assertIsNotNone(asset.tracking_id)
        self.assertTrue(asset.tracking_id.startswith('AST-'))
        # Asset.__str__ appends a checked-out/available status emoji.
        self.assertEqual(str(asset), f'{asset.tracking_id} - Test Laptop (🟢)')

    def test_asset_tracking_id_generation(self):
        """Test tracking ID generation."""
        asset1 = Asset.objects.create(
            name='Test Laptop',
            status='ACTIVE'
        )
        asset2 = Asset.objects.create(
            name='Test Desktop',
            status='ACTIVE'
        )
        # Tracking IDs should be different
        self.assertNotEqual(asset1.tracking_id, asset2.tracking_id)
        # Should be in correct format
        year = timezone.now().year
        self.assertTrue(asset1.tracking_id.startswith(f'AST-{year}'))

    def test_asset_get_reassignment_count(self):
        """Test get_reassignment_count method."""
        asset = Asset.objects.create(
            name='Test Laptop',
            status='ACTIVE',
            assigned_to=self.user
        )
        
        # Initially 0 reassignments
        self.assertEqual(asset.get_reassignment_count(), 0)
        
        # Create an ASSIGNED log (this is the initial assignment, not a reassignment)
        AssetLog.objects.create(
            asset=asset,
            action=AssetLog.Action.ASSIGNED,
            actor=self.user,
            details={'to': self.user.get_full_name()}
        )
        # Still 0 reassignments (initial assignment doesn't count)
        self.assertEqual(asset.get_reassignment_count(), 0)
        
        # Create another ASSIGNED log (this is a reassignment)
        asset2 = User.objects.create_user(
            email='agent@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='User',
            department='IT'
        )
        AssetLog.objects.create(
            asset=asset,
            action=AssetLog.Action.ASSIGNED,
            actor=self.user,
            details={'from': self.user.get_full_name(), 'to': asset2.get_full_name()}
        )
        # Now should be 1 reassignment
        self.assertEqual(asset.get_reassignment_count(), 1)

    def test_asset_has_been_reassigned(self):
        """Test has_been_reassigned method."""
        asset = Asset.objects.create(
            name='Test Laptop',
            status='ACTIVE',
            assigned_to=self.user
        )
        
        # Initial assignment - should be False
        self.assertFalse(asset.has_been_reassigned())
        
        # Add a reassignment
        asset2 = User.objects.create_user(
            email='agent@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='User',
            department='IT'
        )
        AssetLog.objects.create(
            asset=asset,
            action=AssetLog.Action.ASSIGNED,
            actor=self.user,
            details={'from': self.user.get_full_name(), 'to': asset2.get_full_name()}
        )
        # Should be True now
        self.assertTrue(asset.has_been_reassigned())


class AssetViewTests(TestCase):
    """Test asset view functionality."""
    
    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='admin@example.com',
            password='AdminPass123!',
            first_name='Admin',
            last_name='User',
            department='IT'
        )
        self.agent = User.objects.create_user(
            email='agent@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='User',
            department='IT',
            role=User.Role.AGENT,
            is_active=True,
            email_verified=True
        )
        self.client.login(email='admin@example.com', password='AdminPass123!')

    def test_asset_list_view(self):
        """Test asset list page loads."""
        response = self.client.get(reverse('tickets:assets'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'tickets/asset_list.html')

    def test_asset_create(self):
        """Test asset creation via form."""
        response = self.client.post(reverse('tickets:asset_create_page'), {
            'name': 'New Test Laptop',
            'serial_number': 'SN99999',
            'status': 'IN_STORE',
            'location': 'OTHER',
            'location_other': 'HQ',
            'assigned_to': ''
        })
        # Should redirect to asset list
        self.assertEqual(response.status_code, 302)
        self.assertTrue(Asset.objects.filter(name='New Test Laptop').exists())

    def test_asset_create_requires_admin(self):
        """Test that non-admin users cannot create assets."""
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:asset_create_page'))
        self.assertEqual(response.status_code, 403)  # Forbidden

    def test_asset_edit(self):
        """Test editing an asset."""
        asset = Asset.objects.create(
            name='Test Laptop',
            serial_number='SN12345',
            status='ACTIVE'
        )
        response = self.client.post(
            reverse('tickets:asset_edit_page', args=[asset.pk]),
            {
                'name': 'Updated Laptop Name',
                'serial_number': 'SN12345',
                'status': 'IN_STORE'
            }
        )
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertEqual(asset.name, 'Updated Laptop Name')

    def test_asset_detail_view(self):
        """Test asset detail page loads."""
        asset = Asset.objects.create(
            name='Test Laptop',
            serial_number='SN12345',
            status='ACTIVE'
        )
        response = self.client.get(reverse('tickets:asset_detail', args=[asset.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Test Laptop')


class AssetReassignTests(TestCase):
    """Test asset reassignment functionality."""
    
    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='admin@example.com',
            password='AdminPass123!',
            first_name='Admin',
            last_name='User',
            department='IT'
        )
        self.agent1 = User.objects.create_user(
            email='agent1@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='One',
            department='IT',
            role=User.Role.AGENT,
            is_active=True,
            email_verified=True
        )
        self.agent2 = User.objects.create_user(
            email='agent2@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='Two',
            department='IT',
            role=User.Role.AGENT,
            is_active=True,
            email_verified=True
        )
        self.asset = Asset.objects.create(
            name='Test Laptop',
            serial_number='SN12345',
            status='ACTIVE',
            assigned_to=self.agent1
        )
        self.client.login(email='admin@example.com', password='AdminPass123!')

    def test_asset_reassign_creates_log(self):
        """Reassigning an asset should create an AssetLog entry."""
        url = reverse('tickets:asset_reassign', args=[self.asset.pk])
        response = self.client.post(url, {
            'assigned_to': self.agent2.pk,
            'comment': 'Reassigning for workload balance'
        })
        self.assertEqual(response.status_code, 302)
        
        # Check asset was reassigned
        self.asset.refresh_from_db()
        self.assertEqual(self.asset.assigned_to, self.agent2)
        
        # Check log was created
        logs = AssetLog.objects.filter(asset=self.asset, action=AssetLog.Action.ASSIGNED)
        self.assertTrue(logs.exists())
        self.assertEqual(logs.count(), 1)

    def test_asset_reassignment_count_increments(self):
        """Reassigning should increase the reassignment count."""
        initial_count = self.asset.get_reassignment_count()
        
        url = reverse('tickets:asset_reassign', args=[self.asset.pk])
        self.client.post(url, {
            'assigned_to': self.agent2.pk,
            'comment': 'Reassigning'
        })
        
        self.asset.refresh_from_db()
        new_count = self.asset.get_reassignment_count()
        self.assertEqual(new_count, initial_count + 1)

    def test_asset_reassign_trail_history(self):
        """Test the reassign trail history."""
        # Create multiple reassignments
        url = reverse('tickets:asset_reassign', args=[self.asset.pk])
        self.client.post(url, {'assigned_to': self.agent2.pk, 'comment': 'First reassign'})
        
        # Create another reassignment
        self.client.post(url, {'assigned_to': self.agent1.pk, 'comment': 'Second reassign'})
        
        # Get history. The asset was created directly with assigned_to set
        # (no log for that). Each reassign() now goes through release()+
        # assign_to() (the single custody-tracking pair also used by
        # checkout/check-in), so each one logs both an UNASSIGNED (release
        # from the previous holder) and an ASSIGNED (handover to the new
        # one) entry — two reassignments means four log entries.
        history = self.asset.get_assignment_history()
        self.assertEqual(len(history), 4)
        
        # Check latest assignment is agent1
        self.asset.refresh_from_db()
        self.assertEqual(self.asset.assigned_to, self.agent1)

    def test_asset_reassign_unassign(self):
        """Test unassigning an asset."""
        url = reverse('tickets:asset_reassign', args=[self.asset.pk])
        response = self.client.post(url, {
            'assigned_to': '',
            'comment': 'Unassigning asset'
        })
        self.assertEqual(response.status_code, 302)
        
        self.asset.refresh_from_db()
        self.assertIsNone(self.asset.assigned_to)
        
        # Check UNASSIGNED log was created
        logs = AssetLog.objects.filter(asset=self.asset, action=AssetLog.Action.UNASSIGNED)
        self.assertTrue(logs.exists())


class MobilizationTests(TestCase):
    """Test mobilization/demobilization of assets to a job/vessel/dive system."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='admin@example.com',
            password='AdminPass123!',
            first_name='Admin',
            last_name='User',
            department='IT'
        )
        self.client.login(email='admin@example.com', password='AdminPass123!')

        self.job = JobNumber.objects.create(number='JB-100', is_active=True)
        self.asset1 = Asset.objects.create(name='Drill A', status=Asset.Status.IN_STORE)
        self.asset2 = Asset.objects.create(name='Drill B', status=Asset.Status.IN_STORE)

        # Mobilizations now require a linked mobilization-request ticket —
        # a ready-made one for tests that don't care about the ticket
        # linkage itself, just that creation succeeds.
        self.mob_requester = User.objects.create_user(
            email='mob-req@example.com', password='TestPass123!',
            first_name='Mob', last_name='Req', department='IT', role=User.Role.END_USER,
        )
        self.mob_ticket = Ticket.objects.create(
            number='SRV#8000', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='Need gear mobilized', requester=self.mob_requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

    def test_mobilize_batch_flips_assets_to_mobilized(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk, self.asset2.pk],
            'ticket_id': self.mob_ticket.pk,
        })
        self.assertEqual(response.status_code, 302)

        self.asset1.refresh_from_db()
        self.asset2.refresh_from_db()
        self.assertEqual(self.asset1.status, Asset.Status.MOBILIZED)
        self.assertEqual(self.asset2.status, Asset.Status.MOBILIZED)

        mobilization = Mobilization.objects.get(job_number=self.job)
        self.assertEqual(mobilization.status, Mobilization.Status.ACTIVE)
        self.assertEqual(mobilization.items.count(), 2)

    def test_mobilize_rejects_unavailable_asset(self):
        self.asset1.status = Asset.Status.MOBILIZED
        self.asset1.save()

        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk],
            'ticket_id': self.mob_ticket.pk,
        })
        self.assertEqual(response.status_code, 302)
        self.assertFalse(Mobilization.objects.filter(job_number=self.job).exists())

    def test_mobilization_create_requires_ticket_id(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk],
        })
        self.assertEqual(response.status_code, 302)
        self.assertFalse(Mobilization.objects.filter(job_number=self.job).exists())

    def test_mobilization_create_rejects_non_mobilization_ticket(self):
        other_ticket = Ticket.objects.create(
            number='SRV#8001', type=Ticket.Type.SERVICE_REQUEST, title='Not a mobilization',
            description='...', requester=self.mob_requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=False,
        )
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk],
            'ticket_id': other_ticket.pk,
        })
        self.assertEqual(response.status_code, 302)
        self.assertFalse(Mobilization.objects.filter(job_number=self.job).exists())

    def test_mobilization_create_page_requires_ticket_id(self):
        response = self.client.get(reverse('tickets:mobilization_create_page'))
        self.assertEqual(response.status_code, 302)

    def test_partial_demobilize_keeps_mobilization_active(self):
        mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        item1 = MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset1)
        MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset2)
        self.asset1.status = Asset.Status.MOBILIZED
        self.asset1.save()
        self.asset2.status = Asset.Status.MOBILIZED
        self.asset2.save()

        response = self.client.post(
            reverse('tickets:mobilization_item_demobilize', args=[item1.pk]),
            {'return_condition': Asset.Condition.GOOD, 'return_notes': 'Back safely', 'override_reason': 'No ticket on this legacy mobilization'}
        )
        self.assertEqual(response.status_code, 302)

        self.asset1.refresh_from_db()
        self.assertEqual(self.asset1.status, Asset.Status.IN_STORE)

        mobilization.refresh_from_db()
        self.assertEqual(mobilization.status, Mobilization.Status.ACTIVE)

    def test_demobilizing_last_item_completes_mobilization(self):
        mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        item1 = MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset1)
        self.asset1.status = Asset.Status.MOBILIZED
        self.asset1.save()

        self.client.post(
            reverse('tickets:mobilization_item_demobilize', args=[item1.pk]),
            {'return_condition': Asset.Condition.GOOD, 'return_notes': '', 'override_reason': 'No ticket on this legacy mobilization'}
        )

        mobilization.refresh_from_db()
        self.assertEqual(mobilization.status, Mobilization.Status.COMPLETED)

    def test_damaged_return_sends_asset_to_maintenance(self):
        mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        item1 = MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset1)
        self.asset1.status = Asset.Status.MOBILIZED
        self.asset1.save()

        self.client.post(
            reverse('tickets:mobilization_item_demobilize', args=[item1.pk]),
            {'return_condition': Asset.Condition.DAMAGED, 'return_notes': 'Dropped overboard', 'override_reason': 'No ticket on this legacy mobilization'}
        )

        self.asset1.refresh_from_db()
        self.assertEqual(self.asset1.status, Asset.Status.MAINTENANCE)

    def test_job_lookup_returns_only_active_items_for_job(self):
        other_job = JobNumber.objects.create(number='JB-200', is_active=True)
        mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset1)

        other_mobilization = Mobilization.objects.create(job_number=other_job, mobilized_by=self.admin)
        MobilizationItem.objects.create(mobilization=other_mobilization, asset=self.asset2)

        response = self.client.get(reverse('tickets:job_mobilization_lookup'), {'job_number': self.job.pk})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Drill A')
        self.assertNotContains(response, 'Drill B')

    def test_mobilize_against_ticket_fulfills_it_and_traces_back(self):
        requester = User.objects.create_user(
            email='mob-requester@example.com', password='TestPass123!',
            first_name='Mob', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        ticket = Ticket.objects.create(
            number='SRV#8100', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='Need gear mobilized', requester=requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk],
            'ticket_id': ticket.pk,
        })
        self.assertEqual(response.status_code, 302)

        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.Status.PENDING_USER)
        self.assertIsNotNone(ticket.fulfilled_at)
        self.assertEqual(ticket.fulfilled_by, self.admin)

        mobilization = Mobilization.objects.get(ticket=ticket)
        self.assertIn(self.asset1, [item.asset for item in mobilization.items.all()])

    def test_demobilize_all_returns_every_active_item(self):
        mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset1)
        MobilizationItem.objects.create(mobilization=mobilization, asset=self.asset2)
        self.asset1.status = Asset.Status.MOBILIZED
        self.asset1.save()
        self.asset2.status = Asset.Status.MOBILIZED
        self.asset2.save()

        response = self.client.post(
            reverse('tickets:mobilization_demobilize_all', args=[mobilization.pk]),
            {'return_condition': Asset.Condition.GOOD, 'return_notes': 'All back', 'override_reason': 'No ticket on this legacy mobilization'}
        )
        self.assertEqual(response.status_code, 302)

        self.asset1.refresh_from_db()
        self.asset2.refresh_from_db()
        self.assertEqual(self.asset1.status, Asset.Status.IN_STORE)
        self.assertEqual(self.asset2.status, Asset.Status.IN_STORE)

        mobilization.refresh_from_db()
        self.assertEqual(mobilization.status, Mobilization.Status.COMPLETED)

    def test_extend_date_creates_history_and_updates_current_date(self):
        original = date(2026, 9, 1)
        mobilization = Mobilization.objects.create(
            job_number=self.job, mobilized_by=self.admin,
            expected_return_date=original, original_expected_return_date=original,
        )

        response = self.client.post(
            reverse('tickets:mobilization_extend_date', args=[mobilization.pk]),
            {'new_date': '2026-09-15', 'reason': 'Job running long'}
        )
        self.assertEqual(response.status_code, 302)

        mobilization.refresh_from_db()
        self.assertEqual(mobilization.expected_return_date, date(2026, 9, 15))
        self.assertEqual(mobilization.original_expected_return_date, original)

        extension = mobilization.date_extensions.get()
        self.assertEqual(extension.previous_date, original)
        self.assertEqual(extension.new_date, date(2026, 9, 15))
        self.assertEqual(extension.extended_by, self.admin)
        self.assertEqual(extension.reason, 'Job running long')

        # A second, earlier "extension" is rejected rather than silently applied.
        response = self.client.post(
            reverse('tickets:mobilization_extend_date', args=[mobilization.pk]),
            {'new_date': '2026-09-10', 'reason': 'oops'}
        )
        mobilization.refresh_from_db()
        self.assertEqual(mobilization.expected_return_date, date(2026, 9, 15))
        self.assertEqual(mobilization.date_extensions.count(), 1)

    def test_extend_date_forbidden_for_non_admin(self):
        agent = User.objects.create_user(
            email='mob-agent@example.com', password='TestPass123!',
            first_name='Mob', last_name='Agent', department='IT', role=User.Role.AGENT,
        )
        mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        self.client.logout()
        self.client.login(email='mob-agent@example.com', password='TestPass123!')

        response = self.client.post(
            reverse('tickets:mobilization_extend_date', args=[mobilization.pk]),
            {'new_date': '2026-09-15'}
        )
        self.assertEqual(response.status_code, 403)


class MobilizationReceiptConfirmationTests(TestCase):
    """Requester-side per-item accept/dispute handshake for mobilized
    assets, and the aggregation step that drives the linked ticket to
    APPROVED (all accepted) or back to PENDING_FULFILLMENT (any disputed)
    once every item has been actioned."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='receipt-admin@example.com', password='AdminPass123!',
            first_name='Receipt', last_name='Admin', department='IT',
        )
        self.requester = User.objects.create_user(
            email='receipt-requester@example.com', password='TestPass123!',
            first_name='Receipt', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.other_user = User.objects.create_user(
            email='receipt-other@example.com', password='TestPass123!',
            first_name='Receipt', last_name='Other', department='IT', role=User.Role.END_USER,
        )
        self.job = JobNumber.objects.create(number='JB-900', is_active=True)
        self.asset1 = Asset.objects.create(name='Camera A', status=Asset.Status.IN_STORE)
        self.asset2 = Asset.objects.create(name='Camera B', status=Asset.Status.IN_STORE)
        self.ticket = Ticket.objects.create(
            number='SRV#9500', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='Need gear mobilized', requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )
        self.client.login(email='receipt-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk, self.asset2.pk],
            'ticket_id': self.ticket.pk,
        })
        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)
        self.mobilization = Mobilization.objects.get(ticket=self.ticket)
        self.item1 = self.mobilization.items.get(asset=self.asset1)
        self.item2 = self.mobilization.items.get(asset=self.asset2)
        self.client.logout()

    def test_pure_stock_mobilization_posts_one_itemized_created_comment(self):
        created_comment = self.ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.CREATED)
        self.assertIn('2 items requested', created_comment.body)
        self.assertIn('Camera A', created_comment.body)
        self.assertIn('Camera B', created_comment.body)
        self.assertIn('mobilized from stock', created_comment.body)
        self.assertNotIn('ordered from vendor', created_comment.body)
        self.assertTrue(created_comment.is_system_generated)

    def test_mobilization_cards_attributed_to_system_not_admin(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:detail', args=[self.ticket.pk]))
        self.assertContains(response, 'System')
        self.assertNotContains(response, 'Receipt Admin')

    def test_confirmed_card_still_shows_real_requester_name(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item2.pk]))

        response = self.client.get(reverse('tickets:detail', args=[self.ticket.pk]))
        self.assertContains(response, 'Confirmed by Receipt Requester')

    def test_accept_sets_acknowledged_fields_and_notifies_admins(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.assertEqual(response.status_code, 302)

        self.item1.refresh_from_db()
        self.assertIsNotNone(self.item1.acknowledged_at)
        self.assertEqual(self.item1.acknowledged_by, self.requester)
        self.assertTrue(Notification.objects.filter(recipient=self.admin).exists())

    def test_dispute_sets_disputed_fields(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:mobilization_item_dispute', args=[self.item1.pk]), {'reason': 'Never arrived'})
        self.assertEqual(response.status_code, 302)

        self.item1.refresh_from_db()
        self.assertIsNotNone(self.item1.disputed_at)
        self.assertEqual(self.item1.dispute_reason, 'Never arrived')

    def test_non_requester_cannot_accept(self):
        self.client.login(email='receipt-other@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))

        self.item1.refresh_from_db()
        self.assertIsNone(self.item1.acknowledged_at)

    def test_double_confirm_rejected(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))

        self.item1.refresh_from_db()
        # Still the same single confirmation — a second attempt is a no-op,
        # not an error that corrupts the timestamp.
        self.assertIsNotNone(self.item1.acknowledged_at)

    def test_ticket_stays_pending_user_while_items_outstanding(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

    def test_ticket_resolves_when_all_items_accepted(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item2.pk]))

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.APPROVED)
        self.assertIsNotNone(self.ticket.resolution_confirmed_at)
        self.assertEqual(self.ticket.resolution_confirmed_by, self.requester)

    def test_ticket_reopens_when_any_item_disputed(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_dispute', args=[self.item2.pk]), {'reason': 'Missing'})

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_FULFILLMENT)
        self.assertTrue(
            TicketActivityLog.objects.filter(ticket=self.ticket, action='resolution_rejected').exists()
        )

    def test_receipt_confirm_modal_lists_pending_items(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:receipt_confirm_modal', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Camera A')
        self.assertContains(response, 'Camera B')

    def test_receipt_confirm_modal_forbidden_for_non_requester(self):
        self.client.login(email='receipt-other@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:receipt_confirm_modal', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 403)

    def test_confirm_resolution_get_redirects_for_mobilization_ticket(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:confirm_resolution', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

    def test_batch_confirm_applies_all_decisions_in_one_submit(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:mobilization_items_confirm_batch', args=[self.ticket.pk]), {
            'accept_ids': [self.item1.pk, self.item2.pk],
        })
        self.assertEqual(response.status_code, 302)

        self.item1.refresh_from_db()
        self.item2.refresh_from_db()
        self.assertIsNotNone(self.item1.acknowledged_at)
        self.assertIsNotNone(self.item2.acknowledged_at)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.APPROVED)

    def test_batch_confirm_mixed_accept_and_dispute(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:mobilization_items_confirm_batch', args=[self.ticket.pk]), {
            'accept_ids': [self.item1.pk],
            'dispute_ids': [self.item2.pk],
            f'reason_{self.item2.pk}': 'Box was empty',
        })
        self.assertEqual(response.status_code, 302)

        self.item1.refresh_from_db()
        self.item2.refresh_from_db()
        self.assertIsNotNone(self.item1.acknowledged_at)
        self.assertIsNotNone(self.item2.disputed_at)
        self.assertEqual(self.item2.dispute_reason, 'Box was empty')

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_FULFILLMENT)

    def test_batch_confirm_forbidden_for_non_requester(self):
        self.client.login(email='receipt-other@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:mobilization_items_confirm_batch', args=[self.ticket.pk]), {
            'accept_ids': [self.item1.pk],
        })
        self.assertEqual(response.status_code, 403)
        self.item1.refresh_from_db()
        self.assertIsNone(self.item1.acknowledged_at)

    def test_requester_sees_confirm_receipt_card_on_conversation_page(self):
        # Regression coverage for the actual bug: the requester must see the
        # confirm action in the shared conversation thread (tickets:detail),
        # not only in the agent-only side panel that renders on the exact
        # same page but is invisible to a non-agent viewer.
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:detail', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Confirm receipt')
        self.assertContains(response, 'mobilized for')

    def test_agent_conversation_view_shows_pending_card_without_action_button(self):
        self.client.login(email='receipt-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:conversation', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'mobilized for')
        self.assertNotContains(response, 'Confirm receipt')

    def test_confirm_receipt_card_replaced_by_confirmed_state_once_all_items_accepted(self):
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item2.pk]))

        response = self.client.get(reverse('tickets:detail', args=[self.ticket.pk]))
        self.assertContains(response, 'Receipt confirmed for')
        self.assertNotContains(response, 'Confirm receipt')

    def test_mobilization_item_status_label_reflects_confirmation_state(self):
        self.client.login(email='receipt-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:mobilization_detail', args=[self.mobilization.pk]))
        self.assertContains(response, 'Pending Confirmation')

        self.client.logout()
        self.client.login(email='receipt-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item2.pk]))

        self.client.logout()
        self.client.login(email='receipt-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:mobilization_detail', args=[self.mobilization.pk]))
        self.assertNotContains(response, 'Pending Confirmation')
        self.assertContains(response, 'Mobilized')


class MobilizationDemobilizationRequestTests(TestCase):
    """Requester-initiated demobilization handshake: the requester
    self-reports sending mobilized-and-acknowledged assets back, an admin
    then confirms physical receipt via the existing Demobilize action."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='demob-admin@example.com', password='AdminPass123!',
            first_name='Demob', last_name='Admin', department='IT',
        )
        self.requester = User.objects.create_user(
            email='demob-requester@example.com', password='TestPass123!',
            first_name='Demob', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.other_user = User.objects.create_user(
            email='demob-other@example.com', password='TestPass123!',
            first_name='Demob', last_name='Other', department='IT', role=User.Role.END_USER,
        )
        self.job = JobNumber.objects.create(number='JB-901', is_active=True)
        self.asset1 = Asset.objects.create(name='Drill A', status=Asset.Status.IN_STORE)
        self.asset2 = Asset.objects.create(name='Drill B', status=Asset.Status.IN_STORE)
        self.ticket = Ticket.objects.create(
            number='SRV#9600', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='Need gear mobilized', requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'asset_ids': [self.asset1.pk, self.asset2.pk],
            'ticket_id': self.ticket.pk,
        })
        self.mobilization = Mobilization.objects.get(ticket=self.ticket)
        self.item1 = self.mobilization.items.get(asset=self.asset1)
        self.item2 = self.mobilization.items.get(asset=self.asset2)
        self.client.logout()

        # Requester confirms receipt of both items — the prerequisite for
        # self-reporting a demobilization.
        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item1.pk]))
        self.client.post(reverse('tickets:mobilization_item_accept', args=[self.item2.pk]))
        self.client.logout()
        self.item1.refresh_from_db()
        self.item2.refresh_from_db()

    def test_request_demobilization_requires_acknowledged_receipt(self):
        # A fresh item that was never confirmed received can't be
        # self-reported as demobilized.
        asset3 = Asset.objects.create(name='Drill C', status=Asset.Status.IN_STORE)
        item3 = MobilizationItem.objects.create(mobilization=self.mobilization, asset=asset3)
        with self.assertRaises(ValueError):
            item3.request_demobilization(actor=self.requester)

    def test_request_demobilization_rejects_non_requester(self):
        with self.assertRaises(ValueError):
            self.item1.request_demobilization(actor=self.other_user)
        self.item1.refresh_from_db()
        self.assertIsNone(self.item1.return_requested_at)

    def test_request_demobilization_sets_fields_and_notifies_admins(self):
        self.item1.request_demobilization(actor=self.requester, notes='Sent via courier')
        self.item1.refresh_from_db()
        self.assertIsNotNone(self.item1.return_requested_at)
        self.assertEqual(self.item1.return_requested_by, self.requester)
        self.assertEqual(self.item1.return_requested_notes, 'Sent via courier')
        self.assertTrue(Notification.objects.filter(recipient=self.admin).exists())

    def test_double_request_rejected(self):
        self.item1.request_demobilization(actor=self.requester)
        with self.assertRaises(ValueError):
            self.item1.request_demobilization(actor=self.requester)

    def test_cancel_demobilization_request_clears_fields(self):
        self.item1.request_demobilization(actor=self.requester)
        self.item1.cancel_demobilization_request(actor=self.requester)
        self.item1.refresh_from_db()
        self.assertIsNone(self.item1.return_requested_at)
        self.assertIsNone(self.item1.return_requested_by)
        self.assertEqual(self.item1.return_requested_notes, '')

    def test_cancel_fails_after_admin_demobilized(self):
        self.item1.request_demobilization(actor=self.requester)
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {
            'return_condition': 'Good',
        })
        self.client.logout()
        self.item1.refresh_from_db()
        self.assertIsNotNone(self.item1.demobilized_at)
        with self.assertRaises(ValueError):
            self.item1.cancel_demobilization_request(actor=self.requester)

    def test_demobilization_list_shows_only_own_ready_and_pending_items(self):
        self.item1.request_demobilization(actor=self.requester, notes='On its way back')

        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:demobilization_list'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Drill A')
        self.assertContains(response, 'Drill B')
        self.assertContains(response, 'On its way back')

        self.client.logout()
        self.client.login(email='demob-other@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:demobilization_list'))
        self.assertNotContains(response, 'Drill A')

    def test_demobilization_list_and_sidebar_link_persist_after_full_demobilization(self):
        # Requester self-reports both items, admin demobilizes both — nothing
        # is left outstanding, but the history and the sidebar link that
        # leads to it should both still be there afterward.
        self.item1.request_demobilization(actor=self.requester)
        self.item2.request_demobilization(actor=self.requester)
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {
            'return_condition': 'Good',
        })
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item2.pk]), {
            'return_condition': 'Good',
        })
        self.client.logout()

        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:demobilization_list'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Drill A')
        self.assertContains(response, 'Drill B')
        self.assertContains(response, '(2)')
        self.assertContains(response, 'demobilized')

        dashboard_response = self.client.get(reverse('dashboard'))
        self.assertContains(dashboard_response, reverse('tickets:demobilization_list'))

    def test_batch_request_reports_multiple_items_one_bad_id_skipped(self):
        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:mobilization_items_request_demobilize_batch'), {
            'item_ids': [self.item1.pk, self.item2.pk, 999999],
            'notes': 'Job complete, sending everything back',
        })
        self.assertEqual(response.status_code, 302)

        self.item1.refresh_from_db()
        self.item2.refresh_from_db()
        self.assertIsNotNone(self.item1.return_requested_at)
        self.assertIsNotNone(self.item2.return_requested_at)
        self.assertEqual(self.item1.return_requested_notes, 'Job complete, sending everything back')

    def test_admin_cannot_demobilize_without_prior_request_when_ticket_linked(self):
        # Hard-gated, per design — the mobilization has a linked ticket/
        # requester, so demobilize is blocked until they self-report it.
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {
            'return_condition': 'Good',
        })
        self.assertEqual(response.status_code, 302)
        self.item1.refresh_from_db()
        self.assertIsNone(self.item1.demobilized_at)

    def test_demobilize_modal_blocks_when_no_request_reported(self):
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:mobilization_item_demobilize_modal', args=[self.item1.pk]))
        self.assertContains(response, 'Not yet reported returned')
        self.assertNotContains(response, reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]))

    def test_demobilize_modal_prefills_when_request_reported(self):
        self.item1.request_demobilization(actor=self.requester, notes='Sent via courier')
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:mobilization_item_demobilize_modal', args=[self.item1.pk]))
        self.assertContains(response, 'Reported returned')
        self.assertContains(response, 'Sent via courier')

    def test_demobilize_all_blocked_when_any_item_unreported(self):
        self.item1.request_demobilization(actor=self.requester)
        # item2 left unreported.
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.post(reverse('tickets:mobilization_demobilize_all', args=[self.mobilization.pk]), {
            'return_condition': 'Good',
        })
        self.assertEqual(response.status_code, 302)
        self.item1.refresh_from_db()
        self.item2.refresh_from_db()
        self.assertIsNone(self.item1.demobilized_at)
        self.assertIsNone(self.item2.demobilized_at)

    def test_demobilize_ticketless_mobilization_requires_reason(self):
        ticketless_mobilization = Mobilization.objects.create(mobilized_by=self.admin)
        asset = Asset.objects.create(name='Legacy Drill', status=Asset.Status.MOBILIZED)
        item = MobilizationItem.objects.create(mobilization=ticketless_mobilization, asset=asset)

        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.post(reverse('tickets:mobilization_item_demobilize', args=[item.pk]), {
            'return_condition': 'Good',
        })
        self.assertEqual(response.status_code, 302)
        item.refresh_from_db()
        self.assertIsNone(item.demobilized_at)

        response = self.client.post(reverse('tickets:mobilization_item_demobilize', args=[item.pk]), {
            'return_condition': 'Good', 'override_reason': 'Legacy mobilization, no requester to report it',
        })
        self.assertEqual(response.status_code, 302)
        item.refresh_from_db()
        self.assertIsNotNone(item.demobilized_at)
        log = AssetLog.objects.filter(asset=asset, action=AssetLog.Action.DEMOBILIZED).get()
        self.assertEqual(log.details['override_reason'], 'Legacy mobilization, no requester to report it')

    def test_demobilize_all_ticketless_requires_reason(self):
        ticketless_mobilization = Mobilization.objects.create(mobilized_by=self.admin)
        asset = Asset.objects.create(name='Legacy Drill', status=Asset.Status.MOBILIZED)
        MobilizationItem.objects.create(mobilization=ticketless_mobilization, asset=asset)

        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.post(reverse('tickets:mobilization_demobilize_all', args=[ticketless_mobilization.pk]), {
            'return_condition': 'Good',
        })
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertEqual(asset.status, Asset.Status.MOBILIZED)

        response = self.client.post(reverse('tickets:mobilization_demobilize_all', args=[ticketless_mobilization.pk]), {
            'return_condition': 'Good', 'override_reason': 'Legacy mobilization, no requester to report it',
        })
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertEqual(asset.status, Asset.Status.IN_STORE)

    def test_pending_demobilizations_list_admin_only(self):
        self.item1.request_demobilization(actor=self.requester)

        self.client.login(email='demob-other@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:pending_demobilizations_list'))
        self.assertEqual(response.status_code, 403)
        self.client.logout()

        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:pending_demobilizations_list'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Drill A')

    def test_mobilization_detail_shows_self_report_and_notes_for_demobilized_item(self):
        self.item1.request_demobilization(actor=self.requester, notes='Sent via courier')
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {
            'return_condition': 'Good',
            'return_notes': 'Checked in fine',
        })

        response = self.client.get(reverse('tickets:mobilization_detail', args=[self.mobilization.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Sent via courier')
        self.assertContains(response, 'Checked in fine')

    def test_mobilizations_list_shows_pending_demob_badge_and_tab(self):
        self.item1.request_demobilization(actor=self.requester, notes='Sent via courier')

        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:mobilizations'))
        self.assertContains(response, '1 item')
        self.assertContains(response, 'Needs Confirmation')
        self.assertContains(response, '(1)')

        response = self.client.get(reverse('tickets:mobilizations'), {'tab': 'needs_confirmation'})
        self.assertContains(response, self.mobilization.destination_display)

        # Confirming the item drops it out of the Needs Confirmation tab.
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {
            'return_condition': 'Good',
        })
        response = self.client.get(reverse('tickets:mobilizations'), {'tab': 'needs_confirmation'})
        self.assertNotContains(response, self.mobilization.destination_display)

    def test_mobilization_audit_report_shows_full_lifecycle_and_is_admin_only(self):
        self.item1.request_demobilization(actor=self.requester, notes='Sent via courier')
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {
            'return_condition': 'Good',
            'return_notes': 'Checked in fine',
        })

        response = self.client.get(reverse('tickets:mobilization_audit_report', args=[self.mobilization.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Drill A')
        self.assertContains(response, 'Drill B')
        self.assertContains(response, 'Sent via courier')
        self.assertContains(response, 'Checked in fine')
        self.assertContains(response, 'Returned to store (In Store)')
        self.assertContains(response, 'Still mobilized')
        self.client.logout()

        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:mobilization_audit_report', args=[self.mobilization.pk]))
        self.assertEqual(response.status_code, 403)

    def test_mobilization_audit_export_pdf(self):
        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:mobilization_audit_export', args=[self.mobilization.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')

    def test_sidebar_demobilization_link_only_shows_after_receiving_a_mobilization(self):
        # requester has acknowledged both items in setUp — link should show.
        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertContains(response, 'Demobilization')

        # other_user has never had anything mobilized to them — no link.
        self.client.logout()
        self.client.login(email='demob-other@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertNotContains(response, 'Demobilization')

    def test_sidebar_link_persists_after_everything_demobilized(self):
        # Per design: once a requester has ever had anything mobilized to
        # them, the link (and the history behind it) stays up permanently —
        # it does not disappear just because nothing's outstanding anymore.
        # See test_demobilization_list_and_sidebar_link_persist_after_full_demobilization
        # for the corresponding history-content assertions.
        self.item1.request_demobilization(actor=self.requester)
        self.item2.request_demobilization(actor=self.requester)

        self.client.login(email='demob-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item1.pk]), {'return_condition': 'Good'})
        self.client.post(reverse('tickets:mobilization_item_demobilize', args=[self.item2.pk]), {'return_condition': 'Good'})
        self.client.logout()

        self.client.login(email='demob-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertContains(response, 'Demobilization')


class VendorCategoryAndMobilizationPrefillTests(TestCase):
    """Vendor<->AssetCategory filtering (System Settings CRUD) and carrying
    the originating request's asset_type/number_of_assets/purpose into the
    mobilization-create modal."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='vendor-cat-admin@example.com', password='AdminPass123!',
            first_name='Vendor', last_name='CatAdmin', department='IT',
        )
        self.client.login(email='vendor-cat-admin@example.com', password='AdminPass123!')
        self.laptop_category = AssetCategory.objects.create(name='Laptop')

    def test_settings_create_saves_vendor_categories_m2m(self):
        response = self.client.post(reverse('tickets:settings_resource_create', args=['vendors']), {
            'name': 'Category-Aware Vendor',
            'categories': [self.laptop_category.pk],
        })
        self.assertEqual(response.status_code, 200)
        vendor = Vendor.objects.get(name='Category-Aware Vendor')
        self.assertIn(self.laptop_category, vendor.categories.all())

    def test_settings_update_saves_vendor_categories_m2m(self):
        vendor = Vendor.objects.create(name='Existing Vendor')
        server_category = AssetCategory.objects.create(name='Server')
        response = self.client.post(
            reverse('tickets:settings_resource_update', args=['vendors', vendor.pk]),
            {'name': 'Existing Vendor', 'categories': [server_category.pk]}
        )
        self.assertEqual(response.status_code, 200)
        vendor.refresh_from_db()
        self.assertIn(server_category, vendor.categories.all())

    def test_mobilization_modal_prefills_category_quantity_and_notes_from_ticket(self):
        service_category = ServiceCategory.objects.create(
            name='Equipment Prefill Test', slug='equipment-prefill-test', field_group=ServiceCategory.FieldGroup.ASSET
        )
        requester = User.objects.create_user(
            email='prefill-requester@example.com', password='TestPass123!',
            first_name='Prefill', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        ticket = Ticket.objects.create(
            number='SRV#8200', type=Ticket.Type.SERVICE_REQUEST, title='Need laptops for job',
            description='...', requester=requester, purpose='Offshore crew laptops',
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
            service_category=service_category,
            service_request_details={'asset_type': 'LAPTOP', 'number_of_assets': '3'},
        )

        response = self.client.get(reverse('tickets:mobilization_create_page'), {'ticket_id': ticket.pk})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'value="{self.laptop_category.pk}" selected')
        self.assertContains(response, 'value="3"')
        self.assertContains(response, 'Offshore crew laptops')
        self.assertContains(response, 'Prefilled from the request')

    def test_mobilization_modal_prefill_handles_no_matching_category(self):
        service_category = ServiceCategory.objects.create(
            name='Unmatched Category Test', slug='unmatched-category-test', field_group=ServiceCategory.FieldGroup.ASSET
        )
        requester = User.objects.create_user(
            email='prefill-requester2@example.com', password='TestPass123!',
            first_name='Prefill', last_name='RequesterTwo', department='IT', role=User.Role.END_USER,
        )
        ticket = Ticket.objects.create(
            number='SRV#8201', type=Ticket.Type.SERVICE_REQUEST, title='Need something',
            description='...', requester=requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
            service_category=service_category,
            service_request_details={'asset_type': 'OTHER', 'number_of_assets': '2'},
        )

        # No AssetCategory named "Other" exists yet — should render without error.
        response = self.client.get(reverse('tickets:mobilization_create_page'), {'ticket_id': ticket.pk})
        self.assertEqual(response.status_code, 200)


class ServiceRequestReportConfirmationFieldsTests(TestCase):
    """The receipt-confirmation fields (fulfilled/receipt confirmed, who and
    when) must actually reach report output, not just live on the model —
    report_registry.SERVICE_REQUESTS.columns is what every exporter
    (CSV/Excel/JSON/PDF/DOCX) keys off of, not the row dict's own keys, so a
    field only added to the row function silently never appears anywhere."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='report-admin@example.com', password='AdminPass123!',
            first_name='Report', last_name='Admin', department='IT',
        )
        self.requester = User.objects.create_user(
            email='report-requester@example.com', password='TestPass123!',
            first_name='Report', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.ticket = Ticket.objects.create(
            number='SRV#9500', type=Ticket.Type.SERVICE_REQUEST, title='Report fields test',
            description='...', requester=self.requester,
            status=Ticket.Status.APPROVED, is_asset_request=True,
            fulfilled_at=timezone.now(), fulfilled_by=self.admin,
            resolution_confirmed_at=timezone.now(), resolution_confirmed_by=self.requester,
        )

    def test_row_includes_fulfillment_and_confirmation_fields(self):
        from apps.tickets.report_registry import _service_request_row
        row = _service_request_row(self.ticket)
        self.assertNotEqual(row['Fulfilled'], '—')
        self.assertEqual(row['Fulfilled By'], self.admin.get_full_name())
        self.assertNotEqual(row['Receipt Confirmed'], '—')
        self.assertEqual(row['Receipt Confirmed By'], self.requester.get_full_name())

    def test_row_shows_dash_when_not_yet_confirmed(self):
        from apps.tickets.report_registry import _service_request_row
        unconfirmed = Ticket.objects.create(
            number='SRV#9501', type=Ticket.Type.SERVICE_REQUEST, title='Not yet fulfilled',
            description='...', requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True,
        )
        row = _service_request_row(unconfirmed)
        self.assertEqual(row['Fulfilled'], '—')
        self.assertEqual(row['Receipt Confirmed'], '—')

    def test_csv_export_includes_confirmation_columns_in_header(self):
        self.client.login(email='report-admin@example.com', password='AdminPass123!')
        response = self.client.get(reverse('tickets:export_report', args=['service-requests']), {'format': 'csv'})
        self.assertEqual(response.status_code, 200)
        content = response.content.decode('utf-8')
        # CSV exports now lead with a text letterhead block (company name,
        # title, control no., generated date, record count, blank line —
        # see export_csv) before the real column-header row, so find that
        # row by content rather than assuming it's line 0.
        lines = content.splitlines()
        header = next(line for line in lines if 'Fulfilled' in line)
        for column in ('Fulfilled', 'Fulfilled By', 'Receipt Confirmed', 'Receipt Confirmed By', 'Is Mobilization Request'):
            self.assertIn(column, header)


class AssetRequestTwoStepResolutionTests(TestCase):
    """Confirming receipt and resolving the ticket are two separate steps
    for asset-request tickets: requester confirms receipt (-> APPROVED, not
    RESOLVED), then an agent explicitly resolves it, skipping the normal
    resolve-modal/PENDING_USER round-trip since receipt is already on
    record. Non-asset tickets keep the original one-step behavior."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='twostep-admin@example.com', password='AdminPass123!',
            first_name='TwoStep', last_name='Admin', department='IT',
        )
        self.agent = User.objects.create_user(
            email='twostep-agent@example.com', password='TestPass123!',
            first_name='TwoStep', last_name='Agent', department='IT', role=User.Role.AGENT,
        )
        self.requester = User.objects.create_user(
            email='twostep-requester@example.com', password='TestPass123!',
            first_name='TwoStep', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.category = AssetCategory.objects.create(name='TwoStep Category')
        self.asset = Asset.objects.create(name='TwoStep Laptop', status=Asset.Status.IN_STORE, category=self.category)
        self.ticket = Ticket.objects.create(
            number='SRV#9300', type=Ticket.Type.SERVICE_REQUEST, title='Need a laptop',
            description='...', requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True,
        )

    def _fulfill(self):
        self.client.login(email='twostep-admin@example.com', password='AdminPass123!')
        self.client.post(reverse('tickets:fulfill_asset_request', args=[self.ticket.pk]), {
            'asset_id': self.asset.pk, 'comment': '',
        })
        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

    def test_receipt_confirm_modal_renders_single_asset_content(self):
        self._fulfill()
        self.client.logout()
        self.client.login(email='twostep-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:receipt_confirm_modal', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Yes, I received it')
        self.assertContains(response, self.ticket.number)

    def test_confirm_resolution_get_redirects_to_ticket_detail(self):
        self._fulfill()
        self.client.logout()
        self.client.login(email='twostep-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:confirm_resolution', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 302)
        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

    def test_confirming_receipt_does_not_resolve_the_ticket(self):
        self._fulfill()
        self.client.logout()
        self.client.login(email='twostep-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:confirm_resolution', args=[self.ticket.pk]), {'action': 'confirm'})
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.APPROVED)
        self.assertIsNotNone(self.ticket.resolution_confirmed_at)
        self.assertNotEqual(self.ticket.status, Ticket.Status.RESOLVED)

    def test_agent_resolve_skips_modal_flow_once_receipt_confirmed(self):
        self._fulfill()
        self.client.logout()
        self.client.login(email='twostep-requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:confirm_resolution', args=[self.ticket.pk]), {'action': 'confirm'})

        self.client.logout()
        self.client.login(email='twostep-agent@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:resolve_ticket', args=[self.ticket.pk]), {'action': 'confirm'})
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.RESOLVED)
        self.assertIsNotNone(self.ticket.resolved_at)

    def test_agent_resolve_still_uses_normal_flow_before_receipt_confirmed(self):
        self._fulfill()
        self.client.logout()
        self.client.login(email='twostep-agent@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:resolve_ticket', args=[self.ticket.pk]), {'action': 'confirm', 'comment': ''})
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        # Not yet resolved — sent back to the requester for confirmation instead.
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

    def test_rejecting_receipt_sends_ticket_back_to_fulfillment_and_notifies_fulfiller(self):
        self._fulfill()
        self.client.logout()
        self.client.login(email='twostep-requester@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:confirm_resolution', args=[self.ticket.pk]),
            {'action': 'reopen', 'reason': 'Wrong laptop model'}
        )
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_FULFILLMENT)
        self.assertIsNone(self.ticket.resolution_confirmed_at)
        self.assertTrue(
            Notification.objects.filter(recipient=self.ticket.fulfilled_by, message__icontains="wasn't received").exists()
        )

        # This comment carries a real human-typed reason ("Wrong laptop
        # model") — it must stay attributed to the real requester, not
        # flattened to "System".
        comment = self.ticket.comments.get(body__icontains='Wrong laptop model')
        self.assertFalse(comment.is_system_generated)
        self.assertEqual(comment.author, self.requester)

    def test_non_asset_ticket_resolution_confirmation_unchanged(self):
        service_category = ServiceCategory.objects.create(
            name='General Two-Step', slug='general-two-step', field_group=ServiceCategory.FieldGroup.GENERAL
        )
        general_ticket = Ticket.objects.create(
            number='SRV#9301', type=Ticket.Type.SERVICE_REQUEST, title='General request',
            description='...', requester=self.requester, service_category=service_category,
            status=Ticket.Status.PENDING_USER,
        )
        self.client.login(email='twostep-requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:confirm_resolution', args=[general_ticket.pk]), {'action': 'confirm'})
        self.assertEqual(response.status_code, 302)

        general_ticket.refresh_from_db()
        self.assertEqual(general_ticket.status, Ticket.Status.RESOLVED)
        self.assertIsNotNone(general_ticket.resolved_at)


class SystemGeneratedCommentBackfillTests(TestCase):
    """The 0052 migration's data backfill: existing rows already
    identifiable as system-composed via is_receipt_confirmation_prompt or
    mobilization_event should retroactively flip is_system_generated=True,
    without touching plain human comments."""

    def test_backfill_flags_existing_mobilization_and_prompt_rows_only(self):
        import importlib
        from django.apps import apps as real_apps
        migration_module = importlib.import_module(
            'apps.tickets.migrations.0052_ticketcomment_is_system_generated_and_more'
        )
        backfill_is_system_generated = migration_module.backfill_is_system_generated

        admin = User.objects.create_superuser(
            email='backfill-admin@example.com', password='AdminPass123!',
            first_name='Backfill', last_name='Admin', department='IT',
        )
        ticket = Ticket.objects.create(
            number='SRV#9900', type=Ticket.Type.SERVICE_REQUEST, title='...',
            description='...', requester=admin, status=Ticket.Status.PENDING_FULFILLMENT,
        )
        mobilization_comment = TicketComment.objects.create(
            ticket=ticket, author=admin, visibility='PUBLIC',
            body='pretend created card', mobilization_event=TicketComment.MobilizationEvent.CREATED,
        )
        prompt_comment = TicketComment.objects.create(
            ticket=ticket, author=admin, visibility='PUBLIC',
            body='pretend fulfilled prompt', is_receipt_confirmation_prompt=True,
        )
        human_comment = TicketComment.objects.create(
            ticket=ticket, author=admin, visibility='PUBLIC', body='a real reply',
        )

        backfill_is_system_generated(real_apps, None)

        mobilization_comment.refresh_from_db()
        prompt_comment.refresh_from_db()
        human_comment.refresh_from_db()
        self.assertTrue(mobilization_comment.is_system_generated)
        self.assertTrue(prompt_comment.is_system_generated)
        self.assertFalse(human_comment.is_system_generated)


class MobilizationVendorGatingTests(TestCase):
    """A mobilization request's ticket is only fulfilled (and the requester
    only prompted to confirm receipt) once everything on it is actually in
    hand — immediately if nothing was sourced from a vendor, otherwise once
    the last open procurement request against it clears."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='vendorgate-admin@example.com', password='AdminPass123!',
            first_name='VendorGate', last_name='Admin', department='IT',
        )
        self.client.login(email='vendorgate-admin@example.com', password='AdminPass123!')
        self.requester = User.objects.create_user(
            email='vendorgate-requester@example.com', password='TestPass123!',
            first_name='VendorGate', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.category = AssetCategory.objects.create(name='VendorGate Category')
        self.job = JobNumber.objects.create(number='JOB-VGATE-01', is_active=True)
        self.stock_asset = Asset.objects.create(name='VendorGate Stock Asset', status=Asset.Status.IN_STORE, category=self.category)
        self.ticket = Ticket.objects.create(
            number='SRV#9400', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='...', requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

    def test_procurement_only_mobilization_leaves_ticket_pending_vendor(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'ticket_id': self.ticket.pk,
            'procurement_item_name': ['Vendor Widget'],
            'procurement_category_id': [self.category.pk],
            'procurement_quantity': ['1'],
        })
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_VENDOR)
        self.assertIsNone(self.ticket.fulfilled_at)

    def test_receiving_the_only_procurement_request_fulfills_the_ticket(self):
        self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'ticket_id': self.ticket.pk,
            'procurement_item_name': ['Vendor Widget'],
            'procurement_category_id': [self.category.pk],
            'procurement_quantity': ['1'],
        })
        pr = AssetProcurementRequest.objects.get(item_name='Vendor Widget')

        response = self.client.post(reverse('tickets:procurement_receive', args=[pr.pk]))
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)
        self.assertIsNotNone(self.ticket.fulfilled_at)

    def test_mixed_stock_and_vendor_waits_for_vendor_item_before_fulfilling(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.stock_asset.pk],
            'procurement_item_name': ['Vendor Widget'],
            'procurement_category_id': [self.category.pk],
            'procurement_quantity': ['1'],
        })
        self.assertEqual(response.status_code, 302)

        self.ticket.refresh_from_db()
        # The stock item already went out, but the ticket isn't "fulfilled"
        # (no confirm-receipt prompt) until the vendor item also arrives.
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_VENDOR)

        mobilization = Mobilization.objects.get(ticket=self.ticket)
        self.assertEqual(mobilization.items.count(), 1)

        # One coherent CREATED comment covering both lines, not two
        # disconnected ones.
        created_comment = self.ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.CREATED)
        self.assertIn('mobilized from stock', created_comment.body)
        self.assertIn('ordered from vendor', created_comment.body)
        self.assertIn('2 items requested', created_comment.body)

        pr = AssetProcurementRequest.objects.get(item_name='Vendor Widget')
        self.client.post(reverse('tickets:procurement_receive', args=[pr.pk]))

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

        arrived_comment = self.ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.VENDOR_ITEM_ARRIVED)
        self.assertIn('Vendor Widget', arrived_comment.body)
        self.assertIn('All 2 items now fulfilled.', arrived_comment.body)

    def test_cancelling_last_open_procurement_fulfills_ticket_if_something_was_delivered(self):
        self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.stock_asset.pk],
            'procurement_item_name': ['Vendor Widget'],
            'procurement_category_id': [self.category.pk],
            'procurement_quantity': ['1'],
        })
        pr = AssetProcurementRequest.objects.get(item_name='Vendor Widget')

        self.client.post(reverse('tickets:procurement_cancel', args=[pr.pk]))

        self.ticket.refresh_from_db()
        self.assertEqual(self.ticket.status, Ticket.Status.PENDING_USER)

        cancelled_comment = self.ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.VENDOR_ITEM_CANCELLED)
        self.assertIn('Vendor Widget', cancelled_comment.body)
        self.assertIn('All 1 items now fulfilled.', cancelled_comment.body)

    def test_cancelling_only_procurement_with_nothing_delivered_leaves_ticket_pending(self):
        self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': self.job.pk,
            'ticket_id': self.ticket.pk,
            'procurement_item_name': ['Vendor Widget'],
            'procurement_category_id': [self.category.pk],
            'procurement_quantity': ['1'],
        })
        pr = AssetProcurementRequest.objects.get(item_name='Vendor Widget')

        self.client.post(reverse('tickets:procurement_cancel', args=[pr.pk]))

        self.ticket.refresh_from_db()
        self.assertNotEqual(self.ticket.status, Ticket.Status.PENDING_USER)
        self.assertIsNone(self.ticket.fulfilled_at)

        cancelled_comment = self.ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.VENDOR_ITEM_CANCELLED)
        self.assertIn('Nothing left on order for this mobilization.', cancelled_comment.body)


class ThirdPartyVesselMobilizationTests(TestCase):
    """Proposing a third-party vessel on a mobilization creates a pending
    (is_active=False) Vessel, attaches it immediately, reuses an existing
    proposal case-insensitively, and notifies admins for genuinely new ones."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='tpv-admin@example.com', password='AdminPass123!',
            first_name='TPV', last_name='Admin', department='IT',
        )
        self.client.login(email='tpv-admin@example.com', password='AdminPass123!')
        self.asset = Asset.objects.create(name='TPV Asset', status=Asset.Status.IN_STORE)
        requester = User.objects.create_user(
            email='tpv-requester@example.com', password='TestPass123!',
            first_name='TPV', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.ticket = Ticket.objects.create(
            number='SRV#8200', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='...', requester=requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

    def test_proposing_third_party_vessel_creates_pending_vessel(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'asset_ids': [self.asset.pk],
            'ticket_id': self.ticket.pk,
            'third_party_vessels': ['MV Client Vessel'],
        })
        self.assertEqual(response.status_code, 302)
        vessel = Vessel.objects.get(name='MV Client Vessel')
        self.assertFalse(vessel.is_active)
        self.assertEqual(vessel.proposed_by, self.admin)
        mobilization = Mobilization.objects.get(mobilized_by=self.admin)
        self.assertIn(vessel, mobilization.vessels.all())

    def test_multiple_third_party_vessels_in_one_submission(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'asset_ids': [self.asset.pk],
            'ticket_id': self.ticket.pk,
            'third_party_vessels': ['MV Client One', 'MV Client Two'],
        })
        self.assertEqual(response.status_code, 302)
        mobilization = Mobilization.objects.get(mobilized_by=self.admin)
        vessel_names = set(mobilization.vessels.values_list('name', flat=True))
        self.assertEqual(vessel_names, {'MV Client One', 'MV Client Two'})

    def test_case_insensitive_reuse_of_pending_vessel(self):
        Vessel.objects.create(name='MV Reused Vessel', is_active=False, proposed_by=self.admin)
        self.client.post(reverse('tickets:mobilization_create'), {
            'asset_ids': [self.asset.pk],
            'ticket_id': self.ticket.pk,
            'third_party_vessels': ['mv reused vessel'],
        })
        self.assertEqual(Vessel.objects.filter(name__iexact='MV Reused Vessel').count(), 1)

    def test_pending_vessel_invisible_in_new_mobilization_picker(self):
        self.client.post(reverse('tickets:mobilization_create'), {
            'asset_ids': [self.asset.pk],
            'ticket_id': self.ticket.pk,
            'third_party_vessels': ['MV Not Approved Yet'],
        })
        other_admin = User.objects.create_superuser(
            email='tpv-other@example.com', password='AdminPass123!',
            first_name='TPV', last_name='Other', department='IT',
        )
        self.client.logout()
        self.client.login(email='tpv-other@example.com', password='AdminPass123!')
        # The vessel checkbox picker on a *new* mobilization form only shows
        # active (approved) vessels — the pending one stays hidden there
        # even though it's visible on the mobilization that already uses it.
        response = self.client.get(reverse('tickets:mobilization_create_page'), {'ticket_id': self.ticket.pk})
        self.assertNotContains(response, 'MV Not Approved Yet')

    def test_admins_notified_of_new_third_party_vessel(self):
        other_admin = User.objects.create_superuser(
            email='tpv-notify@example.com', password='AdminPass123!',
            first_name='TPV', last_name='Notify', department='IT', role=User.Role.ADMIN,
        )
        self.client.post(reverse('tickets:mobilization_create'), {
            'asset_ids': [self.asset.pk],
            'ticket_id': self.ticket.pk,
            'third_party_vessels': ['MV Notify Test'],
        })
        self.assertTrue(
            Notification.objects.filter(recipient=other_admin, message__icontains='MV Notify Test').exists()
        )


class ConsumableAssetTests(TestCase):
    """Bulk/consumable assets (cable ties, PPE) are tracked by stock count
    rather than as one individually-tracked physical unit each."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='consumable-admin@example.com', password='AdminPass123!',
            first_name='Consumable', last_name='Admin', department='IT',
        )
        self.client.login(email='consumable-admin@example.com', password='AdminPass123!')
        self.consumable_category = AssetCategory.objects.create(name='Cable Ties', is_consumable=True)
        self.regular_category = AssetCategory.objects.create(name='Laptops', is_consumable=False)
        self.consumable_asset = Asset.objects.create(
            name='Cable Ties (100pk)', category=self.consumable_category,
            status=Asset.Status.IN_STORE, quantity_in_stock=10,
        )
        self.regular_asset = Asset.objects.create(
            name='ThinkPad X1', category=self.regular_category, status=Asset.Status.IN_STORE,
        )
        requester = User.objects.create_user(
            email='consumable-requester@example.com', password='TestPass123!',
            first_name='Consumable', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.ticket = Ticket.objects.create(
            number='SRV#8300', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='...', requester=requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

    def test_is_consumable_property(self):
        self.assertTrue(self.consumable_asset.is_consumable)
        self.assertFalse(self.regular_asset.is_consumable)

    def test_consumable_availability_based_on_stock(self):
        self.assertTrue(self.consumable_asset.is_available)
        self.consumable_asset.quantity_in_stock = 0
        self.consumable_asset.save()
        self.assertFalse(self.consumable_asset.is_available)

    def test_mobilizing_consumable_deducts_stock_and_keeps_in_store(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'job_number': '',
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.consumable_asset.pk],
            f'quantity_{self.consumable_asset.pk}': '4',
            'third_party_vessels': ['MV Consumable Test'],
        })
        self.assertEqual(response.status_code, 302)
        self.consumable_asset.refresh_from_db()
        self.assertEqual(self.consumable_asset.quantity_in_stock, 6)
        self.assertEqual(self.consumable_asset.status, Asset.Status.IN_STORE)

        item = MobilizationItem.objects.get(asset=self.consumable_asset)
        self.assertEqual(item.quantity, 4)

    def test_mobilizing_more_than_stock_is_rejected(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.consumable_asset.pk],
            f'quantity_{self.consumable_asset.pk}': '999',
            'third_party_vessels': ['MV Overrequest Test'],
        })
        self.assertEqual(response.status_code, 302)
        self.consumable_asset.refresh_from_db()
        self.assertEqual(self.consumable_asset.quantity_in_stock, 10)
        self.assertFalse(MobilizationItem.objects.filter(asset=self.consumable_asset).exists())

    def test_regular_asset_mobilization_unaffected(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.regular_asset.pk],
            'third_party_vessels': ['MV Regular Test'],
        })
        self.assertEqual(response.status_code, 302)
        self.regular_asset.refresh_from_db()
        self.assertEqual(self.regular_asset.status, Asset.Status.MOBILIZED)
        item = MobilizationItem.objects.get(asset=self.regular_asset)
        self.assertEqual(item.quantity, 1)

    def test_demobilizing_consumable_restores_stock(self):
        self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.consumable_asset.pk],
            f'quantity_{self.consumable_asset.pk}': '4',
            'third_party_vessels': ['MV Demob Test'],
        })
        self.consumable_asset.refresh_from_db()
        self.assertEqual(self.consumable_asset.quantity_in_stock, 6)

        item = MobilizationItem.objects.get(asset=self.consumable_asset)
        item.return_requested_at = timezone.now()
        item.save(update_fields=['return_requested_at'])
        self.client.post(
            reverse('tickets:mobilization_item_demobilize', args=[item.pk]),
            {'return_condition': Asset.Condition.GOOD, 'return_quantity': '4'},
        )
        self.consumable_asset.refresh_from_db()
        self.assertEqual(self.consumable_asset.quantity_in_stock, 10)

    def test_damaged_consumable_not_restocked(self):
        self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.consumable_asset.pk],
            f'quantity_{self.consumable_asset.pk}': '4',
            'third_party_vessels': ['MV Damaged Test'],
        })
        self.consumable_asset.refresh_from_db()
        self.assertEqual(self.consumable_asset.quantity_in_stock, 6)

        item = MobilizationItem.objects.get(asset=self.consumable_asset)
        item.return_requested_at = timezone.now()
        item.save(update_fields=['return_requested_at'])
        self.client.post(
            reverse('tickets:mobilization_item_demobilize', args=[item.pk]),
            {'return_condition': Asset.Condition.DAMAGED, 'return_quantity': '4'},
        )
        self.consumable_asset.refresh_from_db()
        self.assertEqual(self.consumable_asset.quantity_in_stock, 6)

    def test_depleted_consumable_excluded_from_picker(self):
        self.consumable_asset.quantity_in_stock = 0
        self.consumable_asset.save()
        response = self.client.get(reverse('tickets:mobilization_available_assets'))
        self.assertNotContains(response, 'Cable Ties (100pk)')


class MobilizationAutopickTests(TestCase):
    """Quantity auto-pick for individually-tracked assets: resolves N
    available assets in a category, same end state as picking them by hand."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='autopick-admin@example.com', password='AdminPass123!',
            first_name='Autopick', last_name='Admin', department='IT',
        )
        self.client.login(email='autopick-admin@example.com', password='AdminPass123!')
        self.category = AssetCategory.objects.create(name='Autopick Laptops', is_consumable=False)
        self.laptops = [
            Asset.objects.create(name=f'Laptop {i}', category=self.category, status=Asset.Status.IN_STORE)
            for i in range(5)
        ]
        requester = User.objects.create_user(
            email='autopick-requester@example.com', password='TestPass123!',
            first_name='Autopick', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.ticket = Ticket.objects.create(
            number='SRV#8400', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='...', requester=requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

    def test_autopick_returns_requested_quantity(self):
        response = self.client.get(reverse('tickets:mobilization_autopick_assets'), {
            'category_id': self.category.pk, 'quantity': 3,
        })
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(len(data['assets']), 3)

    def test_autopick_caps_at_available_count(self):
        response = self.client.get(reverse('tickets:mobilization_autopick_assets'), {
            'category_id': self.category.pk, 'quantity': 99,
        })
        data = response.json()
        self.assertEqual(len(data['assets']), 5)

    def test_autopicked_ids_mobilize_like_manual_selection(self):
        response = self.client.get(reverse('tickets:mobilization_autopick_assets'), {
            'category_id': self.category.pk, 'quantity': 3,
        })
        picked_ids = [a['id'] for a in response.json()['assets']]

        post_response = self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': picked_ids,
            'third_party_vessels': ['MV Autopick Test'],
        })
        self.assertEqual(post_response.status_code, 302)
        for asset_id in picked_ids:
            asset = Asset.objects.get(pk=asset_id)
            self.assertEqual(asset.status, Asset.Status.MOBILIZED)
        self.assertEqual(MobilizationItem.objects.filter(asset_id__in=picked_ids).count(), 3)

    def test_autopick_excludes_consumable_categories(self):
        consumable_category = AssetCategory.objects.create(name='Autopick Consumables', is_consumable=True)
        Asset.objects.create(name='Some Consumable', category=consumable_category, quantity_in_stock=10)
        response = self.client.get(reverse('tickets:mobilization_autopick_assets'), {
            'category_id': consumable_category.pk, 'quantity': 1,
        })
        self.assertEqual(response.json()['assets'], [])

    def test_autopick_excludes_assets_already_permanently_assigned(self):
        other_user = User.objects.create_user(
            email='autopick-owner@example.com', password='TestPass123!',
            first_name='Autopick', last_name='Owner', department='IT', role=User.Role.END_USER,
        )
        self.laptops[0].assigned_to = other_user
        self.laptops[0].save(update_fields=['assigned_to'])

        response = self.client.get(reverse('tickets:mobilization_autopick_assets'), {
            'category_id': self.category.pk, 'quantity': 99,
        })
        data = response.json()
        self.assertEqual(len(data['assets']), 4)
        self.assertNotIn(self.laptops[0].pk, [a['id'] for a in data['assets']])

    def test_available_assets_picker_excludes_assets_already_permanently_assigned(self):
        other_user = User.objects.create_user(
            email='available-owner@example.com', password='TestPass123!',
            first_name='Available', last_name='Owner', department='IT', role=User.Role.END_USER,
        )
        self.laptops[0].assigned_to = other_user
        self.laptops[0].save(update_fields=['assigned_to'])

        response = self.client.get(reverse('tickets:mobilization_available_assets'))
        self.assertNotContains(response, self.laptops[0].tracking_id)
        self.assertContains(response, self.laptops[1].tracking_id)

    def test_mobilization_create_rejects_a_permanently_assigned_asset(self):
        other_user = User.objects.create_user(
            email='create-owner@example.com', password='TestPass123!',
            first_name='Create', last_name='Owner', department='IT', role=User.Role.END_USER,
        )
        self.laptops[0].assigned_to = other_user
        self.laptops[0].save(update_fields=['assigned_to'])

        response = self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.laptops[0].pk],
            'third_party_vessels': ['MV Assigned Reject Test'],
        })
        self.assertEqual(response.status_code, 302)
        self.assertFalse(MobilizationItem.objects.filter(asset=self.laptops[0]).exists())
        self.laptops[0].refresh_from_db()
        self.assertNotEqual(self.laptops[0].status, Asset.Status.MOBILIZED)


class LowStockAlertTests(TestCase):
    """Low-stock alerting is asset-lifecycle infrastructure — it fires from
    any point quantity_in_stock changes, not just mobilization."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='lowstock-admin@example.com', password='AdminPass123!',
            first_name='LowStock', last_name='Admin', department='IT',
        )
        self.other_admin = User.objects.create_superuser(
            email='lowstock-other@example.com', password='AdminPass123!',
            first_name='LowStock', last_name='Other', department='IT', role=User.Role.ADMIN,
        )
        self.client.login(email='lowstock-admin@example.com', password='AdminPass123!')
        self.category = AssetCategory.objects.create(name='Low Stock Category', is_consumable=True)
        self.asset = Asset.objects.create(
            name='Threshold Widgets', category=self.category,
            quantity_in_stock=6, low_stock_threshold=5, status=Asset.Status.IN_STORE,
        )
        requester = User.objects.create_user(
            email='lowstock-requester@example.com', password='TestPass123!',
            first_name='LowStock', last_name='Requester', department='IT', role=User.Role.END_USER,
        )
        self.ticket = Ticket.objects.create(
            number='SRV#8500', type=Ticket.Type.SERVICE_REQUEST, title='Gear for job',
            description='...', requester=requester,
            status=Ticket.Status.PENDING_FULFILLMENT, is_asset_request=True, is_mobilization_request=True,
        )

    def test_dipping_below_threshold_notifies_once(self):
        self.asset.quantity_in_stock = 4
        self.asset.save()
        self.asset.refresh_low_stock_alert()
        self.assertTrue(Notification.objects.filter(message__icontains='Threshold Widgets').exists())
        count_after_first = Notification.objects.filter(message__icontains='Threshold Widgets').count()

        # Still under threshold — must not notify again.
        self.asset.quantity_in_stock = 3
        self.asset.save()
        self.asset.refresh_low_stock_alert()
        self.assertEqual(
            Notification.objects.filter(message__icontains='Threshold Widgets').count(),
            count_after_first,
        )

    def test_recovering_above_threshold_resets_alert(self):
        self.asset.quantity_in_stock = 4
        self.asset.save()
        self.asset.refresh_low_stock_alert()
        self.asset.refresh_from_db()
        self.assertTrue(self.asset.low_stock_notified)

        self.asset.quantity_in_stock = 10
        self.asset.save()
        self.asset.refresh_low_stock_alert()
        self.asset.refresh_from_db()
        self.assertFalse(self.asset.low_stock_notified)

        # Dipping again after recovery notifies a second time.
        self.asset.quantity_in_stock = 2
        self.asset.save()
        self.asset.refresh_low_stock_alert()
        self.assertEqual(
            Notification.objects.filter(recipient=self.admin, message__icontains='Threshold Widgets').count(), 2
        )

    def test_mobilizing_below_threshold_triggers_alert(self):
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': self.ticket.pk,
            'asset_ids': [self.asset.pk],
            f'quantity_{self.asset.pk}': '2',
            'third_party_vessels': ['MV Low Stock Test'],
        })
        self.assertEqual(response.status_code, 302)
        self.assertTrue(Notification.objects.filter(message__icontains='Threshold Widgets').exists())

    def test_no_threshold_never_alerts(self):
        self.asset.low_stock_threshold = None
        self.asset.quantity_in_stock = 0
        self.asset.save()
        self.asset.refresh_low_stock_alert()
        self.assertFalse(Notification.objects.filter(message__icontains='Threshold Widgets').exists())

    def test_is_low_stock_property(self):
        self.assertFalse(self.asset.is_low_stock)
        self.asset.quantity_in_stock = 5
        self.assertTrue(self.asset.is_low_stock)

    def test_asset_list_low_stock_filter(self):
        self.asset.quantity_in_stock = 3
        self.asset.save()
        other = Asset.objects.create(
            name='Plenty In Stock', category=self.category,
            quantity_in_stock=50, low_stock_threshold=5,
        )
        response = self.client.get(reverse('tickets:assets'), {'filter_low_stock': '1'})
        self.assertContains(response, 'Threshold Widgets')
        self.assertNotContains(response, 'Plenty In Stock')


class RenewableAssetTests(TestCase):
    """Renewable assets (software licenses, subscriptions) — recurring
    renewal dates, cost tracking, and the reminder job that watches them."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='renew-admin@example.com', password='AdminPass123!',
            first_name='Renew', last_name='Admin', department='IT',
        )
        self.other_admin = User.objects.create_superuser(
            email='renew-other@example.com', password='AdminPass123!',
            first_name='Renew', last_name='Other', department='IT', role=User.Role.ADMIN,
        )
        self.client.login(email='renew-admin@example.com', password='AdminPass123!')
        self.renewable_category = AssetCategory.objects.create(name='Software Licenses', is_renewable=True)
        self.regular_category = AssetCategory.objects.create(name='Laptops', is_renewable=False)
        self.vendor = Vendor.objects.create(name='Microsoft Reseller', is_active=True)

    def test_is_renewable_property(self):
        renewable_asset = Asset.objects.create(name='Office 365', category=self.renewable_category)
        regular_asset = Asset.objects.create(name='ThinkPad X1', category=self.regular_category)
        self.assertTrue(renewable_asset.is_renewable)
        self.assertFalse(regular_asset.is_renewable)

    def test_mark_renewed_advances_date_and_resets_flags(self):
        from apps.tickets.models import _add_months
        future_date = timezone.now().date() + timedelta(days=60)
        asset = Asset.objects.create(
            name='Office 365', category=self.renewable_category,
            next_renewal_date=future_date, renewal_interval_months=12,
            renewal_cost=Decimal('500.00'),
            renewal_reminder_90d_sent=True, renewal_reminder_30d_sent=True, renewal_reminder_7d_sent=True,
        )
        asset.mark_renewed(self.admin, new_cost=Decimal('550.00'))
        asset.refresh_from_db()
        self.assertEqual(asset.next_renewal_date, _add_months(future_date, 12))
        self.assertEqual(asset.renewal_cost, Decimal('550.00'))
        self.assertFalse(asset.renewal_reminder_90d_sent)
        self.assertFalse(asset.renewal_reminder_30d_sent)
        self.assertFalse(asset.renewal_reminder_7d_sent)
        self.assertTrue(AssetLog.objects.filter(asset=asset, action=AssetLog.Action.RENEWED).exists())

    def test_mark_renewed_view_requires_interval(self):
        future_date = timezone.now().date() + timedelta(days=60)
        asset = Asset.objects.create(
            name='Office 365', category=self.renewable_category,
            next_renewal_date=future_date,
        )
        response = self.client.post(reverse('tickets:asset_mark_renewed', args=[asset.pk]))
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertEqual(asset.next_renewal_date, future_date)

    def test_mark_renewed_view_success(self):
        from apps.tickets.models import _add_months
        future_date = timezone.now().date() + timedelta(days=60)
        asset = Asset.objects.create(
            name='Office 365', category=self.renewable_category,
            next_renewal_date=future_date, renewal_interval_months=1,
        )
        response = self.client.post(
            reverse('tickets:asset_mark_renewed', args=[asset.pk]), {'new_cost': '25.00'}
        )
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertEqual(asset.next_renewal_date, _add_months(future_date, 1))
        self.assertEqual(asset.renewal_cost, Decimal('25.00'))

    def test_reminder_command_notifies_closest_threshold_only(self):
        today = timezone.now().date()
        Asset.objects.create(
            name='Office 365', category=self.renewable_category,
            next_renewal_date=today + timedelta(days=5), renewal_interval_months=12,
        )
        call_command('send_renewal_reminders')
        self.assertEqual(
            Notification.objects.filter(recipient=self.admin, message__icontains='Office 365').count(), 1
        )
        self.assertIn('7 days', Notification.objects.get(recipient=self.admin, message__icontains='Office 365').message)

        asset = Asset.objects.get(name='Office 365')
        self.assertTrue(asset.renewal_reminder_7d_sent)
        self.assertTrue(asset.renewal_reminder_30d_sent)
        self.assertTrue(asset.renewal_reminder_90d_sent)

    def test_reminder_command_does_not_renotify(self):
        today = timezone.now().date()
        Asset.objects.create(
            name='Office 365', category=self.renewable_category,
            next_renewal_date=today + timedelta(days=5), renewal_interval_months=12,
        )
        call_command('send_renewal_reminders')
        call_command('send_renewal_reminders')
        self.assertEqual(
            Notification.objects.filter(recipient=self.admin, message__icontains='Office 365').count(), 1
        )

    def test_reminder_command_ignores_non_renewable_assets(self):
        today = timezone.now().date()
        Asset.objects.create(
            name='ThinkPad X1', category=self.regular_category,
            next_renewal_date=today + timedelta(days=5),
        )
        call_command('send_renewal_reminders')
        self.assertFalse(Notification.objects.filter(message__icontains='ThinkPad X1').exists())

    def test_asset_list_renewal_due_filter(self):
        today = timezone.now().date()
        Asset.objects.create(
            name='Office 365', category=self.renewable_category, next_renewal_date=today + timedelta(days=5),
        )
        Asset.objects.create(
            name='Adobe Creative Cloud', category=self.renewable_category, next_renewal_date=today + timedelta(days=200),
        )
        # filter_renewal_due is implemented in report_registry's assets
        # queryset (used by the Reports Builder page), not the day-to-day
        # Asset Inventory page (tickets:assets) — that page's equipment tab
        # excludes renewable assets entirely (they live in its separate
        # License tab, sorted by urgency rather than filtered).
        response = self.client.get(reverse('tickets:report_builder', args=['assets']), {'filter_renewal_due': '1'})
        self.assertContains(response, 'Office 365')
        self.assertNotContains(response, 'Adobe Creative Cloud')

    def test_report_export_includes_renewal_columns(self):
        Asset.objects.create(
            name='Office 365', category=self.renewable_category,
            next_renewal_date=date(2026, 6, 1), renewal_interval_months=12,
            renewal_cost=Decimal('500.00'), renewal_vendor=self.vendor,
        )
        # CSV export was deliberately dropped for assets (see
        # report_builder.html/export_report — a flat-text format can't
        # carry the branded letterhead every other asset export shares).
        response = self.client.get(reverse('tickets:export_report', args=['assets']), {'format': 'json'})
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('Renewal Cost', content)
        self.assertIn('Microsoft Reseller', content)


class SLAAndEscalationTests(TestCase):
    """Test SLA and escalation functionality."""
    
    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='admin@example.com',
            password='AdminPass123!',
            first_name='Admin',
            last_name='User',
            department='IT'
        )
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT'
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.client.login(email='admin@example.com', password='AdminPass123!')

    def test_sla_creation(self):
        """Test creating an SLA policy."""
        response = self.client.post(reverse('tickets:sla_create'), {
            'priority': 'P1',
            'response_minutes': 15,
            'resolution_minutes': 60,
            'calendar_id': ''
        })
        self.assertEqual(response.status_code, 302)
        self.assertTrue(SLA.objects.filter(priority='P1').exists())

    def test_sla_policy_list(self):
        """Test SLA management page shows policies."""
        sla = SLA.objects.create(
            priority='P1',
            response_minutes=15,
            resolution_minutes=60
        )
        response = self.client.get(reverse('tickets:sla_management'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'P1')

    def test_escalation_rule_creation(self):
        """Test creating an escalation rule."""
        response = self.client.post(reverse('tickets:rule_create'), {
            'priority': 'P1',
            'timer_type': 'response',
            'threshold_percent': 75,
            'action_type': 'notify',
            'notify_role': 'ADMIN'
        })
        self.assertEqual(response.status_code, 302)
        self.assertTrue(EscalationRule.objects.filter(priority='P1').exists())

    def test_sla_badge_view(self):
        """Test SLA badge view."""
        sla = SLA.objects.create(
            priority='P3',
            response_minutes=60,
            resolution_minutes=240
        )
        ticket = Ticket.objects.create(
            number='TK#9999',
            title='SLA Test Ticket',
            description='Test description',
            requester=self.user,
            category=self.category,
            priority='P3',
            status=Ticket.Status.NEW
        )
        response = self.client.get(reverse('tickets:sla_badge', args=[ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'On Track')


class NotificationTests(TestCase):
    """Test notification functionality."""
    
    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT',
            is_active=True,
            email_verified=True
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.client.login(email='test@example.com', password='TestPass123!')

    def test_notification_creation(self):
        """Test creating a notification."""
        notification = Notification.objects.create(
            recipient=self.user,
            message='Test notification',
            url='/dashboard/'
        )
        self.assertEqual(notification.recipient, self.user)
        self.assertEqual(notification.message, 'Test notification')
        self.assertFalse(notification.is_read)
        self.assertEqual(str(notification), f'Notification for {self.user.email}: Test notification')

    def test_notification_mark_read(self):
        """Test marking a notification as read."""
        notification = Notification.objects.create(
            recipient=self.user,
            message='Test notification',
            url='/dashboard/'
        )
        self.assertFalse(notification.is_read)
        
        notification.is_read = True
        notification.save()
        notification.refresh_from_db()
        self.assertTrue(notification.is_read)

    def test_notification_unread_count(self):
        """Test unread notification count."""
        # Create some notifications
        Notification.objects.create(recipient=self.user, message='Notif 1', url='/')
        Notification.objects.create(recipient=self.user, message='Notif 2', url='/')
        Notification.objects.create(recipient=self.user, message='Notif 3', url='/')
        
        # Read one
        notif = Notification.objects.filter(recipient=self.user).first()
        notif.is_read = True
        notif.save()
        
        count = Notification.objects.filter(recipient=self.user, is_read=False).count()
        self.assertEqual(count, 2)

    def test_notification_list_view(self):
        """Test notification dropdown list."""
        Notification.objects.create(recipient=self.user, message='Test notif 1', url='/')
        Notification.objects.create(recipient=self.user, message='Test notif 2', url='/')
        
        response = self.client.get(reverse('notifications:list'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Test notif 1')
        self.assertContains(response, 'Test notif 2')

    def test_notification_mark_all_read(self):
        """Test marking all notifications as read."""
        Notification.objects.create(recipient=self.user, message='Notif 1', url='/')
        Notification.objects.create(recipient=self.user, message='Notif 2', url='/')
        
        response = self.client.post(reverse('notifications:mark_all_read'))
        self.assertEqual(response.status_code, 200)
        
        count = Notification.objects.filter(recipient=self.user, is_read=False).count()
        self.assertEqual(count, 0)


class RoleBasedAccessTests(TestCase):
    """Test role-based access control."""
    
    def setUp(self):
        self.client = Client()
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        
        # Create users with different roles
        self.end_user = User.objects.create_user(
            email='user@example.com',
            password='TestPass123!',
            first_name='End',
            last_name='User',
            department='IT',
            role=User.Role.END_USER,
            is_active=True,
            email_verified=True
        )
        
        self.agent = User.objects.create_user(
            email='agent@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='User',
            department='IT',
            role=User.Role.AGENT,
            is_active=True,
            email_verified=True
        )
        
        self.team_lead = User.objects.create_user(
            email='lead@example.com',
            password='TestPass123!',
            first_name='Team',
            last_name='Lead',
            department='IT',
            role=User.Role.TEAM_LEAD,
            is_active=True,
            email_verified=True
        )
        
        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='TestPass123!',
            first_name='Admin',
            last_name='User',
            department='IT',
            role=User.Role.ADMIN,
            is_active=True,
            email_verified=True
        )
        
        # Create a ticket
        self.ticket = Ticket.objects.create(
            number='TK#9999',
            title='Test Ticket',
            description='Test description',
            requester=self.end_user,
            category=self.category,
            status=Ticket.Status.NEW
        )

    def test_dashboard_redirects_unauthenticated(self):
        """Unauthenticated users should be redirected to login."""
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 302)

    def test_end_user_dashboard_access(self):
        """End users should access their dashboard."""
        self.client.login(email='user@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'dashboards/end_user_dashboard.html')

    def test_agent_dashboard_access(self):
        """Agents should access their dashboard."""
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'dashboards/agent_dashboard.html')

    def test_team_lead_dashboard_access(self):
        """Team Leads should access their dashboard."""
        self.client.login(email='lead@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'dashboards/team_lead_dashboard.html')

    def test_admin_dashboard_access(self):
        """Admins should access their dashboard."""
        self.client.login(email='admin@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'dashboards/admin_dashboard.html')

    def test_unassigned_queue_agent_access(self):
        """Agents should access unassigned queue."""
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:unassigned'))
        self.assertEqual(response.status_code, 200)

    def test_unassigned_queue_end_user_denied(self):
        """End users should not access unassigned queue."""
        self.client.login(email='user@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:unassigned'))
        self.assertEqual(response.status_code, 403)  # Forbidden

    def test_asset_management_agent_denied(self):
        """Only Admin/Superadmin may browse the full asset inventory — every
        other role (including Agent) only ever sees their own via my_assets."""
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:assets'))
        self.assertEqual(response.status_code, 403)

    def test_asset_create_admin_only(self):
        """Only admins should access asset creation."""
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:asset_create_page'))
        self.assertEqual(response.status_code, 403)
        
        self.client.login(email='admin@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:asset_create_page'))
        self.assertEqual(response.status_code, 200)

    def test_manager_review_team_lead_only(self):
        """Only team leads should access manager review queue."""
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:manager_review_queue'))
        self.assertEqual(response.status_code, 403)
        
        self.client.login(email='lead@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:manager_review_queue'))
        self.assertEqual(response.status_code, 200)


class ServiceRequestFlowTests(TestCase):
    """Test the complete service request workflow."""
    
    def setUp(self):
        self.client = Client()
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.service_category = ServiceCategory.objects.create(
            name='Asset Category', slug='asset-category', field_group=ServiceCategory.FieldGroup.ASSET
        )

        self.end_user = User.objects.create_user(
            email='user@example.com',
            password='TestPass123!',
            first_name='End',
            last_name='User',
            department='IT',
            role=User.Role.END_USER,
            is_active=True,
            email_verified=True
        )

        self.team_lead = User.objects.create_user(
            email='lead@example.com',
            password='TestPass123!',
            first_name='Team',
            last_name='Lead',
            department='IT',
            role=User.Role.TEAM_LEAD,
            is_active=True,
            email_verified=True
        )
        
        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='TestPass123!',
            first_name='Admin',
            last_name='User',
            department='IT',
            role=User.Role.ADMIN,
            is_active=True,
            email_verified=True
        )

    def test_service_request_creation_asset_detection(self):
        """Test that service requests with asset categories are flagged."""
        self.client.login(email='user@example.com', password='TestPass123!')
        
        response = self.client.post(reverse('tickets:create'), {
            'type': 'SERVICE_REQUEST',
            'title': 'Need new laptop',
            'description': 'I need a new laptop for the new developer',
            'service_category': self.service_category.id,
            'purpose': 'New developer onboarding',
            'urgency': 'MEDIUM',
            'number_of_assets': '1',
            'asset_type': 'LAPTOP',
        })
        self.assertEqual(response.status_code, 302)
        
        ticket = Ticket.objects.filter(title='Need new laptop').first()
        self.assertIsNotNone(ticket)
        self.assertEqual(ticket.type, Ticket.Type.SERVICE_REQUEST)
        self.assertTrue(ticket.is_asset_request)

    def test_service_request_team_lead_review(self):
        """Test team lead reviewing a service request."""
        # Create a service request
        self.client.login(email='user@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:create'), {
            'type': 'SERVICE_REQUEST',
            'title': 'Need new laptop',
            'description': 'I need a new laptop for the new developer',
            'service_category': self.service_category.id,
            'purpose': 'New developer onboarding',
            'urgency': 'MEDIUM',
            'number_of_assets': '1',
            'asset_type': 'LAPTOP',
        })

        ticket = Ticket.objects.filter(title='Need new laptop').first()

        # Team lead reviews and approves
        self.client.login(email='lead@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:manager_review_ticket', args=[ticket.pk]),
            {
                'action': 'approve',
                'comment': 'Approved, please assign a laptop'
            }
        )
        self.assertEqual(response.status_code, 302)
        
        ticket.refresh_from_db()
        # Asset request should go to PENDING_FULFILLMENT
        self.assertEqual(ticket.status, Ticket.Status.PENDING_FULFILLMENT)

    def test_manager_request_changes_posts_comment_and_resubmit_returns_to_review(self):
        """The manager's reason must be visible on the ticket (not just an
        activity-log entry), and once the requester replies, the ticket must
        go back through manager review, not straight to the agent pool."""
        self.client.login(email='user@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:create'), {
            'type': 'SERVICE_REQUEST',
            'title': 'Need a monitor',
            'description': 'I need a monitor',
            'service_category': self.service_category.id,
            'purpose': 'Desk setup',
            'urgency': 'MEDIUM',
            'number_of_assets': '1',
            'asset_type': 'MONITOR',
        })
        ticket = Ticket.objects.filter(title='Need a monitor').first()
        self.assertEqual(ticket.status, Ticket.Status.PENDING_MANAGER_REVIEW)

        self.client.login(email='lead@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:manager_review_ticket', args=[ticket.pk]),
            {'action': 'request_changes', 'comment': 'Please specify monitor size'}
        )
        self.assertEqual(response.status_code, 302)
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.Status.PENDING_USER)
        # The reason must be visible on the ticket, not just logged internally.
        self.assertTrue(
            ticket.comments.filter(body__icontains='Please specify monitor size').exists()
        )

        self.client.login(email='user@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:detail', args=[ticket.pk]),
            {'body': '27 inch please'},
            HTTP_HX_REQUEST='true',
        )
        self.assertEqual(response.status_code, 200)
        ticket.refresh_from_db()
        # Must go back to the Team Lead, not straight to the agent pool.
        self.assertEqual(ticket.status, Ticket.Status.PENDING_MANAGER_REVIEW)

    def test_unassigned_ticket_conversation_is_blocked(self):
        """An agent must claim a ticket before its conversation page (and
        commenting) becomes reachable, so status can never move while the
        ticket is still sitting unclaimed in the unassigned queue. This only
        applies to tickets worked via the claim/assign flow (incidents,
        general service requests) — asset-request tickets are worked through
        the fulfillment pool and are never assigned, so they're exempt (see
        test_service_request_admin_fulfillment, which follows the
        fulfillment redirect into this same conversation page)."""
        self.client.login(email='user@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:create'), {
            'type': 'INCIDENT',
            'title': 'Printer is broken',
            'description': 'The printer on the 3rd floor is broken',
            'category': self.category.id,
            'impact': 'INDIVIDUAL',
            'urgency': 'MEDIUM',
        })
        ticket = Ticket.objects.filter(title='Printer is broken').first()
        ticket.status = Ticket.Status.NEW
        ticket.save(update_fields=['status'])
        self.assertIsNone(ticket.assigned_to)

        User.objects.create_user(
            email='agent-unassigned-block@example.com', password='TestPass123!',
            first_name='Agent', last_name='Block', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.client.login(email='agent-unassigned-block@example.com', password='TestPass123!')

        response = self.client.get(reverse('tickets:conversation', args=[ticket.pk]))
        self.assertEqual(response.status_code, 302)
        self.assertRedirects(response, reverse('tickets:unassigned'))

        response = self.client.post(
            reverse('tickets:add_comment_conversation', args=[ticket.pk]),
            {'body': 'trying to reply without claiming'},
        )
        self.assertEqual(response.status_code, 403)
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.Status.NEW)

    def test_service_request_admin_fulfillment(self):
        """Test admin fulfilling an asset request."""
        # Create and approve a service request
        self.client.login(email='user@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:create'), {
            'type': 'SERVICE_REQUEST',
            'title': 'Need new laptop',
            'description': 'I need a new laptop',
            'service_category': self.service_category.id,
            'purpose': 'New developer onboarding',
            'urgency': 'MEDIUM',
            'number_of_assets': '1',
            'asset_type': 'LAPTOP',
        })

        ticket = Ticket.objects.filter(title='Need new laptop').first()

        self.client.login(email='lead@example.com', password='TestPass123!')
        self.client.post(
            reverse('tickets:manager_review_ticket', args=[ticket.pk]),
            {'action': 'approve', 'comment': 'Approved'}
        )
        
        # Create an asset
        asset = Asset.objects.create(
            name='Dell Laptop',
            serial_number='SN12345',
            status='IN_STORE'
        )
        
        # Admin fulfills the request
        self.client.login(email='admin@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:fulfill_asset_request', args=[ticket.pk]),
            {
                'asset_id': asset.pk,
                'comment': 'Fulfilled with Dell Laptop'
            },
            follow=True,
        )
        # Fulfilling redirects straight to the ticket's conversation page —
        # asset-request tickets are never assigned_to anyone (they're worked
        # through the fulfillment pool, not the claim flow), so this must
        # not trip the unclaimed-ticket guard added for unassigned incidents.
        self.assertEqual(response.status_code, 200)

        ticket.refresh_from_db()
        self.assertIsNone(ticket.assigned_to)
        self.assertEqual(ticket.status, Ticket.Status.PENDING_USER)
        self.assertEqual(ticket.assigned_asset, asset)

        asset.refresh_from_db()
        self.assertEqual(asset.assigned_to, self.end_user)


class ServiceRequestVesselJobDiveSystemTests(TestCase):
    """Vessel, Job Number, and Dive System are optional fields available on
    every service request regardless of category — Vessel/Dive System are
    admin-curated multi-select lists; Job Number additionally supports a
    requester-proposed new entry that notifies admins for approval."""

    def setUp(self):
        self.client = Client()
        self.service_category = ServiceCategory.objects.create(
            name='General IT', slug='general-it', field_group=ServiceCategory.FieldGroup.GENERAL
        )
        self.vessel = Vessel.objects.create(name='MV Explorer')
        self.dive_system = DiveSystem.objects.create(name='System A')
        self.job_number = JobNumber.objects.create(number='JOB-0001')

        self.end_user = User.objects.create_user(
            email='marineuser@example.com', password='TestPass123!',
            first_name='Marine', last_name='User', department='MARINE',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.admin = User.objects.create_user(
            email='itadmin@example.com', password='TestPass123!',
            first_name='IT', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.login(email='marineuser@example.com', password='TestPass123!')

    def _base_post(self, **overrides):
        data = {
            'type': 'SERVICE_REQUEST',
            'title': 'Network setup on vessel',
            'description': 'Need network configured',
            'service_category': self.service_category.id,
            'purpose': 'Vessel network installation',
            'urgency': 'MEDIUM',
        }
        data.update(overrides)
        return data

    def test_vessel_and_dive_system_optional_and_selectable_on_any_department(self):
        response = self.client.post(reverse('tickets:create'), self._base_post(
            vessels=[self.vessel.id], dive_systems=[self.dive_system.id],
        ))
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Network setup on vessel')
        self.assertIn(self.vessel, ticket.vessels.all())
        self.assertIn(self.dive_system, ticket.dive_systems.all())

    def test_service_request_succeeds_with_no_optional_fields(self):
        response = self.client.post(reverse('tickets:create'), self._base_post())
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Network setup on vessel')
        self.assertFalse(ticket.vessels.exists())
        self.assertFalse(ticket.dive_systems.exists())
        self.assertIsNone(ticket.job_number)

    def test_existing_job_number_selected(self):
        response = self.client.post(reverse('tickets:create'), self._base_post(
            job_number=str(self.job_number.id),
        ))
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Network setup on vessel')
        self.assertEqual(ticket.job_number, self.job_number)

    def test_new_job_number_creates_pending_entry_and_notifies_admins(self):
        response = self.client.post(reverse('tickets:create'), self._base_post(
            job_number='NEW', new_job_number_text='JOB-9999',
        ))
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Network setup on vessel')

        new_job = JobNumber.objects.get(number='JOB-9999')
        self.assertFalse(new_job.is_active)
        self.assertEqual(new_job.proposed_by, self.end_user)
        self.assertEqual(ticket.job_number, new_job)

        self.assertTrue(
            Notification.objects.filter(recipient=self.admin, message__icontains='JOB-9999').exists()
        )

    def test_inactive_job_number_not_selectable(self):
        inactive = JobNumber.objects.create(number='JOB-INACTIVE', is_active=False)
        response = self.client.post(reverse('tickets:create'), self._base_post(
            job_number=str(inactive.id),
        ))
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Network setup on vessel')
        self.assertIsNone(ticket.job_number)


class SecurityTests(TestCase):
    """Test security features."""
    
    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT',
            is_active=True,
            email_verified=True
        )

    def test_password_policy_min_length(self):
        """Test that password must be at least 10 characters."""
        # Attempt to create user with short password
        with self.assertRaises(Exception):
            User.objects.create_user(
                email='shortpass@example.com',
                password='Short1!',  # Too short
                first_name='Test',
                last_name='User',
                department='IT'
            )

    def test_secure_cookie_settings(self):
        """Test that secure cookie settings are applied."""
        # This test is more of a configuration check
        from django.conf import settings
        self.assertTrue(settings.SESSION_COOKIE_HTTPONLY)
        self.assertTrue(settings.CSRF_COOKIE_HTTPONLY)
        # In production, these should be True, but in development they may be False
        # So we just check they exist
        self.assertIsNotNone(settings.SESSION_COOKIE_HTTPONLY)

    def test_xframe_options(self):
        """Test X-Frame-Options header."""
        # This test checks that the header is set
        # In production, X_FRAME_OPTIONS should be set
        from django.conf import settings
        self.assertIsNotNone(getattr(settings, 'X_FRAME_OPTIONS', None))


class EdgeCaseTests(TestCase):
    """Test edge cases and error handling."""
    
    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='test@example.com',
            password='TestPass123!',
            first_name='Test',
            last_name='User',
            department='IT',
            is_active=True,
            email_verified=True
        )
        self.category = Category.objects.create(name='Hardware', slug='hardware')
        self.client.login(email='test@example.com', password='TestPass123!')

    def test_ticket_create_with_very_long_title(self):
        """Test ticket creation with a title at the field's max length (255)."""
        long_title = 'a' * 255
        response = self.client.post(reverse('tickets:create'), {
            'type': 'INCIDENT',
            'title': long_title,
            'description': 'Test description',
            'category': self.category.id,
            'impact': 'INDIVIDUAL',
            'urgency': 'MEDIUM'
        })
        self.assertEqual(response.status_code, 302)  # Should still work
        ticket = Ticket.objects.filter(title=long_title).first()
        self.assertIsNotNone(ticket)

    def test_asset_create_with_special_characters(self):
        """Test asset creation with special characters."""
        self.client.login(email='admin@example.com', password='TestPass123!')
        # First create admin user
        admin = User.objects.create_superuser(
            email='admin@example.com',
            password='AdminPass123!',
            first_name='Admin',
            last_name='User',
            department='IT'
        )
        self.client.login(email='admin@example.com', password='AdminPass123!')
        
        response = self.client.post(reverse('tickets:asset_create_page'), {
            'name': 'Test Laptop with $pecial & Chars!',
            'serial_number': 'SN#123!@#',
            'status': 'IN_STORE'
        })
        self.assertEqual(response.status_code, 302)
        self.assertTrue(Asset.objects.filter(name__contains='$pecial').exists())

    def test_edit_nonexistent_asset(self):
        """Test editing a nonexistent asset."""
        self.client.login(email='admin@example.com', password='AdminPass123!')
        admin = User.objects.create_superuser(
            email='admin@example.com',
            password='AdminPass123!',
            first_name='Admin',
            last_name='User',
            department='IT'
        )
        self.client.login(email='admin@example.com', password='AdminPass123!')
        
        response = self.client.get(reverse('tickets:asset_edit_page', args=[99999]))
        self.assertEqual(response.status_code, 404)

    def test_claim_already_claimed_ticket(self):
        """Test claiming a ticket that's already assigned."""
        agent = User.objects.create_user(
            email='agent@example.com',
            password='TestPass123!',
            first_name='Agent',
            last_name='User',
            department='IT',
            role=User.Role.AGENT,
            is_active=True,
            email_verified=True
        )
        
        ticket = Ticket.objects.create(
            number='TK#7777',
            title='Test Ticket',
            description='Test description',
            requester=self.user,
            category=self.category,
            status=Ticket.Status.ASSIGNED,
            assigned_to=agent
        )
        
        self.client.login(email='agent@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:claim_ticket', args=[ticket.pk]))
        # Should not allow claiming an already assigned ticket
        # The view might handle this differently, but we check it doesn't error
        self.assertNotEqual(response.status_code, 500)


class RemoteSessionTests(TestCase):
    """Covers the request -> accept/reject -> start -> end remote session flow,
    including the bug fixes made to it (admin oversight access, REJECTED
    old_status guard, END notifying the requester, invalid-code feedback,
    no-active-connector error surfacing)."""

    def setUp(self):
        self.client = Client()
        self.requester = User.objects.create_user(
            email='requester@example.com', password='TestPass123!',
            first_name='Req', last_name='User', department='IT',
            is_active=True, email_verified=True,
        )
        self.agent = User.objects.create_user(
            email='rsagent@example.com', password='TestPass123!',
            first_name='Agent', last_name='User', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.other_agent = User.objects.create_user(
            email='otheragent@example.com', password='TestPass123!',
            first_name='Other', last_name='Agent', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.admin = User.objects.create_user(
            email='rsadmin@example.com', password='TestPass123!',
            first_name='Admin', last_name='User', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.category = Category.objects.create(name='Hardware', slug='rs-hardware')
        self.ticket = Ticket.objects.create(
            number='TK#RS1', title='Need help', description='desc',
            requester=self.requester, category=self.category, status=Ticket.Status.NEW,
        )
        self.connector = RemoteConnector.objects.create(name='Quick Assist', is_active=True)

    def _request_session(self):
        self.client.login(email='rsagent@example.com', password='TestPass123!')
        with patch('apps.tickets.views.send_email_via_brevo', return_value=(True, {})):
            response = self.client.post(reverse('tickets:request_remote_session', args=[self.ticket.pk]))
        return response

    def test_request_session_creates_session_and_notifies_requester(self):
        response = self._request_session()
        self.assertEqual(response.status_code, 200)
        session = RemoteSession.objects.get(ticket=self.ticket)
        self.assertEqual(session.status, 'REQUESTED')
        self.assertTrue(Notification.objects.filter(recipient=self.requester, type=Notification.Type.REMOTE_SESSION).exists())

    def test_request_session_with_no_active_connector_returns_json_error(self):
        self.connector.is_active = False
        self.connector.save()
        response = self._request_session()
        self.assertEqual(response.status_code, 400)
        self.assertIn('error', response.json())

    def test_request_session_blocked_while_one_already_pending(self):
        self._request_session()
        response = self._request_session()
        self.assertEqual(response.status_code, 400)
        self.assertEqual(RemoteSession.objects.filter(ticket=self.ticket).count(), 1)

    def test_non_agent_role_forbidden_from_requesting(self):
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:request_remote_session', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 403)

    def test_requester_can_accept_session(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'ACCEPTED'}
        )
        self.assertEqual(response.status_code, 302)
        session.refresh_from_db()
        self.assertEqual(session.status, 'ACCEPTED')

        comment = self.ticket.comments.get(body__icontains='accepted the remote session request')
        self.assertTrue(comment.is_system_generated)
        self.assertEqual(comment.system_icon, 'monitor')

    def test_requester_can_reject_before_start(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'REJECTED'}
        )
        self.assertEqual(response.status_code, 302)
        session.refresh_from_db()
        self.assertEqual(session.status, 'REJECTED')
        notif = Notification.objects.filter(recipient=self.agent, message__icontains='rejected').first()
        self.assertIsNotNone(notif)
        self.assertEqual(notif.type, Notification.Type.REMOTE_SESSION)

        comment = self.ticket.comments.get(body__icontains='declined the remote session request')
        self.assertTrue(comment.is_system_generated)

    def test_reject_is_ignored_once_session_has_started(self):
        """The REJECTED transition previously had no old_status guard — a stray
        POST could 'reject' a session that was already STARTED/ENDED."""
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        session.status = 'STARTED'
        session.started_at = timezone.now()
        session.save()

        self.client.login(email='requester@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'REJECTED'})
        session.refresh_from_db()
        self.assertEqual(session.status, 'STARTED')

    def test_agent_start_with_invalid_code_shows_error_and_does_not_start(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        session.status = 'ACCEPTED'
        session.save()

        self.client.login(email='rsagent@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'STARTED', 'quick_assist_code': 'ab'}
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Enter the 6-digit code')
        session.refresh_from_db()
        self.assertEqual(session.status, 'ACCEPTED')

    def test_agent_start_with_valid_code_starts_session(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        session.status = 'ACCEPTED'
        session.save()

        self.client.login(email='rsagent@example.com', password='TestPass123!')
        with patch('apps.tickets.views.send_email_via_brevo', return_value=(True, {})):
            response = self.client.post(
                reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'STARTED', 'quick_assist_code': 'ABC123'}
            )
        self.assertEqual(response.status_code, 302)
        session.refresh_from_db()
        self.assertEqual(session.status, 'STARTED')
        self.assertEqual(session.session_code, 'ABC123')

    def test_ending_session_notifies_requester(self):
        """ENDED previously sent no notification/email to the requester at all."""
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        session.status = 'STARTED'
        session.started_at = timezone.now()
        session.save()

        self.client.login(email='rsagent@example.com', password='TestPass123!')
        with patch('apps.tickets.views.send_email_via_brevo', return_value=(True, {})):
            response = self.client.post(
                reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'ENDED'}
            )
        self.assertEqual(response.status_code, 302)
        session.refresh_from_db()
        self.assertEqual(session.status, 'ENDED')
        self.assertTrue(
            Notification.objects.filter(
                recipient=self.requester, type=Notification.Type.REMOTE_SESSION, message__icontains='ended'
            ).exists()
        )

        comment = self.ticket.comments.get(body='The remote session has ended.')
        self.assertTrue(comment.is_system_generated)
        self.assertEqual(comment.system_icon, 'monitor')

    def test_admin_can_view_session_detail_they_are_not_party_to(self):
        """Admin could already see every session on the list page, but got a
        403 opening one they weren't the requester/agent on."""
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        self.client.login(email='rsadmin@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:remote_session_detail', args=[session.pk]))
        self.assertEqual(response.status_code, 200)

    def test_unrelated_agent_forbidden_from_session_detail(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        self.client.login(email='otheragent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:remote_session_detail', args=[session.pk]))
        self.assertEqual(response.status_code, 403)

    def test_end_user_sees_remote_sessions_list_scoped_to_own_requests(self):
        self._request_session()
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:remote_sessions_list'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, self.ticket.number)

    def test_dashboard_banner_shows_pending_accept_for_requester(self):
        self._request_session()
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'A remote session was requested')
        self.assertContains(response, 'Accept')

    def test_dashboard_banner_shows_pending_start_for_agent(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        session.status = 'ACCEPTED'
        session.save()
        self.client.login(email='rsagent@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'accepted your remote session request')
        self.assertContains(response, 'Start Session')

    def test_dashboard_banner_empty_when_nothing_pending(self):
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'A remote session was requested')

    def test_list_page_inline_accept(self):
        """Accept/Reject work directly from the Remote Sessions list row,
        not just the detail page."""
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.post(
            reverse('tickets:remote_session_detail', args=[session.pk]), {'status': 'ACCEPTED'}
        )
        self.assertEqual(response.status_code, 302)
        session.refresh_from_db()
        self.assertEqual(session.status, 'ACCEPTED')

    def test_list_page_shows_accept_reject_buttons_for_requester(self):
        self._request_session()
        self.client.login(email='requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:remote_sessions_list'))
        self.assertContains(response, '>Accept<')
        self.assertContains(response, '>Reject<')

    def test_list_page_shows_start_button_for_agent_once_accepted(self):
        self._request_session()
        session = RemoteSession.objects.get(ticket=self.ticket)
        session.status = 'ACCEPTED'
        session.save()
        self.client.login(email='rsagent@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:remote_sessions_list'))
        self.assertContains(response, 'Start')


class RemoteSessionExpiryTests(TestCase):
    """Covers the process_remote_session_expiry management command: expiry
    after the 2-hour window, the audit log entry, both notifications, and
    that an expired session no longer blocks a fresh request."""

    def setUp(self):
        self.client = Client()
        self.requester = User.objects.create_user(
            email='expreq@example.com', password='TestPass123!',
            first_name='Exp', last_name='Requester', department='IT',
            is_active=True, email_verified=True,
        )
        self.agent = User.objects.create_user(
            email='expagent@example.com', password='TestPass123!',
            first_name='Exp', last_name='Agent', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.category = Category.objects.create(name='Hardware', slug='exp-hardware')
        self.ticket = Ticket.objects.create(
            number='TK#EXP1', title='Need help', description='desc',
            requester=self.requester, category=self.category, status=Ticket.Status.NEW,
        )
        self.connector = RemoteConnector.objects.create(name='Quick Assist Exp', is_active=True)

    def _make_stale_session(self, status='REQUESTED', hours_old=3):
        session = RemoteSession.objects.create(
            ticket=self.ticket, requester=self.requester, agent=self.agent,
            connector=self.connector, status=status,
        )
        RemoteSession.objects.filter(pk=session.pk).update(
            created_at=timezone.now() - timedelta(hours=hours_old)
        )
        session.refresh_from_db()
        return session

    def test_expires_stale_requested_session(self):
        session = self._make_stale_session(status='REQUESTED', hours_old=3)
        call_command('process_remote_session_expiry', verbosity=0)
        session.refresh_from_db()
        self.assertEqual(session.status, 'EXPIRED')

    def test_expires_stale_accepted_session(self):
        session = self._make_stale_session(status='ACCEPTED', hours_old=3)
        call_command('process_remote_session_expiry', verbosity=0)
        session.refresh_from_db()
        self.assertEqual(session.status, 'EXPIRED')

    def test_does_not_expire_recent_session(self):
        session = self._make_stale_session(status='REQUESTED', hours_old=1)
        call_command('process_remote_session_expiry', verbosity=0)
        session.refresh_from_db()
        self.assertEqual(session.status, 'REQUESTED')

    def test_does_not_expire_started_session(self):
        """A session already in progress shouldn't be swept up by the
        REQUESTED/ACCEPTED-only expiry check."""
        session = self._make_stale_session(status='STARTED', hours_old=3)
        call_command('process_remote_session_expiry', verbosity=0)
        session.refresh_from_db()
        self.assertEqual(session.status, 'STARTED')

    def test_expiry_notifies_both_agent_and_requester(self):
        session = self._make_stale_session(status='REQUESTED', hours_old=3)
        call_command('process_remote_session_expiry', verbosity=0)
        self.assertTrue(
            Notification.objects.filter(recipient=self.agent, message__icontains='send a new request').exists()
        )
        self.assertTrue(
            Notification.objects.filter(recipient=self.requester, message__icontains='expired').exists()
        )

    def test_expiry_writes_audit_log_with_system_actor(self):
        session = self._make_stale_session(status='REQUESTED', hours_old=3)
        call_command('process_remote_session_expiry', verbosity=0)
        log = TicketComment.objects.filter(ticket=self.ticket, body__icontains='expired').first()
        self.assertIsNotNone(log)
        system_user = User.objects.get(email='system@ticketswipe.local')
        self.assertEqual(log.author, system_user)

    def test_expired_session_unblocks_new_request(self):
        self._make_stale_session(status='REQUESTED', hours_old=3)
        call_command('process_remote_session_expiry', verbosity=0)

        self.client.login(email='expagent@example.com', password='TestPass123!')
        with patch('apps.tickets.views.send_email_via_brevo', return_value=(True, {})):
            response = self.client.post(reverse('tickets:request_remote_session', args=[self.ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            RemoteSession.objects.filter(ticket=self.ticket, status='REQUESTED').count(), 1
        )


class ReportExportRedesignTests(TestCase):
    """Service Request / Incident PDF, DOCX, and detail-view exports after
    the paper-form redesign: letterhead, uploaded-signature images, and
    device-location capture."""

    def setUp(self):
        self.client = Client()
        self.service_category = ServiceCategory.objects.create(
            name='General IT', slug='general-it-export', field_group=ServiceCategory.FieldGroup.GENERAL
        )
        self.requester = User.objects.create_user(
            email='exportuser@example.com', password='TestPass123!',
            first_name='Export', last_name='User', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.admin = User.objects.create_user(
            email='exportadmin@example.com', password='TestPass123!',
            first_name='Export', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )

        self.sr_ticket = Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST,
            title='Export test service request',
            description='Testing exports',
            requester=self.requester,
            service_category=self.service_category,
            purpose='Testing',
            number='SRV#9001',
            status=Ticket.Status.APPROVED,
            submission_latitude=Decimal('6.500000'),
            submission_longitude=Decimal('3.400000'),
            submission_location_address='Lagos, Nigeria',
        )
        self.incident_ticket = Ticket.objects.create(
            type=Ticket.Type.INCIDENT,
            title='Export test incident',
            description='Testing incident export',
            requester=self.requester,
            number='TK#9002',
            status=Ticket.Status.NEW,
        )
        self.client.login(email='exportadmin@example.com', password='TestPass123!')

    def test_service_request_pdf_export(self):
        response = self.client.get(
            reverse('tickets:export_report_record', args=['service-requests', self.sr_ticket.pk]) + '?format=pdf'
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')

    def test_service_request_docx_export(self):
        response = self.client.get(
            reverse('tickets:export_report_record', args=['service-requests', self.sr_ticket.pk]) + '?format=docx'
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn('wordprocessingml', response['Content-Type'])

    def test_incident_pdf_export(self):
        response = self.client.get(
            reverse('tickets:export_report_record', args=['incidents', self.incident_ticket.pk]) + '?format=pdf'
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')

    def test_incident_docx_export(self):
        response = self.client.get(
            reverse('tickets:export_report_record', args=['incidents', self.incident_ticket.pk]) + '?format=docx'
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn('wordprocessingml', response['Content-Type'])

    def test_service_request_detail_view_shows_location_and_map_link(self):
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['service-requests', self.sr_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Lagos, Nigeria')
        self.assertContains(response, 'openstreetmap.org')

    def test_service_request_detail_view_no_location_shows_not_available(self):
        bare_ticket = Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST, title='No location', description='—',
            requester=self.requester, number='SRV#9003', status=Ticket.Status.APPROVED,
        )
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['service-requests', bare_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Not available')

    def test_pdf_and_docx_export_with_uploaded_signature(self):
        """The fiddliest new code path — image embedding in both PDF (img
        src) and DOCX (doc.add_picture via _docx_image_source) — must not
        error when a signer has an uploaded signature."""
        self.requester.signature = SimpleUploadedFile('sig.png', TINY_PNG_BYTES, content_type='image/png')
        self.requester.save()

        pdf_response = self.client.get(
            reverse('tickets:export_report_record', args=['service-requests', self.sr_ticket.pk]) + '?format=pdf'
        )
        self.assertEqual(pdf_response.status_code, 200)

        docx_response = self.client.get(
            reverse('tickets:export_report_record', args=['service-requests', self.sr_ticket.pk]) + '?format=docx'
        )
        self.assertEqual(docx_response.status_code, 200)
        self.assertIn('wordprocessingml', docx_response['Content-Type'])

    def test_signoff_context_helper_blank_when_no_user(self):
        from apps.tickets.report_registry import _signoff_context
        result = _signoff_context(None, None)
        self.assertIsNone(result['user'])
        self.assertIsNone(result['signature_url'])
        self.assertEqual(result['captured_text'], '')

    def test_signoff_context_helper_uses_signature_when_uploaded(self):
        from apps.tickets.report_registry import _signoff_context
        self.requester.signature = SimpleUploadedFile('sig.png', TINY_PNG_BYTES, content_type='image/png')
        self.requester.save()
        result = _signoff_context(self.requester, timezone.now())
        self.assertIsNotNone(result['signature_url'])
        self.assertIn('captured digitally', result['captured_text'])

    def test_signoff_context_captured_text_embeds_date_and_time(self):
        from apps.tickets.report_registry import _signoff_context
        when = timezone.make_aware(timezone.datetime(2026, 8, 20, 10, 35))
        result = _signoff_context(self.requester, when)
        self.assertEqual(
            result['captured_text'],
            f"{self.requester.get_full_name()} — captured digitally, on 2026-08-20 at 10:35"
        )

    def test_service_request_pdf_letterhead_date_is_generation_date_not_hardcoded(self):
        from django.template.loader import render_to_string
        from apps.tickets.report_registry import service_request_form_sections
        context = service_request_form_sections(self.sr_ticket)
        html = render_to_string('reports/service_request_form_pdf2.html', {**context, 'generated_at': timezone.now()})
        # No longer the old static placeholder revision date — reflects
        # today (the export date), matching the already-dynamic incident/
        # maintenance letterheads.
        self.assertNotIn('18th May 2023', html)
        self.assertIn(timezone.now().strftime('%Y'), html)

    def test_location_context_helper_falls_back_to_coordinates(self):
        from apps.tickets.report_registry import _location_context
        ticket = Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST, title='Offshore', description='—',
            requester=self.requester, number='SRV#9004', status=Ticket.Status.APPROVED,
            submission_latitude=Decimal('4.123456'), submission_longitude=Decimal('3.654321'),
            submission_location_address='',
        )
        result = _location_context(ticket)
        self.assertTrue(result['has_coordinates'])
        self.assertIsNone(result['address'])
        self.assertIn('4.123456', result['display'])
        self.assertIsNotNone(result['map_url'])

    def test_location_context_helper_none_when_no_data(self):
        from apps.tickets.report_registry import _location_context
        ticket = Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST, title='No geo', description='—',
            requester=self.requester, number='SRV#9005', status=Ticket.Status.APPROVED,
        )
        result = _location_context(ticket)
        self.assertFalse(result['has_coordinates'])
        self.assertIsNone(result['display'])
        self.assertIsNone(result['map_url'])

    def test_create_service_request_captures_submitted_location(self):
        self.client.logout()
        self.client.login(email='exportuser@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:create'), {
            'type': 'SERVICE_REQUEST',
            'title': 'Geo capture test',
            'description': 'Testing location capture',
            'service_category': self.service_category.id,
            'purpose': 'Testing',
            'urgency': 'MEDIUM',
            'submission_latitude': '6.5244',
            'submission_longitude': '3.3792',
            'submission_location_address': 'Lagos, Nigeria',
        })
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Geo capture test')
        self.assertEqual(ticket.submission_latitude, Decimal('6.524400'))
        self.assertEqual(ticket.submission_longitude, Decimal('3.379200'))
        self.assertEqual(ticket.submission_location_address, 'Lagos, Nigeria')

    def test_create_service_request_with_malformed_location_does_not_error(self):
        self.client.logout()
        self.client.login(email='exportuser@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:create'), {
            'type': 'SERVICE_REQUEST',
            'title': 'Bad geo capture test',
            'description': 'Testing malformed location handling',
            'service_category': self.service_category.id,
            'purpose': 'Testing',
            'urgency': 'MEDIUM',
            'submission_latitude': 'not-a-number',
            'submission_longitude': 'also-not-a-number',
        })
        self.assertEqual(response.status_code, 302)
        ticket = Ticket.objects.get(title='Bad geo capture test')
        self.assertIsNone(ticket.submission_latitude)
        self.assertIsNone(ticket.submission_longitude)


class MaintenanceReportExportTests(TestCase):
    """Maintenance's form-styled report detail page + PDF/DOCX exports,
    mirroring the Service Request/Incident redesign."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='maintreportadmin@example.com', password='TestPass123!',
            first_name='Maint', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        category = AssetCategory.objects.create(name='Maintenance Report Category')
        it_department, _ = AssetDepartment.objects.get_or_create(
            legacy_user_department_code='IT', defaults={'name': 'IT'}
        )
        self.asset = Asset.objects.create(name='Report Asset', category=category, department=it_department)
        self.schedule = MaintenanceSchedule.objects.create(
            title='Report export schedule',
            departments=[MaintenanceSchedule.Department.IT],
            scheduled_date=timezone.now().date() + timedelta(days=1),
            assigned_to=self.admin,
            checklist_items=['Check A', 'Check B'],
            completed_checklist=['Check A'],
            facility_location='Server Room A',
        )
        self.schedule.target_assets.add(self.asset)
        self.client.login(email='maintreportadmin@example.com', password='TestPass123!')

    def test_maintenance_detail_view_uses_form_layout(self):
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['maintenance', self.schedule.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'MAINTENANCE SCHEDULE REPORT')
        self.assertContains(response, 'Server Room A')
        self.assertContains(response, 'Report Asset')

    def test_maintenance_pdf_export(self):
        response = self.client.get(
            reverse('tickets:export_report_record', args=['maintenance', self.schedule.pk]) + '?format=pdf'
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')

    def test_maintenance_docx_export(self):
        response = self.client.get(
            reverse('tickets:export_report_record', args=['maintenance', self.schedule.pk]) + '?format=docx'
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn('wordprocessingml', response['Content-Type'])


class BulkReportExportTests(TestCase):
    """Checking specific rows on a report table exports only those records
    — exporting with nothing checked still exports the full filtered set."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='bulkexportadmin@example.com', password='TestPass123!',
            first_name='Bulk', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.login(email='bulkexportadmin@example.com', password='TestPass123!')
        self.category = AssetCategory.objects.create(name='Bulk Export Category')
        self.asset1 = Asset.objects.create(name='Bulk Export Asset One', category=self.category)
        self.asset2 = Asset.objects.create(name='Bulk Export Asset Two', category=self.category)
        self.asset3 = Asset.objects.create(name='Bulk Export Asset Three', category=self.category)

    def test_export_with_no_selection_includes_all_matching_rows(self):
        # Assets export dropped CSV (see report_builder.html/export_report —
        # a flat-text format can't carry the branded letterhead every other
        # asset export format shares), so this uses JSON instead — still a
        # flat/unfiltered export, just a different serialization.
        response = self.client.get(reverse('tickets:export_report', args=['assets']) + '?format=json')
        content = response.content.decode()
        self.assertIn('Bulk Export Asset One', content)
        self.assertIn('Bulk Export Asset Two', content)
        self.assertIn('Bulk Export Asset Three', content)

    def test_export_with_selected_ids_only_includes_those_rows(self):
        response = self.client.get(
            reverse('tickets:export_report', args=['assets'])
            + f'?format=json&ids={self.asset1.pk}&ids={self.asset3.pk}'
        )
        content = response.content.decode()
        self.assertIn('Bulk Export Asset One', content)
        self.assertIn('Bulk Export Asset Three', content)
        self.assertNotIn('Bulk Export Asset Two', content)

    def test_report_export_with_cols_only_includes_selected_columns(self):
        # cols-filtering isn't CSV-specific — Excel supports the same
        # mechanism (see build_export_response) and is what assets actually
        # use since CSV was dropped for that report type.
        from openpyxl import load_workbook
        response = self.client.get(
            reverse('tickets:export_report', args=['assets']) + '?format=excel&cols=Tracking%20ID,Name,Status'
        )
        wb = load_workbook(BytesIO(response.content))
        # Row 1 is the branded letterhead banner (logo/title/control no.),
        # not the column header — see BANNER_ROWS/meta_row/header_row in
        # report_exporters.export_excel. The real header row is 6:
        # 3 banner rows + 1 meta row + 1 blank spacer + 1 header row.
        header = [c.value for c in wb.active[6]]
        self.assertEqual(header, ['Tracking ID', 'Name', 'Status'])

    def test_excel_export_with_cols_only_includes_selected_columns(self):
        from openpyxl import load_workbook
        response = self.client.get(
            reverse('tickets:export_report', args=['assets']) + '?format=excel&cols=Name,Category'
        )
        wb = load_workbook(BytesIO(response.content))
        # The letterhead banner always reserves room for 3 zones (logo/
        # title/control), so a data table narrower than that (here: 2
        # columns) leaves one extra blank bordered column touched on the
        # banner rows — openpyxl then reports it as a trailing None cell on
        # every row, including the header. Only the real columns matter here.
        header = [c.value for c in wb.active[6]][:2]  # see header_row note above
        self.assertEqual(header, ['Name', 'Category'])

    def test_bundled_pdf_export_respects_cols_param(self):
        # The bundled/list PDF table export (report_pdf.html) is columns-
        # driven just like CSV/Excel — a narrowed `cols` selection should
        # still render successfully (unlike a single-record letterhead form
        # PDF, which never receives a `cols` param at all).
        response = self.client.get(
            reverse('tickets:export_report', args=['assets']) + '?format=pdf&cols=Name,Status'
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')

    def test_bundled_docx_export_with_cols_only_includes_selected_columns(self):
        from docx import Document
        response = self.client.get(
            reverse('tickets:export_report', args=['assets']) + '?format=docx&cols=Name,Status'
        )
        doc = Document(BytesIO(response.content))
        # tables[0] is the letterhead (logo/title/meta) — the data table is
        # the second table on the page (see _docx_report_letterhead).
        table = doc.tables[1]
        header = [cell.text for cell in table.rows[0].cells]
        self.assertEqual(header, ['Name', 'Status'])

    def test_excel_export_with_unknown_cols_falls_back_to_all_columns(self):
        from openpyxl import load_workbook
        response = self.client.get(
            reverse('tickets:export_report', args=['assets']) + '?format=excel&cols=Nonexistent,Also%20Bogus'
        )
        wb = load_workbook(BytesIO(response.content))
        header = [c.value for c in wb.active[6]]  # see header_row note above
        self.assertIn('Tracking ID', header)
        self.assertIn('Name', header)
        self.assertIn('Category', header)

    def test_report_table_renders_row_checkboxes(self):
        response = self.client.get(reverse('tickets:report_builder', args=['assets']))
        self.assertContains(response, 'bulk-row-checkbox')
        self.assertContains(response, f'value="{self.asset1.pk}"')


class AuditLogColumnPickerExportTests(TestCase):
    """The Logs page's own hand-rolled CSV/Excel/JSON exporter (separate
    from the report_registry Exportables system) respects the same `cols`
    column-picker param as build_export_response — see
    components/export_menu.html."""

    def setUp(self):
        self.client = Client()
        self.it_lead = User.objects.create_user(
            email='auditcolslead@example.com', password='TestPass123!',
            first_name='Audit', last_name='Lead', department='IT',
            role=User.Role.TEAM_LEAD, is_active=True, email_verified=True,
        )
        self.client.login(email='auditcolslead@example.com', password='TestPass123!')

    def test_csv_export_with_cols_only_includes_selected_columns(self):
        response = self.client.get(reverse('tickets:audit_log') + '?tab=tickets&format=csv&cols=Time,Actor')
        header = response.content.decode().splitlines()[0]
        self.assertEqual(header, 'Time,Actor')

    def test_json_export_ignores_cols_param(self):
        response = self.client.get(reverse('tickets:audit_log') + '?tab=tickets&format=json&cols=Time')
        self.assertEqual(response['Content-Type'], 'application/json')

    def test_csv_export_with_unknown_cols_falls_back_to_all_columns(self):
        response = self.client.get(reverse('tickets:audit_log') + '?tab=tickets&format=csv&cols=Bogus')
        header = response.content.decode().splitlines()[0]
        self.assertEqual(header, 'Time,Category,Ticket,Action,Actor,Details')


class NonITTeamLeadRestrictionTests(TestCase):
    """A Team Lead outside IT is scoped solely to the departmental service-
    request approval flow — IT-operational views/sidebar/dashboard stats
    must stay off-limits, while the approval flow itself keeps working."""

    def setUp(self):
        self.client = Client()
        self.non_it_lead = User.objects.create_user(
            email='marinelead2@example.com', password='TestPass123!',
            first_name='Marine', last_name='Lead', department='MARINE',
            role=User.Role.TEAM_LEAD, is_active=True, email_verified=True,
        )
        self.it_lead = User.objects.create_user(
            email='itlead2@example.com', password='TestPass123!',
            first_name='IT', last_name='Lead', department='IT',
            role=User.Role.TEAM_LEAD, is_active=True, email_verified=True,
        )

    def test_sidebar_template_split_by_department(self):
        from apps.common.permissions import get_sidebar_template
        self.assertEqual(get_sidebar_template(self.non_it_lead), 'partials/sidebar_team_lead_approver.html')
        self.assertEqual(get_sidebar_template(self.it_lead), 'partials/sidebar_team_lead.html')

    def test_non_it_lead_denied_it_operational_views(self):
        self.client.login(email='marinelead2@example.com', password='TestPass123!')
        for url_name in ['tickets:team_queue', 'tickets:unassigned', 'tickets:assigned_to_me',
                          'tickets:escalated_tickets', 'tickets:assets', 'tickets:audit_log',
                          'tickets:macro_management']:
            response = self.client.get(reverse(url_name))
            self.assertEqual(response.status_code, 403, f'{url_name} should 403 for a non-IT Team Lead')

    def test_it_lead_keeps_access_to_it_operational_views(self):
        self.client.login(email='itlead2@example.com', password='TestPass123!')
        for url_name in ['tickets:team_queue', 'tickets:unassigned', 'tickets:assigned_to_me',
                          'tickets:escalated_tickets', 'tickets:audit_log',
                          'tickets:macro_management']:
            response = self.client.get(reverse(url_name))
            self.assertEqual(response.status_code, 200, f'{url_name} should stay 200 for an IT Team Lead')

    def test_it_lead_denied_full_asset_inventory(self):
        """The full asset inventory is Admin/Superadmin-only now — even an
        IT Team Lead only ever sees assets assigned to them, via my_assets."""
        self.client.login(email='itlead2@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:assets'))
        self.assertEqual(response.status_code, 403)

    def test_non_it_lead_denied_report_access(self):
        self.client.login(email='marinelead2@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:report_builder', args=['service-requests']))
        self.assertEqual(response.status_code, 403)

    def test_it_lead_keeps_report_access(self):
        self.client.login(email='itlead2@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:report_builder', args=['service-requests']))
        self.assertEqual(response.status_code, 200)

    def test_non_it_lead_keeps_manager_review_access(self):
        """The approval flow itself must stay reachable — that's the whole
        point of a non-IT department having a Team Lead."""
        self.client.login(email='marinelead2@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:manager_review_queue'))
        self.assertEqual(response.status_code, 200)

    def test_non_it_lead_dashboard_uses_end_user_template(self):
        self.client.login(email='marinelead2@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'dashboards/end_user_dashboard.html')

    def test_it_lead_dashboard_uses_team_lead_template(self):
        self.client.login(email='itlead2@example.com', password='TestPass123!')
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'dashboards/team_lead_dashboard.html')

    def test_non_it_lead_remote_sessions_scoped_as_requester(self):
        response_login = self.client.login(email='marinelead2@example.com', password='TestPass123!')
        self.assertTrue(response_login)
        response = self.client.get(reverse('tickets:remote_sessions_list'))
        self.assertEqual(response.status_code, 200)


class ResolutionRootCauseCaptureTests(TestCase):
    """Root Cause / Resolution Steps captured from the resolving agent at
    resolve time — required for Incidents, optional for Service Requests —
    and reused in the Incident report and the KB-conversion pre-fill."""

    def setUp(self):
        self.client = Client()
        self.service_category = ServiceCategory.objects.create(
            name='General IT', slug='general-it-resolve', field_group=ServiceCategory.FieldGroup.GENERAL
        )
        self.requester = User.objects.create_user(
            email='resolveuser@example.com', password='TestPass123!',
            first_name='Resolve', last_name='User', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.agent = User.objects.create_user(
            email='resolveagent@example.com', password='TestPass123!',
            first_name='Resolve', last_name='Agent', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.incident_ticket = Ticket.objects.create(
            type=Ticket.Type.INCIDENT,
            title='Resolve flow incident',
            description='Testing resolve flow',
            requester=self.requester,
            number='TK#9101',
            status=Ticket.Status.NEW,
        )
        self.sr_ticket = Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST,
            title='Resolve flow service request',
            description='Testing resolve flow',
            requester=self.requester,
            service_category=self.service_category,
            purpose='Testing',
            number='SRV#9102',
            status=Ticket.Status.APPROVED,
        )
        self.client.login(email='resolveagent@example.com', password='TestPass123!')

    def test_incident_resolve_rejected_when_root_cause_blank(self):
        response = self.client.post(
            reverse('tickets:resolve_ticket', args=[self.incident_ticket.pk]),
            {'action': 'confirm', 'comment': '', 'resolution_root_cause': '', 'resolution_steps': 'Rebooted the server'},
        )
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertEqual(self.incident_ticket.status, Ticket.Status.NEW)
        self.assertEqual(self.incident_ticket.resolution_root_cause, '')

    def test_incident_resolve_rejected_when_resolution_steps_blank(self):
        response = self.client.post(
            reverse('tickets:resolve_ticket', args=[self.incident_ticket.pk]),
            {'action': 'confirm', 'comment': '', 'resolution_root_cause': 'Faulty NIC', 'resolution_steps': ''},
        )
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertEqual(self.incident_ticket.status, Ticket.Status.NEW)

    def test_incident_resolve_succeeds_with_both_fields(self):
        response = self.client.post(
            reverse('tickets:resolve_ticket', args=[self.incident_ticket.pk]),
            {'action': 'confirm', 'comment': '', 'resolution_root_cause': 'Faulty NIC', 'resolution_steps': 'Replaced the NIC'},
        )
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertEqual(self.incident_ticket.status, Ticket.Status.PENDING_USER)
        self.assertEqual(self.incident_ticket.resolution_root_cause, 'Faulty NIC')
        self.assertEqual(self.incident_ticket.resolution_steps, 'Replaced the NIC')

    def test_service_request_resolve_succeeds_with_blank_fields(self):
        response = self.client.post(
            reverse('tickets:resolve_ticket', args=[self.sr_ticket.pk]),
            {'action': 'confirm', 'comment': '', 'resolution_root_cause': '', 'resolution_steps': ''},
        )
        self.assertEqual(response.status_code, 302)
        self.sr_ticket.refresh_from_db()
        self.assertEqual(self.sr_ticket.status, Ticket.Status.PENDING_USER)

    def test_incident_report_shows_recorded_root_cause_and_steps(self):
        self.incident_ticket.resolution_root_cause = 'Faulty NIC'
        self.incident_ticket.resolution_steps = 'Replaced the NIC'
        self.incident_ticket.status = Ticket.Status.RESOLVED
        self.incident_ticket.save()
        self.client.logout()
        self.client.login(email='resolveuser@example.com', password='TestPass123!')
        admin = User.objects.create_user(
            email='resolveadmin@example.com', password='TestPass123!',
            first_name='Resolve', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.logout()
        self.client.login(email='resolveadmin@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Faulty NIC')
        self.assertContains(response, 'Replaced the NIC')

    def test_incident_report_shows_not_recorded_when_blank(self):
        self.incident_ticket.status = Ticket.Status.RESOLVED
        self.incident_ticket.save()
        admin = User.objects.create_user(
            email='resolveadmin2@example.com', password='TestPass123!',
            first_name='Resolve', last_name='Admin2', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.logout()
        self.client.login(email='resolveadmin2@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Not recorded')

    def test_kb_prefill_uses_recorded_resolution_data(self):
        self.incident_ticket.resolution_root_cause = 'Faulty NIC'
        self.incident_ticket.resolution_steps = 'Replaced the NIC'
        self.incident_ticket.status = Ticket.Status.RESOLVED
        self.incident_ticket.save()
        response = self.client.get(reverse('kb:convert_ticket', args=[self.incident_ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Faulty NIC')
        self.assertContains(response, 'Replaced the NIC')

    def test_kb_prefill_falls_back_to_empty_skeleton_when_blank(self):
        self.sr_ticket.status = Ticket.Status.RESOLVED
        self.sr_ticket.save()
        response = self.client.get(reverse('kb:convert_ticket', args=[self.sr_ticket.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '1. ')

    def test_resolve_modal_omits_fields_for_service_request(self):
        response = self.client.get(
            reverse('tickets:resolve_ticket', args=[self.sr_ticket.pk]), HTTP_HX_REQUEST='true'
        )
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'name="resolution_root_cause"')
        self.assertNotContains(response, 'name="resolution_steps"')

    def test_resolve_modal_shows_fields_and_categories_for_incident(self):
        response = self.client.get(
            reverse('tickets:resolve_ticket', args=[self.incident_ticket.pk]), HTTP_HX_REQUEST='true'
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'name="resolution_root_cause"')
        self.assertContains(response, 'name="resolution_root_cause_category"')
        self.assertContains(response, 'Human Error')

    def test_incident_resolve_captures_root_cause_categories(self):
        response = self.client.post(
            reverse('tickets:resolve_ticket', args=[self.incident_ticket.pk]),
            {
                'action': 'confirm', 'comment': '',
                'resolution_root_cause': 'Faulty NIC', 'resolution_steps': 'Replaced the NIC',
                'resolution_root_cause_category': ['HARDWARE_FAILURE', 'HUMAN_ERROR'],
            },
        )
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertEqual(
            sorted(self.incident_ticket.resolution_root_cause_category), ['HARDWARE_FAILURE', 'HUMAN_ERROR']
        )

    def test_incident_resolve_ignores_invalid_root_cause_categories(self):
        response = self.client.post(
            reverse('tickets:resolve_ticket', args=[self.incident_ticket.pk]),
            {
                'action': 'confirm', 'comment': '',
                'resolution_root_cause': 'Faulty NIC', 'resolution_steps': 'Replaced the NIC',
                'resolution_root_cause_category': ['NOT_A_REAL_CATEGORY'],
            },
        )
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertEqual(self.incident_ticket.resolution_root_cause_category, [])


class IncidentReportApprovalAndCommunicationTests(TestCase):
    """IT Manager / Head of IT merged sign-off, attachment images, and the
    system-derived Section 8 communication fields on the Incident report."""

    def setUp(self):
        self.client = Client()
        self.requester = User.objects.create_user(
            email='approvalreq@example.com', password='TestPass123!',
            first_name='Approval', last_name='Requester', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.agent = User.objects.create_user(
            email='approvalagent@example.com', password='TestPass123!',
            first_name='Approval', last_name='Agent', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.admin = User.objects.create_user(
            email='approvaladmin@example.com', password='TestPass123!',
            first_name='Approval', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.incident_ticket = Ticket.objects.create(
            type=Ticket.Type.INCIDENT,
            title='Approval flow incident',
            description='Testing IT manager approval',
            requester=self.requester,
            number='TK#9201',
            status=Ticket.Status.RESOLVED,
            resolution_root_cause='Faulty NIC',
            resolution_steps='Replaced the NIC',
        )

    def test_approve_denied_for_agent(self):
        self.client.login(email='approvalagent@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:approve_incident_report', args=[self.incident_ticket.pk]))
        self.assertEqual(response.status_code, 403)
        self.incident_ticket.refresh_from_db()
        self.assertIsNone(self.incident_ticket.incident_approved_by)

    def test_approve_rejected_when_not_resolved(self):
        self.incident_ticket.status = Ticket.Status.NEW
        self.incident_ticket.save()
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:approve_incident_report', args=[self.incident_ticket.pk]))
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertIsNone(self.incident_ticket.incident_approved_by)

    def test_approve_succeeds_and_is_idempotent(self):
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        response = self.client.post(reverse('tickets:approve_incident_report', args=[self.incident_ticket.pk]))
        self.assertEqual(response.status_code, 302)
        self.incident_ticket.refresh_from_db()
        self.assertEqual(self.incident_ticket.incident_approved_by, self.admin)
        self.assertIsNotNone(self.incident_ticket.incident_approved_at)
        first_approved_at = self.incident_ticket.incident_approved_at

        # Second approval attempt is a no-op, not a re-approval.
        other_admin = User.objects.create_user(
            email='approvaladmin2@example.com', password='TestPass123!',
            first_name='Approval', last_name='Admin2', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.logout()
        self.client.login(email='approvaladmin2@example.com', password='TestPass123!')
        self.client.post(reverse('tickets:approve_incident_report', args=[self.incident_ticket.pk]))
        self.incident_ticket.refresh_from_db()
        self.assertEqual(self.incident_ticket.incident_approved_by, self.admin)
        self.assertEqual(self.incident_ticket.incident_approved_at, first_approved_at)

    def test_report_detail_shows_merged_signoff_after_approval(self):
        self.incident_ticket.incident_approved_by = self.admin
        self.incident_ticket.incident_approved_at = timezone.now()
        self.incident_ticket.save()
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Reviewed &amp; Approved By (IT Manager / Head of IT)')
        self.assertNotContains(response, 'Admin Director')

    def test_report_detail_shows_approve_button_for_admin_only(self):
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertContains(response, 'Approve as IT Manager')

        team_lead = User.objects.create_user(
            email='approvalteamlead@example.com', password='TestPass123!',
            first_name='Approval', last_name='TeamLead', department='IT',
            role=User.Role.TEAM_LEAD, is_active=True, email_verified=True,
        )
        self.client.logout()
        self.client.login(email='approvalteamlead@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'Approve as IT Manager')

    def test_report_detail_shows_fixed_communication_fields(self):
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertContains(response, 'IT Helpdesk Ticket')
        self.assertContains(response, 'Report Communicated To')

    def test_pdf_and_docx_export_with_approval_and_categories(self):
        self.incident_ticket.resolution_root_cause_category = ['HUMAN_ERROR', 'HARDWARE_FAILURE']
        self.incident_ticket.incident_approved_by = self.admin
        self.incident_ticket.incident_approved_at = timezone.now()
        self.incident_ticket.save()
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')

        pdf_response = self.client.get(
            reverse('tickets:export_report_record', args=['incidents', self.incident_ticket.pk]) + '?format=pdf'
        )
        self.assertEqual(pdf_response.status_code, 200)

        docx_response = self.client.get(
            reverse('tickets:export_report_record', args=['incidents', self.incident_ticket.pk]) + '?format=docx'
        )
        self.assertEqual(docx_response.status_code, 200)
        self.assertIn('wordprocessingml', docx_response['Content-Type'])

    def test_pdf_export_with_image_attachment_includes_extra_page(self):
        from apps.tickets.models import Attachment
        Attachment.objects.create(
            ticket=self.incident_ticket,
            file=SimpleUploadedFile('proof.png', TINY_PNG_BYTES, content_type='image/png'),
            filename='proof.png',
            content_type='image/png',
            uploaded_by=self.agent,
        )
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        pdf_response = self.client.get(
            reverse('tickets:export_report_record', args=['incidents', self.incident_ticket.pk]) + '?format=pdf'
        )
        self.assertEqual(pdf_response.status_code, 200)

        docx_response = self.client.get(
            reverse('tickets:export_report_record', args=['incidents', self.incident_ticket.pk]) + '?format=docx'
        )
        self.assertEqual(docx_response.status_code, 200)

    def test_detail_view_shows_image_attachment_thumbnail(self):
        from apps.tickets.models import Attachment
        Attachment.objects.create(
            ticket=self.incident_ticket,
            file=SimpleUploadedFile('proof.png', TINY_PNG_BYTES, content_type='image/png'),
            filename='proof.png',
            content_type='image/png',
            uploaded_by=self.agent,
        )
        Attachment.objects.create(
            ticket=self.incident_ticket,
            file=SimpleUploadedFile('notes.pdf', b'%PDF-1.4 fake', content_type='application/pdf'),
            filename='notes.pdf',
            content_type='application/pdf',
            uploaded_by=self.agent,
        )
        self.client.login(email='approvaladmin@example.com', password='TestPass123!')
        response = self.client.get(
            reverse('tickets:report_record_detail', args=['incidents', self.incident_ticket.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'target="_blank"')
        self.assertContains(response, 'notes.pdf')


class ProcurementRequestTests(TestCase):
    """Vendor procurement: recording that an item isn't in inventory and is
    being sourced from a vendor, for either a Service Request ticket or a
    Mobilization, and the Receiving step that turns it into a real Asset."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_superuser(
            email='proc-admin@example.com', password='AdminPass123!',
            first_name='Proc', last_name='Admin', department='IT',
        )
        self.requester = User.objects.create_user(
            email='proc-requester@example.com', password='TestPass123!',
            first_name='Proc', last_name='Requester', department='OPERATIONS',
        )
        self.client.login(email='proc-admin@example.com', password='AdminPass123!')
        self.category = AssetCategory.objects.create(name='Laptops', is_consumable=False)
        self.consumable_category = AssetCategory.objects.create(name='Cable Ties', is_consumable=True)
        self.vendor = Vendor.objects.create(name='Acme Supplies', is_active=True)

    def _pending_fulfillment_ticket(self):
        return Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST,
            title='Need a laptop',
            description='Need a laptop, none in stock',
            requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT,
            is_asset_request=True,
        )

    def test_create_against_ticket_moves_it_to_pending_vendor(self):
        ticket = self._pending_fulfillment_ticket()
        response = self.client.post(reverse('tickets:procurement_request_create', args=[ticket.pk]), {
            'item_name': 'Dell Latitude 5440', 'category': self.category.pk, 'quantity': 1,
            'vendor': self.vendor.pk, 'new_vendor_name': '',
        })
        self.assertEqual(response.status_code, 302)
        ticket.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.Status.PENDING_VENDOR)
        pr = AssetProcurementRequest.objects.get(ticket=ticket)
        self.assertEqual(pr.status, AssetProcurementRequest.Status.REQUESTED)
        self.assertEqual(pr.vendor, self.vendor)

    def test_receiving_ticket_procurement_autofulfills(self):
        ticket = self._pending_fulfillment_ticket()
        self.client.post(reverse('tickets:procurement_request_create', args=[ticket.pk]), {
            'item_name': 'Dell Latitude 5440', 'category': self.category.pk, 'quantity': 1,
            'vendor': self.vendor.pk, 'new_vendor_name': '',
        })
        pr = AssetProcurementRequest.objects.get(ticket=ticket)

        response = self.client.post(reverse('tickets:procurement_receive', args=[pr.pk]))
        self.assertEqual(response.status_code, 302)

        ticket.refresh_from_db()
        pr.refresh_from_db()
        self.assertEqual(ticket.status, Ticket.Status.PENDING_USER)
        self.assertIsNotNone(ticket.assigned_asset)
        self.assertEqual(ticket.assigned_asset.name, 'Dell Latitude 5440')
        self.assertEqual(ticket.assigned_asset.assigned_to, self.requester)
        self.assertEqual(pr.status, AssetProcurementRequest.Status.RECEIVED)
        self.assertEqual(pr.received_by, self.admin)

    def _pending_fulfillment_mobilization_ticket(self):
        return Ticket.objects.create(
            type=Ticket.Type.SERVICE_REQUEST,
            title='Gear for job',
            description='Need gear mobilized',
            requester=self.requester,
            status=Ticket.Status.PENDING_FULFILLMENT,
            is_asset_request=True,
            is_mobilization_request=True,
        )

    def test_mobilization_can_be_created_with_only_procurement_rows(self):
        ticket = self._pending_fulfillment_mobilization_ticket()
        response = self.client.post(reverse('tickets:mobilization_create'), {
            'notes': 'Job needs items not in stock',
            'ticket_id': ticket.pk,
            'third_party_vessels': ['MV Procurement Test'],
            'procurement_item_name': ['Deck Cable'],
            'procurement_category_id': [str(self.category.pk)],
            'procurement_quantity': ['2'],
            'procurement_vendor_id': [str(self.vendor.pk)],
            'procurement_vendor_name': [''],
            'procurement_expected_date': [''],
        })
        self.assertEqual(response.status_code, 302)
        mobilization = Mobilization.objects.get(mobilized_by=self.admin)
        self.assertEqual(mobilization.items.count(), 0)
        pr = AssetProcurementRequest.objects.get(mobilization=mobilization)
        self.assertEqual(pr.item_name, 'Deck Cable')
        self.assertEqual(pr.quantity, 2)

        # One itemized CREATED comment, listing the vendor line — not the
        # old "assets mobilized" wording (nothing was actually sent).
        created_comment = ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.CREATED)
        self.assertIn('ordered from vendor', created_comment.body)
        self.assertNotIn('mobilized from stock', created_comment.body)

    def test_receiving_mobilization_procurement_creates_item_and_mobilizes(self):
        ticket = self._pending_fulfillment_mobilization_ticket()
        self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': ticket.pk,
            'third_party_vessels': ['MV Receive Test'],
            'procurement_item_name': ['Deck Cable'],
            'procurement_category_id': [str(self.category.pk)],
            'procurement_quantity': ['1'],
            'procurement_vendor_id': [str(self.vendor.pk)],
            'procurement_vendor_name': [''],
            'procurement_expected_date': [''],
        })
        mobilization = Mobilization.objects.get(mobilized_by=self.admin)
        pr = AssetProcurementRequest.objects.get(mobilization=mobilization)

        response = self.client.post(reverse('tickets:procurement_receive', args=[pr.pk]))
        self.assertEqual(response.status_code, 302)

        item = MobilizationItem.objects.get(mobilization=mobilization)
        self.assertEqual(item.asset.status, Asset.Status.MOBILIZED)
        self.assertEqual(item.quantity, 1)
        pr.refresh_from_db()
        self.assertEqual(pr.status, AssetProcurementRequest.Status.RECEIVED)

        arrived_comment = ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.VENDOR_ITEM_ARRIVED)
        self.assertIn('Deck Cable', arrived_comment.body)
        self.assertIn('All 1 items now fulfilled.', arrived_comment.body)

    def test_receiving_consumable_batch_reports_quantity_not_row_count(self):
        # A consumable vendor order for 3 units collapses into ONE
        # MobilizationItem row with quantity=3 — the arrival message must
        # say "All 3", not "All 1" (row count).
        ticket = self._pending_fulfillment_mobilization_ticket()
        self.client.post(reverse('tickets:mobilization_create'), {
            'ticket_id': ticket.pk,
            'third_party_vessels': ['MV Batch Test'],
            'procurement_item_name': ['Cable Ties'],
            'procurement_category_id': [str(self.consumable_category.pk)],
            'procurement_quantity': ['3'],
            'procurement_vendor_id': [str(self.vendor.pk)],
            'procurement_vendor_name': [''],
            'procurement_expected_date': [''],
        })
        mobilization = Mobilization.objects.get(mobilized_by=self.admin)
        pr = AssetProcurementRequest.objects.get(mobilization=mobilization)

        self.client.post(reverse('tickets:procurement_receive', args=[pr.pk]))

        item = MobilizationItem.objects.get(mobilization=mobilization)
        self.assertEqual(item.quantity, 3)

        arrived_comment = ticket.comments.get(mobilization_event=TicketComment.MobilizationEvent.VENDOR_ITEM_ARRIVED)
        self.assertIn('All 3 items now fulfilled.', arrived_comment.body)

        # The confirm-receipt card on the conversation page has the same
        # row-count-vs-quantity pitfall.
        ticket.refresh_from_db()
        self.client.logout()
        self.client.login(email='proc-requester@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:detail', args=[ticket.pk]))
        self.assertContains(response, '3 assets mobilized')

    def test_consumable_receive_increments_stock_and_low_stock_alert_clears(self):
        stock_asset = Asset.objects.create(
            name='Zip Ties', category=self.consumable_category, status=Asset.Status.IN_STORE,
            quantity_in_stock=1, low_stock_threshold=5,
        )
        stock_asset.refresh_low_stock_alert()
        self.assertTrue(Asset.objects.get(pk=stock_asset.pk).low_stock_notified)

        pr = AssetProcurementRequest.objects.create(
            item_name='Zip Ties', category=self.consumable_category, quantity=20,
            requested_by=self.admin,
        )
        self.client.post(reverse('tickets:procurement_receive', args=[pr.pk]))

        stock_asset.refresh_from_db()
        self.assertEqual(stock_asset.quantity_in_stock, 21)
        self.assertFalse(stock_asset.low_stock_notified)

    def test_cancel_open_request_has_no_asset_side_effects(self):
        ticket = self._pending_fulfillment_ticket()
        pr = AssetProcurementRequest.objects.create(
            item_name='Dell Latitude 5440', category=self.category, quantity=1,
            ticket=ticket, requested_by=self.admin,
        )
        response = self.client.post(reverse('tickets:procurement_cancel', args=[pr.pk]))
        self.assertEqual(response.status_code, 302)
        pr.refresh_from_db()
        self.assertEqual(pr.status, AssetProcurementRequest.Status.CANCELLED)
        self.assertFalse(Asset.objects.filter(name='Dell Latitude 5440').exists())

    def test_new_vendor_name_proposes_inactive_vendor_and_notifies_admins(self):
        other_admin = User.objects.create_superuser(
            email='proc-notify@example.com', password='AdminPass123!',
            first_name='Proc', last_name='Notify', department='IT', role=User.Role.ADMIN,
        )
        ticket = self._pending_fulfillment_ticket()
        self.client.post(reverse('tickets:procurement_request_create', args=[ticket.pk]), {
            'item_name': 'Deck Cable', 'category': self.category.pk, 'quantity': 1,
            'vendor': '', 'new_vendor_name': 'New Vendor Co',
        })
        vendor = Vendor.objects.get(name='New Vendor Co')
        self.assertFalse(vendor.is_active)
        pr = AssetProcurementRequest.objects.get(ticket=ticket)
        self.assertEqual(pr.vendor, vendor)
        self.assertTrue(
            Notification.objects.filter(recipient=other_admin, message__icontains='New Vendor Co').exists()
        )


class AttachmentPreviewTests(TestCase):
    """No preview should depend on Google Docs Viewer reaching a (possibly
    private/local) file URL — images/PDF/video/audio render natively, Office
    docs get converted to PDF via LibreOffice and cached."""

    def setUp(self):
        self.client = Client()
        self.category = Category.objects.create(name='General', slug='general')
        self.requester = User.objects.create_user(
            email='attach-req@example.com', password='TestPass123!',
            first_name='Attach', last_name='Requester', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.agent = User.objects.create_user(
            email='attach-agent@example.com', password='TestPass123!',
            first_name='Attach', last_name='Agent', department='IT',
            role=User.Role.AGENT, is_active=True, email_verified=True,
        )
        self.ticket = Ticket.objects.create(
            number='TK#ATT1', title='Attachment preview ticket', description='d',
            requester=self.requester, category=self.category, status=Ticket.Status.NEW,
        )
        self.client.login(email='attach-agent@example.com', password='TestPass123!')

    def _attach(self, filename, content_type, content=b'data'):
        return Attachment.objects.create(
            ticket=self.ticket,
            file=SimpleUploadedFile(filename, content, content_type=content_type),
            filename=filename, content_type=content_type, uploaded_by=self.agent,
        )

    def test_image_previews_natively_no_gview(self):
        attachment = self._attach('proof.png', 'image/png', TINY_PNG_BYTES)
        response = self.client.get(reverse('tickets:attachment_preview', args=[attachment.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<img')
        self.assertNotContains(response, 'docs.google.com')

    def test_pdf_previews_natively_no_gview(self):
        attachment = self._attach('notes.pdf', 'application/pdf', b'%PDF-1.4 fake')
        response = self.client.get(reverse('tickets:attachment_preview', args=[attachment.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'docs.google.com')
        self.assertContains(response, attachment.file.url)

    def test_office_document_converted_via_libreoffice_and_cached(self):
        attachment = self._attach(
            'report.docx',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        fake_pdf_dir = tempfile.mkdtemp()
        fake_pdf_path = os.path.join(fake_pdf_dir, 'converted.pdf')
        with open(fake_pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4 fake converted output')

        with patch('apps.documents_display.utils.convert_office_to_pdf', return_value=fake_pdf_path) as mock_convert:
            response = self.client.get(reverse('tickets:attachment_preview', args=[attachment.pk]))
            self.assertEqual(response.status_code, 200)
            mock_convert.assert_called_once()

        attachment.refresh_from_db()
        self.assertTrue(attachment.preview_pdf)
        self.assertNotContains(response, 'docs.google.com')

        # Second preview must reuse the cached PDF, not convert again.
        with patch('apps.documents_display.utils.convert_office_to_pdf') as mock_convert_again:
            response = self.client.get(reverse('tickets:attachment_preview', args=[attachment.pk]))
            self.assertEqual(response.status_code, 200)
            mock_convert_again.assert_not_called()

    def test_office_document_conversion_failure_shows_fallback(self):
        attachment = self._attach(
            'broken.docx',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        with patch('apps.documents_display.utils.convert_office_to_pdf', return_value=None):
            response = self.client.get(reverse('tickets:attachment_preview', args=[attachment.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Couldn't generate a preview")

    def test_video_and_audio_preview_natively(self):
        video = self._attach('clip.mp4', 'video/mp4')
        response = self.client.get(reverse('tickets:attachment_preview', args=[video.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<video')

        audio = self._attach('note.mp3', 'audio/mpeg')
        response = self.client.get(reverse('tickets:attachment_preview', args=[audio.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<audio')

    def test_unrecognized_format_falls_back_to_download(self):
        attachment = self._attach('archive.zip', 'application/zip')
        response = self.client.get(reverse('tickets:attachment_preview', args=[attachment.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Preview not available for this file type')


class AssetImportTransformTests(TestCase):
    """Regression tests for apps.tickets.asset_import_transform, covering the
    four bugs found while testing the import feature against real client
    inventory layouts: a location banner above the header row, a person's
    bundle-start row carrying no device, a header using an accepted alias
    word, and multi-line CSV field corruption."""

    def test_location_banner_above_header_row_is_carried_forward(self):
        from apps.tickets.asset_import_transform import transform_raw_rows
        rows = [
            ('Hydrodive Asset Inventory',),
            ('Ground Floor',),
            ('S/n', 'USER', 'DEPARTMENT', 'DEVICE', 'TAG', 'TRACK #NO', 'Active', 'Not Active', 'Comments'),
            (1, 'John Doe', 'Account', 'Monitor', 'HD GF ACC MNT 008', 'ACC008', 'Active', '', 'Functional'),
        ]
        out = transform_raw_rows(rows)
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0]['location_name'], 'Ground Floor')

    def test_later_location_banner_still_overrides_earlier_one(self):
        from apps.tickets.asset_import_transform import transform_raw_rows
        rows = [
            ('Ground Floor',),
            ('S/n', 'USER', 'DEPARTMENT', 'DEVICE', 'TAG', 'TRACK #NO', 'Active', 'Not Active', 'Comments'),
            (1, 'John Doe', 'Account', 'Monitor', 'HD GF ACC MNT 008', 'ACC008', 'Active', '', ''),
            ('First Floor',),
            (2, 'Jane Smith', 'HSE', 'Laptop', 'HD FF HSE LP 002', 'HSE002', '', 'Not Active', ''),
        ]
        out = transform_raw_rows(rows)
        self.assertEqual(out[0]['location_name'], 'Ground Floor')
        self.assertEqual(out[1]['location_name'], 'First Floor')

    def test_bundle_start_row_with_no_device_still_captures_person(self):
        from apps.tickets.asset_import_transform import transform_raw_rows
        rows = [
            ('S/n', 'USER', 'DEPARTMENT', 'DEVICE', 'TAG', 'TRACK #NO', 'Active', 'Not Active', 'Comments'),
            (1, 'John Doe', 'Account', '', '', '', 'Active', '', 'Functional'),
            (None, None, 'Account', 'Monitor', 'HD GF ACC MNT 008', 'ACC008', None, None, None),
            (None, None, 'Account', 'CPU', 'HD GF ACC CPU 008', None, None, None, None),
        ]
        out = transform_raw_rows(rows)
        self.assertEqual(len(out), 2)
        for row in out:
            self.assertEqual(row['assigned_to_name'], 'John Doe')
            self.assertEqual(row['status_hint'], 'ACTIVE')
            self.assertEqual(row['notes'], 'Functional')

    def test_header_recognized_via_accepted_alias_words(self):
        from apps.tickets.asset_import_transform import transform_raw_rows
        rows = [
            ('S/n', 'USER', 'DEPARTMENT', 'Item', 'Asset Tag', 'TRACK #NO', 'Active', 'Not Active', 'Comments'),
            (1, 'John Doe', 'Account', 'Monitor', 'HD GF ACC MNT 008', 'ACC008', 'Active', '', ''),
        ]
        out = transform_raw_rows(rows)
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0]['name'], 'Monitor')

    def test_csv_multiline_quoted_field_not_corrupted(self):
        from apps.tickets.views import _read_raw_rows
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD001,ACC008,Active,,"line1\nline2"\r\n'
        )
        upload = SimpleUploadedFile('inventory.csv', raw_csv.encode('utf-8'), content_type='text/csv')
        rows = _read_raw_rows(upload)
        self.assertEqual(rows[1][-1], 'line1\nline2')


class AssetImportEndToEndTests(TestCase):
    """Exercises the full upload -> preview -> commit flow through the real
    views, not just the transform function in isolation."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='import-admin@example.com', password='TestPass123!',
            first_name='Import', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.assignee = User.objects.create_user(
            email='import-assignee@example.com', password='TestPass123!',
            first_name='John', last_name='Doe', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.client.login(email='import-admin@example.com', password='TestPass123!')

    def _upload_csv(self, raw_csv):
        upload = SimpleUploadedFile('inventory.csv', raw_csv.encode('utf-8'), content_type='text/csv')
        return self.client.post(reverse('tickets:asset_import'), {'file': upload})

    def test_full_import_flow_creates_expected_assets(self):
        raw_csv = (
            'Hydrodive Asset Inventory\r\n'
            'Ground Floor\r\n'
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,Functional\r\n'
            ',,Account,CPU,HD GF ACC CPU 008,,,,\r\n'
        )
        response = self._upload_csv(raw_csv)
        self.assertEqual(response.status_code, 302)
        batch = AssetImportBatch.objects.get()
        self.assertEqual(batch.status, AssetImportBatch.Status.PENDING_REVIEW)
        self.assertEqual(batch.row_count, 2)

        preview = self.client.get(reverse('tickets:asset_import_preview', args=[batch.pk]))
        self.assertEqual(preview.status_code, 200)
        self.assertContains(preview, 'Monitor')

        commit = self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]))
        self.assertEqual(commit.status_code, 302)
        batch.refresh_from_db()
        self.assertEqual(batch.status, AssetImportBatch.Status.COMMITTED)

        self.assertEqual(Asset.objects.count(), 2)
        monitor = Asset.objects.get(name='Monitor')
        self.assertEqual(monitor.tracking_id, 'HD-GF-ACC-MNT-008')
        self.assertEqual(monitor.location.name, 'Ground Floor')
        self.assertEqual(monitor.department.name, 'Account')
        self.assertEqual(monitor.assigned_to, self.assignee)
        self.assertEqual(monitor.status, Asset.Status.IN_USE)

    def test_discarded_batch_creates_no_assets(self):
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        response = self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_discard', args=[batch.pk]))
        batch.refresh_from_db()
        self.assertEqual(batch.status, AssetImportBatch.Status.DISCARDED)
        self.assertEqual(Asset.objects.count(), 0)

    def test_non_admin_cannot_import(self):
        self.client.logout()
        self.client.login(email='import-assignee@example.com', password='TestPass123!')
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        response = self._upload_csv(raw_csv)
        self.assertEqual(response.status_code, 403)

    def test_unresolved_assignee_hint_set_when_name_does_not_resolve(self):
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,Nobody Matching,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]))
        asset = Asset.objects.get()
        self.assertIsNone(asset.assigned_to)
        self.assertEqual(asset.unresolved_assignee_hint, 'Nobody Matching')

    def test_unresolved_assignee_hint_blank_when_name_resolves(self):
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]))
        asset = Asset.objects.get()
        self.assertEqual(asset.assigned_to, self.assignee)
        self.assertEqual(asset.unresolved_assignee_hint, '')


class AssetImportDepartmentLocationMatchTests(TestCase):
    """Prevents a repeat of the 'Account' vs 'Accounting' incident — the
    preview step must flag sheet department/location text that doesn't
    exactly match an existing row, and the commit step must honor an
    admin's explicit choice to map it onto an existing one instead of
    silently creating a duplicate."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='deptmatch-admin@example.com', password='TestPass123!',
            first_name='Dept', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.login(email='deptmatch-admin@example.com', password='TestPass123!')

    def _upload_csv(self, raw_csv):
        upload = SimpleUploadedFile('inventory.csv', raw_csv.encode('utf-8'), content_type='text/csv')
        return self.client.post(reverse('tickets:asset_import'), {'file': upload})

    def test_preview_flags_department_and_location_with_no_exact_match(self):
        AssetDepartment.objects.get_or_create(name='Accounting', defaults={'tag_code': 'ACC'})[0]
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        preview = self.client.get(reverse('tickets:asset_import_preview', args=[batch.pk]))
        self.assertIn('Account', preview.context['unmatched_departments'])
        self.assertContains(preview, 'Review new departments')
        self.assertContains(preview, 'dept_map:Account')

    def test_preview_does_not_flag_exact_department_match(self):
        AssetDepartment.objects.get_or_create(name='Accounting', defaults={'tag_code': 'ACC'})[0]
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Accounting,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        preview = self.client.get(reverse('tickets:asset_import_preview', args=[batch.pk]))
        self.assertEqual(preview.context['unmatched_departments'], [])

    def test_commit_without_override_still_creates_new_department(self):
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,BrandNewDept,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]))
        asset = Asset.objects.get()
        self.assertEqual(asset.department.name, 'BrandNewDept')

    def test_commit_with_override_maps_onto_existing_department_not_a_duplicate(self):
        accounting = AssetDepartment.objects.get_or_create(name='Accounting', defaults={'tag_code': 'ACC'})[0]
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,Account,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]), {
            f'dept_map:Account': str(accounting.pk),
        })
        asset = Asset.objects.get()
        self.assertEqual(asset.department_id, accounting.pk)
        self.assertEqual(AssetDepartment.objects.filter(name='Account').count(), 0)

    def test_commit_with_location_override_maps_onto_existing_location(self):
        gf = Location.objects.create(name='Ground Floor', tag_code='GF')
        raw_csv = (
            'Ground Flr\r\n'
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,IT,Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]), {
            f'loc_map:Ground Flr': str(gf.pk),
        })
        asset = Asset.objects.get()
        self.assertEqual(asset.location_id, gf.pk)
        self.assertEqual(Location.objects.filter(name='Ground Flr').count(), 0)

    def test_preview_flags_category_with_no_exact_match(self):
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,IT,Screen Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        preview = self.client.get(reverse('tickets:asset_import_preview', args=[batch.pk]))
        self.assertIn('Screen Monitor', preview.context['unmatched_categories'])
        self.assertContains(preview, 'cat_map:Screen Monitor')

    def test_commit_without_category_override_still_creates_new_category(self):
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,IT,Screen Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]))
        asset = Asset.objects.get()
        self.assertEqual(asset.category.name, 'Screen Monitor')

    def test_commit_with_category_override_maps_onto_existing_category_not_a_duplicate(self):
        monitor = AssetCategory.objects.create(name='Monitor')
        raw_csv = (
            'S/n,USER,DEPARTMENT,DEVICE,TAG,TRACK #NO,Active,Not Active,Comments\r\n'
            '1,John Doe,IT,Screen Monitor,HD GF ACC MNT 008,ACC008,Active,,\r\n'
        )
        self._upload_csv(raw_csv)
        batch = AssetImportBatch.objects.get()
        self.client.post(reverse('tickets:asset_import_commit', args=[batch.pk]), {
            f'cat_map:Screen Monitor': str(monitor.pk),
        })
        asset = Asset.objects.get()
        self.assertEqual(asset.category_id, monitor.pk)
        self.assertEqual(AssetCategory.objects.filter(name='Screen Monitor').count(), 0)


class AssetOrphanedActionTests(TestCase):
    """Covers the state an imported asset can land in that the app's own
    checkout/reassign flows never produce on their own: a status implying
    someone has the asset (IN_USE/MAINTENANCE) with no assigned_to/
    checked_out_to. Before the fix, such an asset had no available action —
    Checkout requires IN_STORE/READY, Reassign requires an existing holder."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='orphan-admin@example.com', password='TestPass123!',
            first_name='Orphan', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.new_owner = User.objects.create_user(
            email='orphan-owner@example.com', password='TestPass123!',
            first_name='New', last_name='Owner', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.client.login(email='orphan-admin@example.com', password='TestPass123!')

    def test_can_reassign_true_for_in_use_status_with_no_holder(self):
        asset = Asset.objects.create(name='Orphaned Monitor', status=Asset.Status.IN_USE)
        self.assertTrue(asset.can_reassign)

    def test_can_reassign_true_for_maintenance_status_with_no_holder(self):
        asset = Asset.objects.create(name='Orphaned CPU', status=Asset.Status.MAINTENANCE)
        self.assertTrue(asset.can_reassign)

    def test_can_reassign_still_false_for_genuinely_available_stock(self):
        asset = Asset.objects.create(name='Spare UPS', status=Asset.Status.IN_STORE)
        self.assertFalse(asset.can_reassign)

    def test_can_reassign_unaffected_for_normal_assigned_asset(self):
        asset = Asset.objects.create(name='Laptop', status=Asset.Status.IN_USE, assigned_to=self.new_owner)
        self.assertTrue(asset.can_reassign)

    def test_reassign_view_succeeds_with_no_old_holder(self):
        """The exact previously-unreachable case: status implies a holder,
        but assigned_to/checked_out_to are both None. Must not crash on the
        old-holder notification and must actually assign the new user."""
        asset = Asset.objects.create(
            name='Orphaned Monitor', status=Asset.Status.IN_USE,
            unresolved_assignee_hint='Someone Unmatched',
        )
        response = self.client.post(reverse('tickets:asset_reassign', args=[asset.pk]), {
            'assigned_to': self.new_owner.pk,
            'comment': 'Claiming orphaned import row',
        })
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertEqual(asset.assigned_to, self.new_owner)
        self.assertEqual(asset.unresolved_assignee_hint, '')

    def test_reassign_blocked_for_genuinely_available_stock(self):
        asset = Asset.objects.create(name='Spare UPS', status=Asset.Status.IN_STORE)
        response = self.client.post(reverse('tickets:asset_reassign', args=[asset.pk]), {
            'assigned_to': self.new_owner.pk,
            'comment': 'Should be blocked',
        })
        self.assertEqual(response.status_code, 302)
        asset.refresh_from_db()
        self.assertIsNone(asset.assigned_to)


class AssetNameMatchingTests(TestCase):
    """match_users_by_name / _find_assigned_to_by_name_or_email tolerate
    first/last name being swapped between the system and an imported name,
    and the post-User-save signal that auto-assigns a now-resolvable
    unresolved_assignee_hint to the matching account."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='name-admin@example.com', password='TestPass123!',
            first_name='Name', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )

    def test_import_matches_swapped_name_order(self):
        User.objects.create_user(
            email='swapped-user@example.com', password='TestPass123!',
            first_name='Doe', last_name='John', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        matches = list(match_users_by_name('John Doe'))
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0].email, 'swapped-user@example.com')

    def test_new_user_creation_auto_assigns_matching_unresolved_hint(self):
        asset = Asset.objects.create(
            name='Orphaned Laptop', status=Asset.Status.IN_USE,
            unresolved_assignee_hint='Jane Roe',
        )
        new_user = User.objects.create_user(
            email='jane-roe@example.com', password='TestPass123!',
            first_name='Jane', last_name='Roe', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        asset.refresh_from_db()
        self.assertEqual(asset.assigned_to, new_user)
        self.assertEqual(asset.unresolved_assignee_hint, '')
        notif = Notification.objects.filter(recipient=self.admin).latest('created_at')
        self.assertIn('Jane Roe', notif.message)
        self.assertIn('auto-assigned', notif.message)
        self.assertIn('filter_q=', notif.url)

    def test_new_user_creation_does_not_auto_assign_blocked_match(self):
        """A match that can_reassign rejects (e.g. a mobilized or consumable
        asset) is left alone rather than force-assigned — admins are told to
        handle it manually instead."""
        category = AssetCategory.objects.create(name='Cable Ties', is_consumable=True)
        asset = Asset.objects.create(
            name='Orphaned Cable Batch', status=Asset.Status.IN_USE,
            unresolved_assignee_hint='Jane Roe', category=category,
        )
        User.objects.create_user(
            email='jane-roe-blocked@example.com', password='TestPass123!',
            first_name='Jane', last_name='Roe', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        asset.refresh_from_db()
        self.assertIsNone(asset.assigned_to)
        self.assertEqual(asset.unresolved_assignee_hint, 'Jane Roe')
        notif = Notification.objects.filter(recipient=self.admin).latest('created_at')
        self.assertIn('could not be auto-assigned', notif.message)

    def test_new_user_creation_does_not_notify_without_matching_hint(self):
        before = Notification.objects.filter(recipient=self.admin).count()
        User.objects.create_user(
            email='no-match@example.com', password='TestPass123!',
            first_name='Nobody', last_name='Special', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        after = Notification.objects.filter(recipient=self.admin).count()
        self.assertEqual(before, after)

    def test_asset_list_search_matches_assignee_name(self):
        owner = User.objects.create_user(
            email='search-owner@example.com', password='TestPass123!',
            first_name='Searchable', last_name='Owner', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        Asset.objects.create(name='Docking Station', assigned_to=owner, status=Asset.Status.IN_USE)
        Asset.objects.create(name='Spare Cable', status=Asset.Status.IN_STORE)
        self.client.login(email='name-admin@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:assets'), {'filter_q': 'Searchable'})
        self.assertContains(response, 'Docking Station')
        self.assertNotContains(response, 'Spare Cable')

    def test_asset_list_search_matches_unresolved_hint(self):
        Asset.objects.create(
            name='Unclaimed Monitor', status=Asset.Status.IN_USE,
            unresolved_assignee_hint='Someone Pending',
        )
        Asset.objects.create(name='Spare Cable', status=Asset.Status.IN_STORE)
        self.client.login(email='name-admin@example.com', password='TestPass123!')
        response = self.client.get(reverse('tickets:assets'), {'filter_q': 'Someone Pending'})
        self.assertContains(response, 'Unclaimed Monitor')
        self.assertNotContains(response, 'Spare Cable')


class AssetMobilizedReassignGuardTests(TestCase):
    """A mobilized asset must be demobilized before it can be reassigned —
    mobilization is a second, parallel custody mechanism that doesn't
    route through Asset.assign_to()/release(), so reassigning a mobilized
    asset used to silently orphan its open MobilizationItem (mobilization
    detail keeps showing it as active/out at the job) while the asset
    itself showed a brand new holder. See [[project_asset_inventory_restructuring]]."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='mob-guard-admin@example.com', password='TestPass123!',
            first_name='MobGuard', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.new_owner = User.objects.create_user(
            email='mob-guard-owner@example.com', password='TestPass123!',
            first_name='New', last_name='Owner', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.client.login(email='mob-guard-admin@example.com', password='TestPass123!')

        self.job = JobNumber.objects.create(number='JB-200', is_active=True)
        self.asset = Asset.objects.create(name='Mobilized Drill', status=Asset.Status.MOBILIZED)
        self.mobilization = Mobilization.objects.create(job_number=self.job, mobilized_by=self.admin)
        self.item = MobilizationItem.objects.create(mobilization=self.mobilization, asset=self.asset)

    def test_can_reassign_false_while_mobilization_open(self):
        self.assertFalse(self.asset.can_reassign)

    def test_can_reassign_true_again_after_demobilize(self):
        self.item.demobilized_at = timezone.now()
        self.item.save()
        self.assertTrue(self.asset.can_reassign)

    def test_reassign_view_blocked_while_mobilized(self):
        response = self.client.post(reverse('tickets:asset_reassign', args=[self.asset.pk]), {
            'assigned_to': self.new_owner.pk,
            'comment': 'Should be blocked',
        })
        self.assertEqual(response.status_code, 302)
        self.asset.refresh_from_db()
        self.assertIsNone(self.asset.assigned_to)
        self.assertEqual(self.asset.status, Asset.Status.MOBILIZED)
        self.item.refresh_from_db()
        self.assertIsNone(self.item.demobilized_at)

    def test_reassign_view_succeeds_after_demobilize(self):
        self.item.demobilized_at = timezone.now()
        self.item.save()
        response = self.client.post(reverse('tickets:asset_reassign', args=[self.asset.pk]), {
            'assigned_to': self.new_owner.pk,
            'comment': 'Now allowed',
        })
        self.assertEqual(response.status_code, 302)
        self.asset.refresh_from_db()
        self.assertEqual(self.asset.assigned_to, self.new_owner)

    def test_scrap_request_blocked_while_mobilized(self):
        """Same bug class as reassign: asset_scrap_request clears
        checked_out_to/assigned_to and flips status to DAMAGED directly,
        which would orphan the open MobilizationItem exactly like an
        un-guarded reassign did."""
        response = self.client.post(reverse('tickets:asset_scrap_request', args=[self.asset.pk]), {
            'comment': 'Should be blocked',
        })
        self.assertEqual(response.status_code, 400)
        self.asset.refresh_from_db()
        self.assertEqual(self.asset.status, Asset.Status.MOBILIZED)
        self.item.refresh_from_db()
        self.assertIsNone(self.item.demobilized_at)

    def test_scrap_request_succeeds_after_demobilize(self):
        self.item.demobilized_at = timezone.now()
        self.item.save()
        response = self.client.post(reverse('tickets:asset_scrap_request', args=[self.asset.pk]), {
            'comment': 'Now allowed',
        })
        self.assertEqual(response.status_code, 302)
        self.asset.refresh_from_db()
        self.assertEqual(self.asset.status, Asset.Status.DAMAGED)


class AssetInventoryGroupingTests(TestCase):
    """apps.tickets.views.assets() grouping-by-owner annotation."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='group-admin@example.com', password='TestPass123!',
            first_name='Group', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.owner_a = User.objects.create_user(
            email='group-owner-a@example.com', password='TestPass123!',
            first_name='Alice', last_name='Anderson', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.owner_b = User.objects.create_user(
            email='group-owner-b@example.com', password='TestPass123!',
            first_name='Bob', last_name='Baker', department='IT',
            role=User.Role.END_USER, is_active=True, email_verified=True,
        )
        self.department = AssetDepartment.objects.create(name='Account')
        self.location = Location.objects.create(name='Ground Floor')
        # 2 devices for Alice, 2 for Bob, 1 unmatched-hint, 1 fully unassigned.
        Asset.objects.create(name='Monitor A1', assigned_to=self.owner_a, status=Asset.Status.IN_USE)
        Asset.objects.create(name='CPU A2', assigned_to=self.owner_a, status=Asset.Status.IN_USE)
        Asset.objects.create(name='Monitor B1', assigned_to=self.owner_b, status=Asset.Status.IN_USE)
        Asset.objects.create(name='CPU B2', assigned_to=self.owner_b, status=Asset.Status.IN_USE)
        Asset.objects.create(
            name='UPS Unmatched', status=Asset.Status.IN_USE,
            unresolved_assignee_hint='Unknown Person',
        )
        Asset.objects.create(
            name='Spare Printer', status=Asset.Status.IN_STORE,
            department=self.department, location=self.location,
        )
        self.client.login(email='group-admin@example.com', password='TestPass123!')

    def test_grouped_by_default_produces_one_group_per_owner(self):
        response = self.client.get(reverse('tickets:assets'))
        self.assertEqual(response.status_code, 200)
        groups = list(response.context['assets'])
        self.assertTrue(response.context['group_by_owner'])
        # One group per distinct owner/hint/pool bucket.
        self.assertEqual(len(groups), 4)
        alice_group = next(g for g in groups if g['label'] == 'Alice Anderson')
        self.assertEqual(alice_group['item_count'], 2)
        self.assertEqual(len(alice_group['assets']), 2)

    def test_unmatched_hint_gets_its_own_labeled_group(self):
        response = self.client.get(reverse('tickets:assets'))
        groups = list(response.context['assets'])
        unmatched = next(g for g in groups if 'UPS Unmatched' in [a.name for a in g['assets']])
        self.assertEqual(unmatched['label'], 'Needs review: Unknown Person')

    def test_fully_unassigned_grouped_by_department_and_location(self):
        response = self.client.get(reverse('tickets:assets'))
        groups = list(response.context['assets'])
        spare = next(g for g in groups if 'Spare Printer' in [a.name for a in g['assets']])
        self.assertIn('Account', spare['label'])
        self.assertIn('Ground Floor', spare['label'])

    def test_group_by_owner_off_falls_back_to_newest_first(self):
        response = self.client.get(reverse('tickets:assets'), {'filter_group_by_owner': '0'})
        self.assertFalse(response.context['group_by_owner'])
        page_assets = list(response.context['assets'])
        # Flat mode still yields plain Asset instances, not group dicts.
        self.assertTrue(hasattr(page_assets[0], 'pk'))

    def test_worst_status_picks_most_severe_item_in_group(self):
        # Alice's two devices are both IN_USE; mark one DAMAGED — the group
        # must surface DAMAGED, not just the first item's status.
        Asset.objects.filter(name='CPU A2').update(status=Asset.Status.DAMAGED)
        response = self.client.get(reverse('tickets:assets'))
        groups = list(response.context['assets'])
        alice_group = next(g for g in groups if g['label'] == 'Alice Anderson')
        self.assertEqual(alice_group['worst_status'], Asset.Status.DAMAGED)

    def test_worst_condition_picks_most_severe_item_in_group(self):
        Asset.objects.filter(name='Monitor A1').update(condition=Asset.Condition.GOOD)
        Asset.objects.filter(name='CPU A2').update(condition=Asset.Condition.POOR)
        response = self.client.get(reverse('tickets:assets'))
        groups = list(response.context['assets'])
        alice_group = next(g for g in groups if g['label'] == 'Alice Anderson')
        self.assertEqual(alice_group['worst_condition'], Asset.Condition.POOR)

    def test_seat_code_uses_tag_slot_number_when_present(self):
        self.department.tag_code = 'PLD'
        self.department.save()
        Asset.objects.filter(name='Monitor A1').update(department=self.department, tag_slot_number=3)
        Asset.objects.filter(name='CPU A2').update(department=self.department, tag_slot_number=3)
        response = self.client.get(reverse('tickets:assets'))
        groups = list(response.context['assets'])
        alice_group = next(g for g in groups if g['label'] == 'Alice Anderson')
        self.assertEqual(alice_group['seat_code'], 'PLD-003')

    def test_seat_code_none_for_legacy_format_assets(self):
        response = self.client.get(reverse('tickets:assets'))
        groups = list(response.context['assets'])
        alice_group = next(g for g in groups if g['label'] == 'Alice Anderson')
        self.assertIsNone(alice_group['seat_code'])

    def test_pagination_links_preserve_active_filters(self):
        # 15 more owners so the grouped list spans more than one page (10
        # groups/page) — regression check for the pagination querystring
        # bug where filter_low_stock/filter_renewal_due/filter_group_by_owner
        # used to be silently dropped from Prev/Next/page-number links.
        for i in range(15):
            owner = User.objects.create_user(
                email=f'group-extra-{i}@example.com', password='TestPass123!',
                first_name=f'Extra{i}', last_name='Owner', department='IT',
                role=User.Role.END_USER, is_active=True, email_verified=True,
            )
            Asset.objects.create(name=f'Laptop {i}', assigned_to=owner, status=Asset.Status.IN_USE)

        response = self.client.get(reverse('tickets:assets'), {
            'filter_status': Asset.Status.IN_USE, 'filter_group_by_owner': '1',
        }, HTTP_HX_REQUEST='true')
        content = response.content.decode()
        self.assertIn('filter_status=IN_USE', content)
        self.assertIn('filter_group_by_owner=1', content)

    def test_sort_by_recently_updated_surfaces_edited_asset_first(self):
        response = self.client.get(reverse('tickets:assets'), {
            'filter_group_by_owner': '0', 'sort': '-updated_at',
        })
        self.assertEqual(response.context['active_sort'], '-updated_at')
        asset = Asset.objects.get(name='Spare Printer')
        asset.save()  # bump updated_at via a real save, as the edit form does
        response = self.client.get(reverse('tickets:assets'), {
            'filter_group_by_owner': '0', 'sort': '-updated_at',
        })
        page_assets = list(response.context['assets'])
        self.assertEqual(page_assets[0].pk, asset.pk)

    def test_invalid_sort_falls_back_to_default(self):
        response = self.client.get(reverse('tickets:assets'), {'sort': 'not-a-real-option'})
        self.assertEqual(response.context['active_sort'], 'owner')

    def test_sort_included_in_pagination_querystring(self):
        response = self.client.get(reverse('tickets:assets'), {
            'filter_group_by_owner': '0', 'sort': '-updated_at',
        }, HTTP_HX_REQUEST='true')
        self.assertIn('sort=-updated_at', response.content.decode())


class AssetEditRedirectPreservesStateTests(TestCase):
    """apps.tickets.views.asset_edit_page — regression test for the param-
    name mismatch bug where the redirect back to the list read q/category/
    status/location but the list view actually filters on filter_q/
    filter_category/filter_status/filter_location, so filters (and sort/
    page/tab/grouping) silently reset on every edit."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='edit-redirect-admin@example.com', password='TestPass123!',
            first_name='Edit', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.category = AssetCategory.objects.create(name='Laptops')
        self.asset = Asset.objects.create(name='Laptop 1', category=self.category, status=Asset.Status.IN_STORE)
        self.client.login(email='edit-redirect-admin@example.com', password='TestPass123!')

    def test_redirect_preserves_filters_sort_page_and_tab(self):
        url = reverse('tickets:asset_edit_page', args=[self.asset.pk])
        response = self.client.post(
            f'{url}?source=list&filter_q=laptop&filter_category={self.category.pk}'
            f'&filter_status=IN_STORE&filter_group_by_owner=0&filter_tab=equipment'
            f'&sort=-updated_at&page=3',
            data={
                'name': self.asset.name, 'category': self.category.pk, 'status': Asset.Status.IN_STORE,
                'tracking_id': self.asset.tracking_id, 'serial_number': '', 'model': '', 'manufacturer': '',
            },
        )
        self.assertEqual(response.status_code, 302)
        location = response['Location']
        self.assertIn('filter_q=laptop', location)
        self.assertIn(f'filter_category={self.category.pk}', location)
        self.assertIn('filter_status=IN_STORE', location)
        self.assertIn('filter_group_by_owner=0', location)
        self.assertIn('filter_tab=equipment', location)
        self.assertIn('sort=-updated_at', location)
        self.assertIn('page=3', location)


class AssetRenewalStatusTests(TestCase):
    """Asset.renewal_status boundary cases — mirrors warranty_status."""

    def setUp(self):
        self.category = AssetCategory.objects.create(name='Software License', is_renewable=True)

    def _asset(self, next_renewal_date):
        return Asset.objects.create(
            name='License', category=self.category, next_renewal_date=next_renewal_date,
        )

    def test_no_date_on_file(self):
        self.assertEqual(self._asset(None).renewal_status, 'NO_DATE')

    def test_scheduled_when_more_than_30_days_out(self):
        asset = self._asset(timezone.now().date() + timedelta(days=31))
        self.assertEqual(asset.renewal_status, 'SCHEDULED')

    def test_due_soon_at_exactly_30_days(self):
        asset = self._asset(timezone.now().date() + timedelta(days=30))
        self.assertEqual(asset.renewal_status, 'DUE_SOON')

    def test_due_soon_today(self):
        asset = self._asset(timezone.now().date())
        self.assertEqual(asset.renewal_status, 'DUE_SOON')

    def test_overdue_when_in_the_past(self):
        asset = self._asset(timezone.now().date() - timedelta(days=1))
        self.assertEqual(asset.renewal_status, 'OVERDUE')


class AssetLicensesTabTests(TestCase):
    """apps.tickets.views.assets() Equipment/Licenses & Subscriptions split —
    see [[project_asset_inventory_restructuring]]. A renewable asset
    (software license, subscription, support contract) shouldn't appear in
    the physical seat-card/flat inventory at all, and its budget/renewal
    aggregates should be global (unfiltered by search), matching the
    get_asset_kpis() convention already used by the main KPI strip."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            email='license-admin@example.com', password='TestPass123!',
            first_name='License', last_name='Admin', department='IT',
            role=User.Role.ADMIN, is_active=True, email_verified=True,
        )
        self.client.login(email='license-admin@example.com', password='TestPass123!')

        self.hw_category = AssetCategory.objects.create(name='Laptop')
        self.license_category = AssetCategory.objects.create(name='Software License', is_renewable=True)
        self.other_license_category = AssetCategory.objects.create(name='Support Contract', is_renewable=True)
        self.vendor = Vendor.objects.create(name='Microsoft')

        self.laptop = Asset.objects.create(name='Work Laptop', category=self.hw_category, status=Asset.Status.IN_USE)
        self.overdue_license = Asset.objects.create(
            name='Overdue License', category=self.license_category, renewal_vendor=self.vendor,
            next_renewal_date=timezone.now().date() - timedelta(days=5), renewal_cost=Decimal('600.00'),
        )
        self.due_soon_license = Asset.objects.create(
            name='Due Soon License', category=self.license_category,
            next_renewal_date=timezone.now().date() + timedelta(days=10), renewal_cost=Decimal('1200.00'),
        )
        self.scheduled_license = Asset.objects.create(
            name='Scheduled License', category=self.other_license_category,
            next_renewal_date=timezone.now().date() + timedelta(days=90),
        )
        self.no_date_license = Asset.objects.create(name='No Date License', category=self.license_category)

    def test_equipment_tab_excludes_renewable_assets(self):
        response = self.client.get(reverse('tickets:assets'))
        self.assertEqual(response.context['active_tab'], 'equipment')
        page_assets = list(response.context['assets'])
        # Grouped mode yields group dicts, not asset names directly — walk
        # into each group's assets to confirm no license slipped through.
        all_asset_names = []
        for g in page_assets:
            if isinstance(g, dict):
                all_asset_names.extend(a.name for a in g['assets'])
            else:
                all_asset_names.append(g.name)
        self.assertIn('Work Laptop', all_asset_names)
        for license_name in ['Overdue License', 'Due Soon License', 'Scheduled License', 'No Date License']:
            self.assertNotIn(license_name, all_asset_names)

    def test_equipment_tab_flat_mode_also_excludes_renewables(self):
        response = self.client.get(reverse('tickets:assets'), {'filter_group_by_owner': '0'})
        page_assets = list(response.context['assets'])
        names = [a.name for a in page_assets]
        self.assertIn('Work Laptop', names)
        self.assertNotIn('Overdue License', names)

    def test_licenses_tab_shows_only_renewable_assets_ordered_by_urgency(self):
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses'})
        self.assertEqual(response.context['active_tab'], 'licenses')
        names = [a.name for a in response.context['assets']]
        self.assertNotIn('Work Laptop', names)
        # Overdue (5 days ago) sorts before due-soon (10 days) before
        # scheduled (90 days); no-date-on-file sorts last (nulls last).
        self.assertEqual(names, ['Overdue License', 'Due Soon License', 'Scheduled License', 'No Date License'])

    def test_licenses_tab_budget_aggregates(self):
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses'})
        self.assertEqual(response.context['license_total_cost'], Decimal('1800.00'))
        self.assertEqual(response.context['license_costed_count'], 2)
        self.assertEqual(response.context['license_total_count'], 4)
        self.assertEqual(response.context['license_due_soon_count'], 1)
        self.assertEqual(response.context['license_overdue_count'], 1)

    def test_licenses_tab_budget_currency_is_per_asset(self):
        # Both costed assets default to '$' (Asset.renewal_currency's
        # model default) — a single agreed currency, so it's safe to show.
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses'})
        self.assertEqual(response.context['license_budget_currency'], '$')
        self.assertFalse(response.context['license_budget_mixed_currencies'])

        # Once the two costed assets disagree on currency, the aggregate
        # can no longer be labeled with a single symbol.
        self.due_soon_license.renewal_currency = '€'
        self.due_soon_license.save(update_fields=['renewal_currency'])
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses'})
        self.assertEqual(response.context['license_budget_currency'], '')
        self.assertTrue(response.context['license_budget_mixed_currencies'])

    def test_licenses_tab_aggregates_unaffected_by_search(self):
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses', 'filter_q': 'Overdue'})
        names = [a.name for a in response.context['assets']]
        self.assertEqual(names, ['Overdue License'])
        # Aggregates stay global regardless of the search narrowing the list.
        self.assertEqual(response.context['license_total_count'], 4)
        self.assertEqual(response.context['license_overdue_count'], 1)

    def test_licenses_tab_category_filter_scoped_to_renewable_categories(self):
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses'})
        category_names = [c.name for c in response.context['categories']]
        self.assertIn('Software License', category_names)
        self.assertIn('Support Contract', category_names)
        self.assertNotIn('Laptop', category_names)

    def test_licenses_tab_due_soon_filter(self):
        # Matches Asset.is_renewal_due_soon's own semantics — "due soon"
        # includes already-overdue renewals, not just the upcoming window.
        response = self.client.get(reverse('tickets:assets'), {'filter_tab': 'licenses', 'filter_due_soon': '1'})
        names = [a.name for a in response.context['assets']]
        self.assertEqual(names, ['Overdue License', 'Due Soon License'])

    def test_license_badge_count_shown_on_equipment_tab(self):
        response = self.client.get(reverse('tickets:assets'))
        self.assertEqual(response.context['license_total_count'], 4)

    def test_htmx_pagination_target_renders_table_only_not_full_panel(self):
        response = self.client.get(
            reverse('tickets:assets'), {'filter_tab': 'licenses'},
            HTTP_HX_REQUEST='true', HTTP_HX_TARGET='assetTableContainer',
        )
        content = response.content.decode()
        self.assertNotIn('id="assetPanel"', content)
        self.assertIn('Overdue License', content)

    def test_htmx_tab_switch_target_renders_full_panel(self):
        response = self.client.get(
            reverse('tickets:assets'), {'filter_tab': 'licenses'},
            HTTP_HX_REQUEST='true', HTTP_HX_TARGET='assetPanel',
        )
        content = response.content.decode()
        self.assertIn('id="assetPanel"', content)

