# apps/tickets/report_exporters.py
"""
Format writers shared by every report type in the Exportables feature.
Each report type's queryset/row-builder (report_registry.py) is generic
over these — adding a new export format means adding one function here and
registering it in EXPORTERS, not touching every report type.
"""
import csv
import json
import os
from io import BytesIO

import requests
from django.conf import settings
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from playwright.sync_api import sync_playwright
from xhtml2pdf import pisa

PRIMARY_RGB = RGBColor(0x0D, 0x94, 0x88)
META_GRAY_RGB = RGBColor(0x47, 0x55, 0x69)


def _filename(base, include_timestamp=True):
    """Every exported file's name — the one place this is formatted, so
    every export format/report type gets it for free. Shape:
    {COMPANY-INITIALS}-{CONTEXT}-{DATE}-{TIME}, e.g.
    "HDG-ASSETS-BY-PERSON-20260827-143022" — company_initials is
    blank-safe (no prefix/leading hyphen if unset). `include_timestamp`
    is dropped for single-record form exports (incident/service
    request/maintenance/asset/procurement), since the record's own
    number/tracking_id already makes the filename unique — e.g.
    "HDG-INCIDENT-INC-2026-00001". List/report exports (CSV/Excel/PDF/
    DOCX of a filtered set) keep the timestamp, since re-running the same
    filter has no other unique identifier."""
    from apps.accounts.models import ClientSettings
    prefix = ClientSettings.objects.get_or_create(id=1)[0].company_initials.strip().upper()
    context = base.upper().replace('_', '-').replace(' ', '-')
    parts = [prefix, context]
    if include_timestamp:
        parts.append(timezone.now().strftime('%Y%m%d-%H%M%S'))
    return '-'.join(part for part in parts if part)


def _pdf_link_callback(uri, rel):
    """Resolve /media/ and /static/ URLs in the PDF template to absolute
    filesystem paths — xhtml2pdf can't fetch relative URLs itself (it has
    no request context), so without this the letterhead logo silently
    fails to render."""
    if uri.startswith(settings.MEDIA_URL):
        path = os.path.join(settings.MEDIA_ROOT, uri.replace(settings.MEDIA_URL, '', 1))
    elif uri.startswith(settings.STATIC_URL):
        static_root = settings.STATIC_ROOT or os.path.join(settings.BASE_DIR, 'static')
        path = os.path.join(static_root, uri.replace(settings.STATIC_URL, '', 1))
    else:
        return uri
    return path if os.path.isfile(path) else uri


def _csv_safe(value):
    """Neutralize spreadsheet formula injection: a string value starting
    with =, +, -, or @ is interpreted as a live formula by Excel/Sheets
    when the exported file is later opened by another admin. Leaves
    non-string values (numbers, None, ...) untouched so exports keep their
    real type in Excel rather than becoming text."""
    if isinstance(value, str) and value and value[0] in ('=', '+', '-', '@'):
        return "'" + value
    return value


def export_csv(rows, columns, title, filename_base, control_number='', **kwargs):
    from apps.accounts.models import ClientSettings
    client_settings = ClientSettings.objects.first()

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base)}.csv"'
    writer = csv.writer(response)
    # Document header block, same info every other export format carries in
    # its letterhead — company name, a plain-language title so the file's
    # contents are obvious without opening it, and the org's document
    # control number for audit/filing purposes.
    writer.writerow([client_settings.company_name if client_settings else 'My Company'])
    writer.writerow([f'{(title or "Report").upper()} - DATA EXPORT REPORT'])
    if control_number:
        writer.writerow([f'Control No: {control_number}'])
    writer.writerow([f'Generated: {timezone.now().strftime("%B %d, %Y %H:%M")}'])
    writer.writerow([f'Records: {len(rows)}'])
    writer.writerow([])

    # extrasaction='ignore': `rows` dicts always carry every field from
    # row_from_obj(), but `columns` may be a caller-narrowed subset (the
    # export_menu.html column picker) — DictWriter otherwise raises on the
    # extra keys instead of just dropping them.
    dict_writer = csv.DictWriter(response, fieldnames=columns, extrasaction='ignore')
    dict_writer.writeheader()
    for row in rows:
        dict_writer.writerow({key: _csv_safe(value) for key, value in row.items()})
    return response


_EXCEL_FONT = 'Calibri'  # matches every DOCX export's body font (report_exporters._docx_*)
_EXCEL_HEADER_FILL = PatternFill(start_color='0D9488', end_color='0D9488', fill_type='solid')
_EXCEL_BANDED_FILL = PatternFill(start_color='F8FAFC', end_color='F8FAFC', fill_type='solid')
_EXCEL_BORDER_COLOR = 'CBD5E1'


def export_excel(rows, columns, title, filename_base, control_number='', **kwargs):
    from openpyxl.styles import Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from apps.accounts.models import ClientSettings

    wb = Workbook()
    ws = wb.active
    ws.title = (title or 'Report')[:31]

    thin = Side(style='thin', color=_EXCEL_BORDER_COLOR)
    cell_border = Border(top=thin, bottom=thin, left=thin, right=thin)
    last_col_letter = get_column_letter(len(columns))

    # Document letterhead — a boxed 3-zone banner (logo | title | control
    # number) over a meta line, matching the org's actual paper-form layout
    # (logo box, centered title, highlighted control-number box) rather than
    # a few bare unstyled text rows. Same visual language as the PDF/DOCX
    # exports' own boxed letterheads (report_pdf.html, _docx_report_letterhead).
    client_settings = ClientSettings.objects.first()
    total_cols = len(columns)
    logo_cols = max(1, min(3, total_cols // 6))
    control_cols = max(1, min(3, total_cols // 6))
    title_cols = max(1, total_cols - logo_cols - control_cols)
    logo_end = get_column_letter(logo_cols)
    title_start = get_column_letter(logo_cols + 1)
    title_end = get_column_letter(logo_cols + title_cols)
    control_start = get_column_letter(logo_cols + title_cols + 1)

    black_thin = Side(style='thin', color='000000')
    box_border = Border(top=black_thin, bottom=black_thin, left=black_thin, right=black_thin)
    BANNER_ROWS = 3  # logo/title/control zone height, in sheet rows

    for row in range(1, BANNER_ROWS + 1):
        ws.row_dimensions[row].height = 22

    ws.merge_cells(f'A1:{logo_end}{BANNER_ROWS}')
    logo_cell = ws['A1']
    logo_cell.border = box_border
    logo_cell.alignment = Alignment(horizontal='center', vertical='center')
    logo_source = _docx_image_source(client_settings.logo) if client_settings and client_settings.logo else None
    if logo_source:
        try:
            from openpyxl.drawing.image import Image as XLImage
            img = XLImage(logo_source)
            img.height, img.width = 40, int(40 * img.width / img.height)
            img.anchor = 'A1'
            ws.add_image(img)
        except Exception:
            logo_source = None
    if not logo_source:
        logo_cell.value = client_settings.company_name if client_settings else 'My Company'
        logo_cell.font = Font(name=_EXCEL_FONT, size=11, bold=True, color='0D9488')

    ws.merge_cells(f'{title_start}1:{title_end}{BANNER_ROWS}')
    title_cell = ws[f'{title_start}1']
    title_cell.value = f'{(title or "Report").upper()}\nDATA EXPORT REPORT'
    title_cell.font = Font(name=_EXCEL_FONT, size=15, bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    title_cell.border = box_border
    # Border every cell in the merged range too — Excel only draws the
    # anchor cell's border by default, leaving the rest of a merged box
    # looking unbordered.
    for col_idx in range(logo_cols + 1, logo_cols + title_cols + 1):
        ws.cell(row=1, column=col_idx).border = box_border
        ws.cell(row=BANNER_ROWS, column=col_idx).border = box_border

    ws.merge_cells(f'{control_start}1:{last_col_letter}{BANNER_ROWS}')
    control_cell = ws[f'{control_start}1']
    control_cell.value = f'CONTROL NO.\n{control_number}' if control_number else 'CONTROL NO.\n—'
    control_cell.font = Font(name=_EXCEL_FONT, size=11, bold=True, color='FFFFFF')
    control_cell.fill = _EXCEL_HEADER_FILL
    control_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    control_cell.border = box_border
    for col_idx in range(logo_cols + title_cols + 1, total_cols + 1):
        ws.cell(row=1, column=col_idx).border = box_border
        ws.cell(row=BANNER_ROWS, column=col_idx).border = box_border
        ws.cell(row=1, column=col_idx).fill = _EXCEL_HEADER_FILL

    meta_row = BANNER_ROWS + 1
    ws.merge_cells(f'A{meta_row}:{last_col_letter}{meta_row}')
    meta_cell = ws[f'A{meta_row}']
    meta_cell.value = f'Generated: {timezone.now().strftime("%B %d, %Y %H:%M")}   |   Records: {len(rows)}'
    meta_cell.font = Font(name=_EXCEL_FONT, size=9, color='475569')
    meta_cell.alignment = Alignment(horizontal='left', vertical='center')

    header_row = meta_row + 2  # one blank spacer row before the data table

    # Header row: bold white-on-brand-teal, matching the PDF/DOCX table
    # exports' own header styling (report_pdf.html, _docx_report_letterhead)
    # so every export format reads as the same product.
    for col, header in enumerate(columns, 1):
        cell = ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(name=_EXCEL_FONT, size=11, bold=True, color='FFFFFF')
        cell.fill = _EXCEL_HEADER_FILL
        cell.alignment = Alignment(horizontal='left', vertical='center')
        cell.border = cell_border
    ws.row_dimensions[header_row].height = 20

    # Data rows: consistent font/size throughout, thin borders, and light
    # banding on alternating rows — the same "professional inventory sheet"
    # look a reader would expect, not a bare unstyled dump of values.
    for offset, row_data in enumerate(rows):
        row_idx = header_row + 1 + offset
        banded = offset % 2 == 1
        for col_idx, key in enumerate(columns, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=_csv_safe(row_data.get(key, '')))
            cell.font = Font(name=_EXCEL_FONT, size=10.5)
            cell.alignment = Alignment(horizontal='left', vertical='center')
            cell.border = cell_border
            if banded:
                cell.fill = _EXCEL_BANDED_FILL

    # Auto-size each column from its header/content length so the sheet is
    # readable without the user manually resizing every column first.
    for col_idx, header in enumerate(columns, 1):
        letter = get_column_letter(col_idx)
        max_len = len(str(header))
        for row_data in rows[:500]:
            value = row_data.get(header, '')
            max_len = max(max_len, len(str(value)) if value is not None else 0)
        ws.column_dimensions[letter].width = min(max(max_len + 2, 10), 50)

    # Frozen header row + AutoFilter dropdowns — standard spreadsheet
    # conveniences for a table a reader will scroll and filter through.
    ws.freeze_panes = f'A{header_row + 1}'
    if rows:
        ws.auto_filter.ref = f'A{header_row}:{last_col_letter}{header_row + len(rows)}'

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base)}.xlsx"'
    wb.save(response)
    return response


def export_json(rows, columns, title, filename_base, **kwargs):
    response = HttpResponse(json.dumps(rows, indent=2, default=str), content_type='application/json')
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base)}.json"'
    return response


def export_pdf(rows, columns, title, filename_base, request=None, filter_summary='', control_number='', **kwargs):
    html = render_to_string('reports/report_pdf.html', {
        'title': title,
        'columns': columns,
        'rows': rows,
        'generated_at': timezone.now(),
        'filter_summary': filter_summary,
        'record_count': len(rows),
        'control_number': control_number,
    }, request=request)

    pdf_bytes = _html_to_pdf(html, request, paginate=True)
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base)}.pdf"'
    return response


def _render_form_pdf(template_name, context, request, filename_base):
    html = render_to_string(template_name, {**context, 'generated_at': timezone.now()}, request=request)
    result = BytesIO()
    pisa.CreatePDF(src=html, dest=result, encoding='utf-8', link_callback=_pdf_link_callback)
    response = HttpResponse(result.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base, include_timestamp=False)}.pdf"'
    return response


_PAGE_NUMBER_FOOTER = """
<div style="width:100%; font-size:7.5pt; color:#94A3B8; text-align:center; font-family:Arial,Helvetica,sans-serif;">
  Page <span class="pageNumber"></span> of <span class="totalPages"></span>
</div>
"""


def _html_to_pdf(html, request, paginate=False):
    """Render HTML to PDF bytes via headless Chromium (Playwright), so
    templates can rely on real JS/CSS execution (Tailwind CDN, flexbox)
    instead of xhtml2pdf's limited/buggy renderer. `page.set_content()`
    doesn't carry page origin like `page.goto()`, so a <base href> pointing
    at this request's host must already be present in the template's <head>
    for relative /media//static/ URLs to resolve. `paginate=True` adds a
    Chromium-native "Page X of Y" footer — useful for multi-page tabular
    reports where the page count isn't known until render time."""
    with sync_playwright() as p:
        browser = p.chromium.launch()
        try:
            page = browser.new_page()
            page.set_content(html, wait_until='networkidle')
            pdf_kwargs = {'prefer_css_page_size': True, 'print_background': True}
            if paginate:
                # An explicit `margin` here overrides the template's own
                # @page margin (Chromium ignores CSS page-margins once this
                # option is passed), so all four sides must be restated —
                # matching report_pdf.html's own @page margin, with extra
                # bottom room reserved for the page-number footer.
                pdf_kwargs.update(
                    display_header_footer=True,
                    header_template='<span></span>',
                    footer_template=_PAGE_NUMBER_FOOTER,
                    margin={'top': '1.6cm', 'left': '1.6cm', 'right': '1.6cm', 'bottom': '1.9cm'},
                )
            pdf_bytes = page.pdf(**pdf_kwargs)
        finally:
            browser.close()
    return pdf_bytes


def _render_form_pdf_chromium(template_name, context, request, filename_base):
    html = render_to_string(template_name, {**context, 'generated_at': timezone.now()}, request=request)
    pdf_bytes = _html_to_pdf(html, request)
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base, include_timestamp=False)}.pdf"'
    return response


def _docx_heading(doc, text):
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = PRIMARY_RGB
    return heading


def _docx_simple_table(doc, headers, rows, placeholder='To be completed by the IT team'):
    """Bordered table with a bold header row, matching the "blank form
    section" tables the paper-form PDFs ship (Systems Affected,
    Recommendations, etc.) — used so the DOCX has a real table there too
    instead of a placeholder sentence. `rows` empty renders one italic
    placeholder row spanning all columns, same as the PDF's blank row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    if rows:
        for row_values in rows:
            row_cells = table.add_row().cells
            for cell, value in zip(row_cells, row_values):
                cell.text = str(value) if value not in (None, '') else ''
    else:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[-1])
        row_cells[0].paragraphs[0].add_run(placeholder).italic = True
    doc.add_paragraph()
    return table


def _docx_field(doc, label, value):
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(str(value) if value not in (None, '') else '—')


def _docx_form_section(doc, title, fields, body_paragraph=None):
    """Boxed section matching the paper-form PDFs' `.section-box`/`.cell`
    styling: a shaded title bar followed by one bordered label|value row per
    field, inside a single outer-bordered table — instead of a plain heading
    with unboxed 'Label: value' paragraphs. `fields` is a list of
    (label, value) pairs; a falsy value renders as '—'. Pass `body_paragraph`
    (free text, e.g. a description) to render it as one full-width row below
    the fields, still inside the same box, instead of a separate field row."""
    rows = 1 + len(fields) + (1 if body_paragraph is not None else 0)
    table = doc.add_table(rows=rows, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(5.7)

    title_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    title_cell.text = ''
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = title_cell.paragraphs[0].add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = PRIMARY_RGB
    _shade_cell(title_cell, 'F0FDFA')
    _set_cell_border(title_cell, top={'sz': 16, 'val': 'single', 'color': '0F172A'},
                      bottom={'sz': 8, 'val': 'single', 'color': '0F172A'},
                      left={'sz': 16, 'val': 'single', 'color': '0F172A'},
                      right={'sz': 16, 'val': 'single', 'color': '0F172A'})

    row_edge = {'sz': 4, 'val': 'single', 'color': '94A3B8'}
    outer_edge = {'sz': 16, 'val': 'single', 'color': '0F172A'}
    for idx, (label, value) in enumerate(fields, start=1):
        label_cell, value_cell = table.rows[idx].cells
        label_cell.text = ''
        label_cell.paragraphs[0].add_run(label).bold = True
        value_cell.text = str(value) if value not in (None, '') else '—'
        is_last = idx == len(fields) and body_paragraph is None
        for cell in (label_cell, value_cell):
            _set_cell_border(cell, top=row_edge,
                              left=outer_edge if cell is label_cell else row_edge,
                              right=outer_edge if cell is value_cell else row_edge,
                              bottom=outer_edge if is_last else None)

    if body_paragraph is not None:
        body_cell = table.rows[-1].cells[0].merge(table.rows[-1].cells[1])
        body_cell.text = str(body_paragraph) if body_paragraph else '—'
        _set_cell_border(body_cell, top=row_edge, left=outer_edge, right=outer_edge, bottom=outer_edge)

    doc.add_paragraph()
    return table


def _set_cell_border(cell, **edges):
    """python-docx has no high-level cell-border API — draws thin borders
    on a table cell via direct OXML manipulation. edges: top/bottom/left/
    right, each a dict like {'sz': 8, 'val': 'single', 'color': '0F172A'}."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge_name, edge_data in edges.items():
        if not edge_data:
            continue
        tag = f'w:{edge_name}'
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def _box_cell(cell, sz=8, color='0F172A'):
    edge = {'sz': sz, 'val': 'single', 'color': color, 'space': '0'}
    _set_cell_border(cell, top=edge, bottom=edge, left=edge, right=edge)


def _docx_image_source(field):
    """Returns something doc.add_picture()/run.add_picture() can consume
    for a Django ImageField, regardless of storage backend: a local path
    in dev (FileSystemStorage), or fetched bytes in production (Cloudinary,
    which has no filesystem path). Returns None on any failure so callers
    can gracefully fall back to text — never let a bad/missing image crash
    an export."""
    if not field:
        return None
    try:
        return field.path
    except NotImplementedError:
        pass
    except Exception:
        return None
    try:
        resp = requests.get(field.url, timeout=5)
        resp.raise_for_status()
        return BytesIO(resp.content)
    except Exception:
        return None


def _docx_letterhead(doc, form_code, rev, form_date, title, page):
    """Boxed 3-column header matching the paper forms' layout: logo box,
    centered form title + code, and a Page/Rev/Date meta box."""
    from apps.accounts.models import ClientSettings

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    logo_cell, title_cell, meta_cell = table.rows[0].cells
    for cell, width in zip((logo_cell, title_cell, meta_cell), (Inches(1.3), Inches(3.4), Inches(1.8))):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _box_cell(cell)

    client_settings = ClientSettings.objects.first()
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_source = _docx_image_source(client_settings.logo) if client_settings and client_settings.logo else None
    if logo_source:
        logo_p.add_run().add_picture(logo_source, width=Inches(1.1))
    else:
        brand_run = logo_p.add_run(client_settings.company_name if client_settings else 'My Company')
        brand_run.bold = True
        brand_run.font.color.rgb = PRIMARY_RGB

    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    code_p = title_cell.add_paragraph()
    code_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_run = code_p.add_run(form_code)
    code_run.font.size = Pt(8)
    code_run.font.color.rgb = META_GRAY_RGB

    meta_rows = [('Page:', page), ('Rev.:', rev), ('Date:', form_date)]
    for index, (label, value) in enumerate(meta_rows):
        meta_p = meta_cell.paragraphs[0] if index == 0 else meta_cell.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(0)
        label_run = meta_p.add_run(f'{label} ')
        label_run.bold = True
        label_run.font.size = Pt(8)
        value_run = meta_p.add_run(str(value))
        value_run.font.size = Pt(8)

    doc.add_paragraph()


def _docx_report_letterhead(doc, title, generated_at, filter_summary, record_count, usable_width_in=7.8, control_number=''):
    """Boxed letterhead for the generic tabular exports (group/batch report
    exports), matching the same visual language as `_docx_letterhead` and
    the Excel/PDF exports' own letterheads: logo | title | a distinctly
    highlighted Control No. box | Generated/Filters/Records meta. Spans
    `usable_width_in` so its border aligns with the data table beneath it."""
    from apps.accounts.models import ClientSettings

    logo_w, title_w, control_w = 1.5, 3.2, 1.5
    meta_w = max(usable_width_in - logo_w - title_w - control_w, 1.6)

    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    logo_cell, title_cell, control_cell, meta_cell = table.rows[0].cells
    widths = (Inches(logo_w), Inches(title_w), Inches(control_w), Inches(meta_w))
    for cell, width in zip((logo_cell, title_cell, control_cell, meta_cell), widths):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _box_cell(cell)

    # Control No. gets its own filled box — same brand-teal highlight the
    # Excel/PDF letterheads use — rather than sitting as just another line
    # in the meta cell, so it can't be missed at a glance.
    _shade_cell(control_cell, '0D9488')
    control_p = control_cell.paragraphs[0]
    control_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    control_label_run = control_p.add_run('CONTROL NO.')
    control_label_run.bold = True
    control_label_run.font.size = Pt(7)
    control_label_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    control_value_p = control_cell.add_paragraph()
    control_value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    control_value_run = control_value_p.add_run(control_number or '—')
    control_value_run.bold = True
    control_value_run.font.size = Pt(9)
    control_value_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    client_settings = ClientSettings.objects.first()
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_source = _docx_image_source(client_settings.logo) if client_settings and client_settings.logo else None
    if logo_source:
        logo_p.add_run().add_picture(logo_source, width=Inches(1.3))
    else:
        brand_run = logo_p.add_run(client_settings.company_name if client_settings else 'My Company')
        brand_run.bold = True
        brand_run.font.color.rgb = PRIMARY_RGB

    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run((title or 'Report').upper())
    title_run.bold = True
    title_run.font.size = Pt(13)
    subtitle_p = title_cell.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run('DATA EXPORT REPORT')
    subtitle_run.font.size = Pt(8)
    subtitle_run.font.color.rgb = META_GRAY_RGB

    meta_rows = [
        ('Generated:', generated_at.strftime('%B %d, %Y %H:%M')),
        ('Filters:', filter_summary or 'None applied'),
        ('Records:', f'{record_count} record{"s" if record_count != 1 else ""}'),
    ]
    for index, (label, value) in enumerate(meta_rows):
        meta_p = meta_cell.paragraphs[0] if index == 0 else meta_cell.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(0)
        label_run = meta_p.add_run(f'{label} ')
        label_run.bold = True
        label_run.font.size = Pt(8)
        value_run = meta_p.add_run(str(value))
        value_run.font.size = Pt(8)

    doc.add_paragraph()


def _docx_signoff_table(doc, headers, rows):
    """Role | Name & Signature | Date table matching the paper form's
    Sign-off & Approvals table — each row a (role_label, signoff) pair, an
    uploaded signature embedded as an image in the signature cell when
    present, else the 'captured digitally' text fallback ('Pending' when
    nobody has signed off yet)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    for role_label, signoff in rows:
        role_cell, sig_cell, date_cell = table.add_row().cells
        role_cell.paragraphs[0].add_run(role_label).bold = True

        sig_p = sig_cell.paragraphs[0]
        if not signoff['captured_text']:
            sig_p.add_run('Pending').italic = True
        elif signoff['signature_url'] and signoff['user']:
            src = _docx_image_source(signoff['user'].signature)
            if src:
                sig_p.add_run().add_picture(src, height=Inches(0.3))
            else:
                sig_p.add_run(signoff['captured_text']).italic = True
        else:
            sig_p.add_run(signoff['captured_text']).italic = True

        date_cell.text = signoff['date'].strftime('%Y-%m-%d') if signoff['date'] else ''
    doc.add_paragraph()
    return table


def _docx_signoff_field(doc, label, signoff):
    """Renders a sign-off line: the user's uploaded signature image when
    available, else the existing 'captured digitally' text fallback."""
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    if not signoff['captured_text']:
        p.add_run('Pending')
        return
    if signoff['signature_url'] and signoff['user']:
        src = _docx_image_source(signoff['user'].signature)
        if src:
            doc.add_picture(src, height=Inches(0.35))
    caption = doc.add_paragraph()
    caption.add_run(signoff['captured_text']).italic = True


def export_incident_pdf(ticket, request):
    from .report_registry import incident_form_sections
    context = incident_form_sections(ticket)
    return _render_form_pdf_chromium('reports/incident_form_pdf.html', context, request, f'incident_{ticket.number}')


def export_incident_docx(ticket):
    from .report_registry import incident_form_sections
    ctx = incident_form_sections(ticket)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    _docx_letterhead(doc, 'HDG-IT-FRM-086', 'A01', timezone.now().strftime('%d %B %Y'), 'IT INCIDENT REPORT FORM', page='1 of 4')

    title = doc.add_heading(f'IT Incident Report — {ticket.number}', level=1)
    for run in title.runs:
        run.font.color.rgb = PRIMARY_RGB
    doc.add_paragraph(ticket.title).runs[0].italic = True

    _docx_form_section(doc, 'Section 1 — Incident Details', [
        ('Incident Reference No.', ticket.number),
        ('Date & Time of Incident', ticket.incident_datetime.strftime('%Y-%m-%d %H:%M') if ticket.incident_datetime else None),
        ('Incident Category', ctx['incident_category_display']),
        ('Severity', ticket.get_urgency_display()),
        ('Business Impact', ticket.get_business_impact_display() if ticket.business_impact else None),
        ('How Discovered', ctx['how_discovered_display']),
        ('Location / IP / Hostname', ticket.location_hostname),
        ('System / Service / Asset Affected', ticket.location_hostname),
    ])

    _docx_form_section(doc, 'Section 2 — Reporter Information', [
        ('Reported By', ticket.requester.get_full_name() or ticket.requester.email),
        ('Department / Unit', ticket.requester.get_department_display()),
        ('Job Title / Role', ticket.requester.position),
        ('Email', ticket.requester.email),
        ('Submitted', ticket.created_at.strftime('%Y-%m-%d %H:%M')),
    ])

    fields_3 = [('Immediate/Initial Actions', ticket.immediate_actions)] if ticket.immediate_actions else []
    _docx_form_section(doc, 'Section 3 — Detailed Incident Description', fields_3, body_paragraph=ticket.description)

    _docx_form_section(doc, 'Section 5 — Root Cause Analysis', [
        ('Root Cause Category', ', '.join(ctx['root_cause_category_display']) or None),
        ('Detailed Root Cause / Contributing Factors', ctx['resolution_root_cause']),
    ])

    _docx_heading(doc, 'Section 4 — Systems / Users Affected')
    _docx_simple_table(doc, ['System / Application / Device', 'No. of Users / Devices Impacted', 'Nature of Impact'], [])

    resolved = ticket.status in ('RESOLVED', 'CLOSED')
    _docx_form_section(doc, 'Section 6 — Resolution & Corrective Actions', [
        ('Resolution Status', 'Fully Resolved' if resolved else 'Pending / Escalated'),
        ('Date & Time of Resolution', ticket.resolution_confirmed_at.strftime('%Y-%m-%d %H:%M') if ticket.resolution_confirmed_at else 'Pending'),
        ('Steps Taken to Resolve the Incident', ctx['resolution_steps']),
    ])

    _docx_heading(doc, 'Section 7 — Recommendations & Preventive Actions')
    _docx_simple_table(doc, ['Recommended Action', 'Responsible Person', 'Target Date', 'Status'], [])

    _docx_form_section(doc, 'Section 8 — Communication & Distribution', [
        ('Report Communicated To', 'IT Manager'),
        ('Method of Communication', 'IT Helpdesk Ticket'),
    ])

    _docx_heading(doc, 'Section 9 — Supporting Documentation Attached')
    if ctx['attachments'] or ctx['image_attachments']:
        for a in ctx['attachments']:
            doc.add_paragraph(f'• {a.filename}')
        for a in ctx['image_attachments']:
            doc.add_paragraph(f'{a.filename}:')
            image_source = _docx_image_source(a.file)
            if image_source:
                doc.add_picture(image_source, width=Inches(6))
                doc.add_page_break()
            else:
                doc.add_paragraph('(image could not be embedded)').runs[0].italic = True
    else:
        doc.add_paragraph('No attachments.').runs[0].italic = True

    _docx_heading(doc, 'Section 10 — Sign-off & Approvals')
    _docx_signoff_table(doc, ['Role', 'Name & Signature', 'Date'], [
        ('Prepared By (Reporter)', ctx['prepared_by_signoff']),
        ('Reviewed & Approved By (IT Manager / Head of IT)', ctx['it_manager_signoff']),
        ('Confirmed Resolved By', ctx['confirmed_resolved_signoff']),
    ])

    buf = BytesIO()
    doc.save(buf)
    response = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{_filename("incident_" + ticket.number, include_timestamp=False)}.docx"'
    return response


def export_service_request_pdf(ticket, request):
    from .report_registry import service_request_form_sections
    context = service_request_form_sections(ticket)
    return _render_form_pdf_chromium('reports/service_request_form_pdf2.html', context, request, f'service_request_{ticket.number}')


def export_service_request_docx(ticket):
    from .report_registry import service_request_form_sections
    ctx = service_request_form_sections(ticket)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    _docx_letterhead(doc, 'HDG-IT-FRM-081', '1', timezone.now().strftime('%d %B %Y'), 'SERVICE REQUEST FORM', page='1 of 1')

    title = doc.add_heading(f'Service Request — {ticket.number}', level=1)
    for run in title.runs:
        run.font.color.rgb = PRIMARY_RGB
    doc.add_paragraph(ticket.title).runs[0].italic = True

    request_fields = [
        ("Requester's Name", ctx['requester_name']),
        ('Date', ticket.created_at.strftime('%Y-%m-%d')),
        ('Department', ctx['requester_department']),
        ('Location', ctx['location']['display'] or 'Not available'),
        ('Reported To', ctx['reported_to'] or 'Unassigned'),
        ('Service Category', ctx['service_category_name'] or '—'),
        ('Purpose', ctx['purpose'] or '—'),
    ] + list(ctx['dynamic_fields'])
    _docx_form_section(doc, 'Request Details', request_fields)

    _docx_form_section(doc, 'Description of Work Order Request', [], body_paragraph=ticket.description)

    if ctx['vessels'] or ctx['job_number'] or ctx['dive_systems']:
        vessel_fields = []
        if ctx['vessels']:
            vessel_fields.append(('Vessel(s)', ', '.join(v.name for v in ctx['vessels'])))
        if ctx['job_number']:
            vessel_fields.append(('Job Number', ctx['job_number'].number))
        if ctx['dive_systems']:
            vessel_fields.append(('Dive System(s)', ', '.join(s.name for s in ctx['dive_systems'])))
        _docx_form_section(doc, 'Marine / Job Details', vessel_fields)

    _docx_heading(doc, 'Confirmation & Sign-off')
    _docx_signoff_field(doc, "Requester's Confirmation", ctx['requester_signoff'])
    _docx_signoff_field(doc, 'IT Officer Signature', ctx['it_officer_signoff'])

    _docx_heading(doc, "Requester's Feedback")
    if ctx['feedback_comment'] or ctx['feedback_rating']:
        rating = f' (Rating: {ctx["feedback_rating"]}/5)' if ctx['feedback_rating'] else ''
        doc.add_paragraph(f'{ctx["feedback_comment"] or "—"}{rating}')
    else:
        doc.add_paragraph('No feedback submitted yet.').runs[0].italic = True

    _docx_signoff_field(doc, "Requester's Signature", ctx['requester_feedback_signoff'])
    _docx_signoff_field(doc, 'IT Manager Signature', ctx['it_manager_signoff'])

    buf = BytesIO()
    doc.save(buf)
    response = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{_filename("service_request_" + ticket.number, include_timestamp=False)}.docx"'
    return response


def export_procurement_request_pdf(procurement_request, request):
    # The item's own ticket, or — if it's routed through a mobilization —
    # the mobilization-request ticket that originally justified sending it
    # out (legacy ad-hoc mobilizations have no ticket, and a standalone
    # restock request has neither, in which case there's nothing to pull
    # background/delivery context from).
    source_ticket = procurement_request.ticket or (
        procurement_request.mobilization.ticket if procurement_request.mobilization_id else None
    )

    if source_ticket:
        origin_display = f'Service Request {source_ticket.number}'
        contact = source_ticket.requester
        destination_display = contact.get_department_display() or '—'
    elif procurement_request.mobilization_id:
        origin_display = f'Mobilization — {procurement_request.mobilization.destination_display}'
        contact = procurement_request.mobilization.mobilized_by
        destination_display = procurement_request.mobilization.destination_display
    else:
        origin_display = 'Stock replenishment (no linked request)'
        contact = procurement_request.requested_by
        destination_display = '—'

    # Split like the incident report does: photos get embedded so a vendor
    # reading the PDF standalone can actually see them — a login-gated
    # download link is useless to someone with no account on this system.
    # Anything else is just named, since it can't be inlined. Two photos
    # share each appended page (rather than one-per-page) purely to keep
    # the page count — and so the print cost — down.
    all_attachments = list(source_ticket.attachments.all()) if source_ticket else []
    image_attachments = [a for a in all_attachments if a.content_type.startswith('image/')]
    other_attachments = [a for a in all_attachments if a not in image_attachments]
    image_attachment_pairs = [image_attachments[i:i + 2] for i in range(0, len(image_attachments), 2)]

    # Sections 3 (background) and 4 (notes) are each only shown when there's
    # actually something to put in them — numbered on the fly so the ones
    # that do appear stay in unbroken sequence instead of e.g. jumping from
    # "Section 2" straight to "Section 5" when both are skipped.
    next_section = 3
    background_section = None
    if source_ticket:
        background_section = next_section
        next_section += 1
    notes_section = None
    if procurement_request.notes:
        notes_section = next_section
        next_section += 1
    documentation_section = next_section

    context = {
        'procurement_request': procurement_request,
        'source_ticket': source_ticket,
        'origin_display': origin_display,
        'contact': contact,
        'destination_display': destination_display,
        'attachments': other_attachments,
        'image_attachments': image_attachments,
        'image_attachment_pairs': image_attachment_pairs,
        'background_section': background_section,
        'notes_section': notes_section,
        'documentation_section': documentation_section,
    }
    return _render_form_pdf_chromium(
        'reports/procurement_request_form_pdf.html', context, request,
        f'procurement_request_{procurement_request.pk}'
    )


def export_maintenance_pdf(schedule, request):
    from .report_registry import maintenance_form_sections
    context = maintenance_form_sections(schedule)
    return _render_form_pdf_chromium('reports/maintenance_form_pdf.html', context, request, f'maintenance_{context["schedule_code"]}')


def export_maintenance_docx(schedule):
    from .report_registry import maintenance_form_sections
    ctx = maintenance_form_sections(schedule)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    _docx_letterhead(doc, 'HDG-IT-FRM-090', '1', timezone.now().strftime('%d %B %Y'), 'MAINTENANCE SCHEDULE REPORT', page='1 of 1')

    title = doc.add_heading(f'Maintenance Schedule — {ctx["schedule_code"]}', level=1)
    for run in title.runs:
        run.font.color.rgb = PRIMARY_RGB
    doc.add_paragraph(schedule.title).runs[0].italic = True

    schedule_fields = [
        ('Target Department(s)', ctx['department_display']),
        ('Status', ctx['status_display']),
        ('Scheduled Date', schedule.scheduled_date.strftime('%Y-%m-%d')),
    ]
    if schedule.start_time:
        schedule_fields.append(('Time', f'{schedule.start_time.strftime("%H:%M")}' + (f' – {schedule.end_time.strftime("%H:%M")}' if schedule.end_time else '')))
    schedule_fields.append(('Facility / Location', schedule.facility_location))
    if ctx['target_assets']:
        schedule_fields.append(('Target Asset(s)', ', '.join(f'{a.name} ({a.tracking_id})' for a in ctx['target_assets'])))
    if ctx['vendors']:
        schedule_fields.append(('Third-Party Vendor(s)', ', '.join(v.name for v in ctx['vendors'])))
    _docx_form_section(doc, 'Section 1 — Schedule Details', schedule_fields)

    _docx_form_section(doc, 'Section 2 — Description', [], body_paragraph=schedule.description or '—')

    _docx_heading(doc, 'Section 3 — Assigned Personnel')
    if ctx['assignees']:
        for person in ctx['assignees']:
            doc.add_paragraph(f'• {person.get_full_name() or person.email}')
    else:
        doc.add_paragraph('Unassigned').runs[0].italic = True

    _docx_heading(doc, f'Section 4 — Checklist ({ctx["checklist_progress"]}% complete)')
    if ctx['checklist']:
        for item, done in ctx['checklist']:
            doc.add_paragraph(f'{"☑" if done else "☐"} {item}')
    else:
        doc.add_paragraph('No checklist items.').runs[0].italic = True

    _docx_heading(doc, 'Section 5 — Timeline')
    timeline_rows = [
        ('Created', schedule.created_at.strftime('%Y-%m-%d %H:%M')),
        ('Started', ctx['started_at'].strftime('%Y-%m-%d %H:%M') if ctx['started_at'] else 'Not started'),
        ('Completed', schedule.completed_at.strftime('%Y-%m-%d %H:%M') if schedule.completed_at else 'Not completed'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for cell, label in zip(header_cells, ('Stage', 'Timestamp')):
        cell.text = ''
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
    for stage, timestamp in timeline_rows:
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(stage).bold = True
        row_cells[1].text = timestamp
    doc.add_paragraph()

    doc.add_page_break()
    multi_department = len(ctx['department_sections']) > 1
    _docx_heading(doc, 'Section 6 — Owner Confirmation by Department' if multi_department else 'Section 6 — Owner Confirmation')
    if ctx['department_sections']:
        for section in ctx['department_sections']:
            if multi_department:
                doc.add_paragraph(section['department_display']).runs[0].bold = True
            if section['confirmations']:
                conf_table = doc.add_table(rows=1, cols=4)
                conf_table.style = 'Table Grid'
                for cell, label in zip(conf_table.rows[0].cells, ('Asset', 'Status', 'Confirmed By', 'Confirmed At')):
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(label)
                    run.bold = True
                for c in section['confirmations']:
                    row_cells = conf_table.add_row().cells
                    row_cells[0].text = f'{c.asset.name} ({c.asset.tracking_id})'
                    row_cells[1].text = c.get_status_display()
                    row_cells[2].text = c.confirmed_by.get_full_name() if c.confirmed_by else '—'
                    row_cells[3].text = c.confirmed_at.strftime('%Y-%m-%d %H:%M') if c.confirmed_at else '—'
                doc.add_paragraph()
            else:
                doc.add_paragraph('No target assets in this department.').runs[0].italic = True
    else:
        doc.add_paragraph('No target assets recorded for this schedule.').runs[0].italic = True

    buf = BytesIO()
    doc.save(buf)
    response = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{_filename("maintenance_" + ctx["schedule_code"], include_timestamp=False)}.docx"'
    return response


def export_asset_pdf(asset, request):
    from .report_registry import asset_form_sections
    context = asset_form_sections(asset)
    return _render_form_pdf_chromium('reports/asset_form_pdf.html', context, request, f'asset_{asset.tracking_id}')


def export_mobilization_audit_pdf(mobilization, request):
    from .report_registry import mobilization_audit_sections
    context = mobilization_audit_sections(mobilization)
    return _render_form_pdf_chromium(
        'reports/mobilization_audit_pdf.html', context, request, f'mobilization_{mobilization.pk}_audit'
    )


def export_asset_docx(asset):
    from .report_registry import asset_form_sections
    ctx = asset_form_sections(asset)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    _docx_letterhead(doc, f'Asset {asset.tracking_id}', '1', timezone.now().strftime('%d %B %Y'), 'ASSET RECORD', page='1 of 1')

    title = doc.add_heading(f'Asset Record — {asset.tracking_id}', level=1)
    for run in title.runs:
        run.font.color.rgb = PRIMARY_RGB
    doc.add_paragraph(asset.name).runs[0].italic = True

    _docx_form_section(doc, 'Section 1 — Identification', [
        ('Tracking ID', asset.tracking_id),
        ('Name', asset.name),
        ('Category', asset.category.name if asset.category else None),
        ('Serial Number', asset.serial_number),
        ('Model', asset.model),
        ('Manufacturer', asset.manufacturer),
    ])

    custody_fields = [
        ('Status', asset.get_status_display()),
        ('Condition', asset.get_condition_display()),
        ('Location', asset.location.full_name() if asset.location else None),
        ('Department', asset.department.name if asset.department else None),
    ]
    if asset.is_consumable:
        custody_fields.append(('Quantity In Stock', asset.quantity_in_stock))
    else:
        custody_fields.append(('Assigned To', asset.assigned_to.get_full_name() if asset.assigned_to else None))
        if asset.checked_out_to_id:
            custody_fields.append(('Checked Out To', asset.checked_out_to.get_full_name()))
            custody_fields.append(('Checked Out At', asset.checked_out_at.strftime('%Y-%m-%d %H:%M') if asset.checked_out_at else None))
            custody_fields.append(('Expected Return', asset.expected_return_date.strftime('%Y-%m-%d') if asset.expected_return_date else None))
    _docx_form_section(doc, 'Section 2 — Status & Custody', custody_fields)

    lifecycle_fields = [
        ('Purchase Date', asset.purchase_date.strftime('%Y-%m-%d') if asset.purchase_date else None),
        ('Warranty Expiry', asset.warranty_expiry.strftime('%Y-%m-%d') if asset.warranty_expiry else None),
    ]
    if asset.warranty_provider:
        lifecycle_fields.append(('Warranty Provider', asset.warranty_provider))
    if asset.is_renewable:
        lifecycle_fields.append(('Next Renewal Date', asset.next_renewal_date.strftime('%Y-%m-%d') if asset.next_renewal_date else None))
        lifecycle_fields.append(('Renewal Vendor', asset.renewal_vendor.name if asset.renewal_vendor else None))
        lifecycle_fields.append(('Auto-Renews', 'Yes' if asset.auto_renews else 'No'))
    _docx_form_section(doc, 'Section 3 — Lifecycle', lifecycle_fields)

    if ctx['notes_section']:
        _docx_form_section(doc, f'Section {ctx["notes_section"]} — Notes', [], body_paragraph=asset.notes)

    _docx_heading(doc, f'Section {ctx["activity_section"]} — Recent Activity')
    _docx_simple_table(
        doc,
        ['Date', 'Action', 'By', 'Details'],
        [
            (
                log.created_at.strftime('%Y-%m-%d %H:%M'),
                log.get_action_display(),
                log.actor.get_full_name() if log.actor else 'System',
                log.get_details_display() or '—',
            )
            for log in ctx['recent_activity']
        ],
        placeholder='No activity recorded.',
    )

    buf = BytesIO()
    doc.save(buf)
    response = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{_filename("asset_" + asset.tracking_id, include_timestamp=False)}.docx"'
    return response


def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _weighted_column_widths(columns, rows, usable_width_in, min_in=0.75, max_in=2.3):
    """Word/LibreOffice both render docx tables with naive equal-width
    columns unless widths are stated explicitly — 'autofit' only kicks in
    once a human opens and nudges the table in Word itself. Estimate a
    reasonable width per column from header + a sample of cell content
    lengths, clip to a sane range, then rescale so the row still fills the
    page exactly."""
    sample_rows = rows[:200]
    weights = []
    for col in columns:
        content_len = max((len(str(r.get(col, ''))) for r in sample_rows), default=0)
        weights.append(max(len(col), content_len * 0.6, 6))
    raw = [usable_width_in * w / sum(weights) for w in weights]
    clipped = [min(max(w, min_in), max_in) for w in raw]
    scale = usable_width_in / sum(clipped)
    return [Inches(w * scale) for w in clipped]


def export_docx(rows, columns, title, filename_base, filter_summary='', control_number='', **kwargs):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(9.5)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = Cm(21.0)
    section.page_width = Cm(29.7)
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)
    usable_width_in = (section.page_width - section.left_margin - section.right_margin) / 914400

    generated_at = timezone.now()
    _docx_report_letterhead(doc, title, generated_at, filter_summary, len(rows), usable_width_in, control_number=control_number)

    table = doc.add_table(rows=1, cols=max(len(columns), 1))
    table.style = 'Table Grid'
    table.autofit = False

    widths = _weighted_column_widths(columns, rows, usable_width_in)

    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        header_cells[i].text = col
        header_cells[i].width = widths[i]
        _shade_cell(header_cells[i], '0D9488')
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)

    for row_data in rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row_data.get(col, '—'))
            cells[i].width = widths[i]

    for i, w in enumerate(widths):
        table.columns[i].width = w

    from apps.accounts.models import ClientSettings
    client_settings = ClientSettings.objects.first()
    company_name = client_settings.company_name if client_settings else 'HydroDive'

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f'Generated by {company_name} IT Service Management System — Confidential, for internal use only.'
    )
    footer_run.italic = True
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = META_GRAY_RGB

    buf = BytesIO()
    doc.save(buf)
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{_filename(filename_base)}.docx"'
    return response


EXPORTERS = {
    'csv': export_csv,
    'excel': export_excel,
    'json': export_json,
    'pdf': export_pdf,
    'docx': export_docx,
}


def build_export_response(report_config, request, export_format, filter_summary=''):
    queryset = report_config.get_queryset(request)
    # Bulk row selection: if the report table's checkboxes selected specific
    # records, export only those — every report type's get_queryset()
    # returns a plain QuerySet keyed on its model's default pk, so pk__in
    # works uniformly here without touching any individual report type.
    ids = request.GET.getlist('ids')
    if ids:
        queryset = queryset.filter(pk__in=ids)
    rows = [report_config.row_from_obj(obj) for obj in queryset]
    columns = report_config.columns

    # Column picker (list exports only — see components/export_menu.html):
    # `cols` is a comma-separated subset of report_config.columns chosen in
    # the modal. Applies to every flat/table export format here (CSV, Excel,
    # the generic PDF table, the generic DOCX table) — this function is only
    # ever reached from the bundled list export endpoint, never a single-
    # record letterhead form export, so there's no form layout to break.
    # Re-filtered against the real column list (never trust the querystring
    # as-is) and re-ordered to match it, so a tampered/garbled value just
    # falls back to every column rather than erroring.
    if export_format in ('csv', 'excel', 'pdf', 'docx') and request.GET.get('cols'):
        requested = set(c.strip() for c in request.GET['cols'].split(','))
        selected = [c for c in report_config.columns if c in requested]
        if selected:
            columns = selected

    exporter = EXPORTERS.get(export_format)
    if not exporter:
        return HttpResponse('Invalid format', status=400)

    return exporter(
        rows, columns, report_config.label, report_config.slug,
        request=request, filter_summary=filter_summary, control_number=report_config.control_number,
    )
